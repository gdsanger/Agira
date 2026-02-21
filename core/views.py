from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse, JsonResponse, Http404
from django.views.decorators.http import require_POST, require_http_methods
from django.core.exceptions import ValidationError
from django.core.paginator import Paginator
from django.db import models, transaction, IntegrityError
from django.db.models import Q, Count
from django.contrib.contenttypes.models import ContentType
from django.contrib.auth import authenticate, login as auth_login, logout as auth_logout
from django.contrib.auth.decorators import login_required
from django.utils import timezone
from django.urls import reverse
from django.conf import settings
from decimal import Decimal, InvalidOperation
from datetime import datetime, date, time
import logging
import os
import re
import json
import openai
from google import genai
from django.utils.safestring import mark_safe
import markdown
import bleach
from .models import (
    Project, Item, ItemStatus, ItemComment, User, Release, Node, ItemType, Organisation,
    Attachment, AttachmentLink, AttachmentRole, Activity, ProjectStatus, NodeType, ReleaseStatus,
    AIProvider, AIModel, AIProviderType, AIJobsHistory, UserOrganisation, UserRole,
    ExternalIssueMapping, ExternalIssueKind, Change, ChangeStatus, ChangeApproval, ApprovalStatus, RiskLevel, ReleaseType,
    MailTemplate, MailActionMapping, IssueOpenQuestion, IssueStandardAnswer, OpenQuestionStatus, OpenQuestionSource,
    GlobalSettings, SystemSetting, ChangePolicy, ChangePolicyRole)


from .services.workflow import ItemWorkflowGuard
from .services.activity import ActivityService
from .services.storage import AttachmentStorageService
from .services.storage.errors import AttachmentTooLarge
from .services.agents import AgentService
from .services.mail import check_mail_trigger, prepare_mail_preview
from .services.change_policy_service import ChangePolicyService
from .backends.azuread import AzureADAuth, AzureADAuthError

# Configure logging
logger = logging.getLogger(__name__)

# Create markdown parser once at module level for better performance
MARKDOWN_PARSER = markdown.Markdown(extensions=['extra', 'fenced_code'])

# Allowed HTML tags and attributes for sanitization
ALLOWED_TAGS = [
    'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
    'p', 'br', 'strong', 'em', 'u', 'strike',
    'ul', 'ol', 'li',
    'blockquote', 'code', 'pre',
    'a', 'img',
    'table', 'thead', 'tbody', 'tr', 'th', 'td',
    'div', 'span'
]

ALLOWED_ATTRIBUTES = {
    'a': ['href', 'title', 'target', 'rel'],
    'img': ['src', 'alt', 'title', 'width', 'height'],
    'code': ['class'],
    'pre': ['class'],
    'div': ['class'],
    'span': ['class'],
}

# RAG context constants
RAG_NO_CONTEXT_MESSAGE = "No additional context found."

def home(request):
    """Home page view."""
    return render(request, 'home.html')

def login_view(request):
    """Login page view."""
    # If user is already authenticated, redirect to dashboard
    if request.user.is_authenticated:
        return redirect('dashboard')
    
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        next_url = request.POST.get('next', 'dashboard')
        
        # Authenticate user
        user = authenticate(request, username=username, password=password)
        
        if user is not None:
            if user.active:
                # Login the user
                auth_login(request, user)
                # Redirect to next page or dashboard
                return redirect(next_url if next_url else 'dashboard')
            else:
                # User account is inactive
                error = 'Ihr Konto ist deaktiviert.'
                return render(request, 'login.html', {
                    'error': error, 
                    'next': next_url,
                    'azure_ad_enabled': settings.AZURE_AD_ENABLED
                })
        else:
            # Invalid credentials
            error = 'Benutzername oder Passwort ist falsch.'
            return render(request, 'login.html', {
                'error': error, 
                'next': next_url,
                'azure_ad_enabled': settings.AZURE_AD_ENABLED
            })
    
    # GET request - show login form
    next_url = request.GET.get('next', '')
    return render(request, 'login.html', {
        'next': next_url,
        'azure_ad_enabled': settings.AZURE_AD_ENABLED
    })

def logout_view(request):
    """Logout view with Azure AD single sign-out support."""
    # Check if user was logged in via Azure AD (has azure_ad_object_id)
    azure_ad_user = False
    if request.user.is_authenticated:
        azure_ad_user = hasattr(request.user, 'azure_ad_object_id') and request.user.azure_ad_object_id
    
    # Log out from Agira
    auth_logout(request)
    
    # If user was logged in via Azure AD and Azure AD is enabled, offer single logout
    if azure_ad_user and settings.AZURE_AD_ENABLED:
        # Redirect to Azure AD logout
        try:
            azure_ad = AzureADAuth()
            post_logout_uri = request.build_absolute_uri(reverse('login'))
            logout_url = azure_ad.get_logout_url(post_logout_uri)
            return redirect(logout_url)
        except AzureADAuthError as e:
            logger.warning(f"Azure AD logout failed: {str(e)}")
            # Continue with local logout page
        except Exception as e:
            logger.error(f"Unexpected error during Azure AD logout: {str(e)}", exc_info=True)
            # Continue with local logout page
    
    return render(request, 'logged_out.html')

@login_required
def dashboard(request):
    """Dashboard page view with KPIs and activity overview."""
    from datetime import timedelta, date
    from django.db.models.functions import TruncDate
    
    # Calculate KPIs
    kpis = {
        'inbox_count': Item.objects.filter(status=ItemStatus.INBOX).count(),
        'backlog_count': Item.objects.filter(status=ItemStatus.BACKLOG).count(),
        'in_progress_count': Item.objects.filter(
            status__in=[ItemStatus.WORKING, ItemStatus.TESTING, ItemStatus.READY_FOR_RELEASE]
        ).count(),
        'closed_7d_count': Item.objects.filter(
            status=ItemStatus.CLOSED,
            updated_at__gte=timezone.now() - timedelta(days=7)
        ).count(),
        'changes_open_count': Change.objects.exclude(status__in=[ChangeStatus.DEPLOYED, ChangeStatus.CANCELED]).count(),
        'ai_jobs_24h_count': AIJobsHistory.objects.filter(
            timestamp__gte=timezone.now() - timedelta(hours=24)
        ).count(),
    }
    
    # Calculate AI jobs cost (24h)
    ai_jobs_24h = AIJobsHistory.objects.filter(
        timestamp__gte=timezone.now() - timedelta(hours=24),
        costs__isnull=False
    ).aggregate(total_costs=models.Sum('costs'))
    kpis['ai_jobs_24h_costs'] = ai_jobs_24h['total_costs'] or Decimal('0')
    
    # Calculate closed items by day for the last 7 days
    now = timezone.now()
    today = timezone.localdate()
    days_ago_7 = today - timedelta(days=6)  # Include today = 7 days total
    
    # Get closed items grouped by date
    closed_by_day = Item.objects.filter(
        status=ItemStatus.CLOSED,
        updated_at__gte=timezone.make_aware(
            timezone.datetime.combine(days_ago_7, timezone.datetime.min.time())
        )
    ).annotate(
        date=TruncDate('updated_at')
    ).values('date').annotate(
        count=Count('id')
    ).order_by('date')
    
    # Create a dictionary for quick lookup
    closed_dict = {item['date']: item['count'] for item in closed_by_day}
    
    # Generate all 7 days with counts (0 if no items closed that day)
    closed_items_chart = []
    for i in range(6, -1, -1):  # 6 days ago to today
        day = today - timedelta(days=i)
        count = closed_dict.get(day, 0)
        closed_items_chart.append({
            'date': day.strftime('%Y-%m-%d'),
            'date_display': day.strftime('%d.%m'),  # Format for display
            'count': count
        })
    
    context = {
        'kpis': kpis,
        'closed_items_chart_json': mark_safe(json.dumps(closed_items_chart)),
    }
    return render(request, 'dashboard.html', context)

@login_required
def projects(request):
    """Projects page view."""
    projects_list = Project.objects.all()
    
    # Annotate with item counts by status
    projects_list = projects_list.annotate(
        inbox_count=Count('items', filter=Q(items__status=ItemStatus.INBOX)),
        backlog_count=Count('items', filter=Q(items__status=ItemStatus.BACKLOG)),
        working_count=Count('items', filter=Q(items__status=ItemStatus.WORKING)),
        testing_count=Count('items', filter=Q(items__status=ItemStatus.TESTING)),
        ready_for_release_count=Count('items', filter=Q(items__status=ItemStatus.READY_FOR_RELEASE))
    )
    
    # Server-side search filter
    q = request.GET.get('q', '')
    if q:
        projects_list = projects_list.filter(name__icontains=q)
    
    # Filter by organisation
    org_filter = request.GET.get('organisation', '')
    if org_filter:
        projects_list = projects_list.filter(clients__id=org_filter)
    
    # Filter by status
    status_filter = request.GET.get('status', '')
    if status_filter:
        projects_list = projects_list.filter(status=status_filter)
    
    # Get all organisations and statuses for filter dropdowns
    organisations = Organisation.objects.all().order_by('name')
    statuses = ProjectStatus.choices
    
    context = {
        'projects': projects_list,
        'search_query': q,
        'organisations': organisations,
        'statuses': statuses,
        'selected_organisation': org_filter,
        'selected_status': status_filter,
    }
    return render(request, 'projects.html', context)

@login_required
def project_create(request):
    """Project create page view."""
    if request.method == 'GET':
        # Show the create form
        statuses = ProjectStatus.choices
        context = {
            'project': None,
            'statuses': statuses,
        }
        return render(request, 'project_form.html', context)
    
    # Handle POST request (HTMX form submission)
    try:
        project = Project.objects.create(
            name=request.POST.get('name'),
            description=request.POST.get('description', ''),
            status=request.POST.get('status', ProjectStatus.NEW),
            github_owner=request.POST.get('github_owner', ''),
            github_repo=request.POST.get('github_repo', '')
        )
        return JsonResponse({
            'success': True,
            'message': 'Project created successfully',
            'project_id': project.id,
            'redirect': f'/projects/{project.id}/'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)

@login_required
def project_edit(request, id):
    """Project edit page view."""
    project = get_object_or_404(Project, id=id)
    
    if request.method == 'GET':
        # Show the edit form
        statuses = ProjectStatus.choices
        context = {
            'project': project,
            'statuses': statuses,
        }
        return render(request, 'project_form.html', context)
    
    # Handle POST request (HTMX form submission)
    try:
        project.name = request.POST.get('name', project.name)
        project.description = request.POST.get('description', project.description)
        project.status = request.POST.get('status', project.status)
        project.github_owner = request.POST.get('github_owner', project.github_owner)
        project.github_repo = request.POST.get('github_repo', project.github_repo)
        project.save()
        
        return JsonResponse({
            'success': True,
            'message': 'Project updated successfully',
            'redirect': f'/projects/{project.id}/'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)

@login_required
def project_detail(request, id):
    """Project detail page view."""
    project = get_object_or_404(
        Project.objects.annotate(
            inbox_count=Count('items', filter=Q(items__status=ItemStatus.INBOX)),
            backlog_count=Count('items', filter=Q(items__status=ItemStatus.BACKLOG)),
            working_count=Count('items', filter=Q(items__status=ItemStatus.WORKING)),
        ),
        id=id
    )
    
    # Render markdown description to HTML with sanitization
    description_html = None
    if project.description:
        # Reset parser state for clean conversion
        MARKDOWN_PARSER.reset()
        # Convert markdown to HTML
        html = MARKDOWN_PARSER.convert(project.description)
        # Sanitize HTML to prevent XSS attacks
        sanitized_html = bleach.clean(
            html,
            tags=ALLOWED_TAGS,
            attributes=ALLOWED_ATTRIBUTES,
            strip=True
        )
        description_html = mark_safe(sanitized_html)
    
    # Get all organisations for the client management
    all_organisations = Organisation.objects.all().order_by('name')
    
    # Get all item types for adding new items
    item_types = ItemType.objects.filter(is_active=True).order_by('name')
    
    # Get node types for adding new nodes
    node_types = NodeType.choices
    
    # Count attachments for this project
    content_type = ContentType.objects.get_for_model(Project)
    attachments_count = AttachmentLink.objects.filter(
        target_content_type=content_type,
        target_object_id=project.id,
        role=AttachmentRole.PROJECT_FILE,
        attachment__is_deleted=False
    ).count()
    
    context = {
        'project': project,
        'description_html': description_html,
        'all_organisations': all_organisations,
        'item_types': item_types,
        'node_types': node_types,
        'attachments_count': attachments_count,
    }
    return render(request, 'project_detail.html', context)

@login_required
def project_items_tab(request, id):
    """Project Items tab with pagination, filtering, and sorting."""
    project = get_object_or_404(Project, id=id)
    
    # Start with all items for this project
    items = Item.objects.filter(project=project).select_related(
        'type', 'assigned_to', 'solution_release'
    ).order_by('-updated_at')
    
    # Get filter parameters
    search_query = request.GET.get('q', '').strip()
    status_filter_type = request.GET.get('status_filter', 'not_closed')  # Default to not_closed
    type_filter = request.GET.getlist('type')
    
    # Convert type_filter to integers for proper comparison
    type_filter_ints = []
    for t in type_filter:
        try:
            type_filter_ints.append(int(t))
        except (ValueError, TypeError):
            pass
    
    # Apply search filter (title and description)
    if search_query:
        items = items.filter(
            Q(title__icontains=search_query) | Q(description__icontains=search_query)
        )
    
    # Apply status filter - simplified to closed vs not closed
    if status_filter_type == 'closed':
        items = items.filter(status=ItemStatus.CLOSED)
    else:  # not_closed (default)
        items = items.exclude(status=ItemStatus.CLOSED)
    
    # Apply type filter
    if type_filter_ints:
        items = items.filter(type_id__in=type_filter_ints)
    
    # Pagination - 25 items per page
    paginator = Paginator(items, 25)
    page_number = request.GET.get('page', 1)
    page_obj = paginator.get_page(page_number)
    
    # Get all item types for filter dropdown
    item_types = ItemType.objects.filter(is_active=True).order_by('name')
    
    # Get all status choices
    status_choices = ItemStatus.choices
    
    # Get project releases for solution_release dropdown
    project_releases = project.releases.all().order_by('-update_date', 'name')
    
    context = {
        'project': project,
        'page_obj': page_obj,
        'search_query': search_query,
        'status_filter_type': status_filter_type,
        'selected_types': type_filter_ints,
        'item_types': item_types,
        'status_choices': status_choices,
        'closed_status_value': ItemStatus.CLOSED,  # Pass the constant to template
        'project_releases': project_releases,
    }
    return render(request, 'partials/project_items_tab.html', context)

@login_required
def items_inbox(request):
    """Items Inbox page view."""
    items = Item.objects.filter(status=ItemStatus.INBOX).select_related(
        'project', 'type', 'organisation', 'requester', 'assigned_to'
    ).order_by('-created_at')
    
    context = {
        'items': items,
    }
    return render(request, 'items_inbox.html', context)

@login_required
def items_backlog(request):
    """Items Backlog page view."""
    items = Item.objects.filter(status=ItemStatus.BACKLOG).select_related(
        'project', 'type', 'organisation', 'requester', 'assigned_to'
    ).order_by('-updated_at')
    
    # Search filter
    q = request.GET.get('q', '')
    if q:
        items = items.filter(
            models.Q(title__icontains=q) | models.Q(description__icontains=q)
        )
    
    # Project filter
    project_id = request.GET.get('project', '')
    if project_id:
        try:
            items = items.filter(project_id=int(project_id))
        except (ValueError, TypeError):
            # Invalid project_id, ignore filter
            project_id = ''
    
    # Get all projects for filter dropdown
    projects = Project.objects.all().order_by('name')
    
    context = {
        'items': items,
        'search_query': q,
        'project_id': project_id,
        'projects': projects,
    }
    return render(request, 'items_backlog.html', context)

@login_required
def items_working(request):
    """Items Working page view."""
    items = Item.objects.filter(status=ItemStatus.WORKING).select_related(
        'project', 'type', 'organisation', 'requester', 'assigned_to'
    ).order_by('-updated_at')
    
    # Search filter
    q = request.GET.get('q', '')
    if q:
        items = items.filter(
            models.Q(title__icontains=q) | models.Q(description__icontains=q)
        )
    
    # Project filter
    project_id = request.GET.get('project', '')
    if project_id:
        try:
            items = items.filter(project_id=int(project_id))
        except (ValueError, TypeError):
            # Invalid project_id, ignore filter
            project_id = ''
    
    # Get all projects for filter dropdown
    projects = Project.objects.all().order_by('name')
    
    context = {
        'items': items,
        'search_query': q,
        'project_id': project_id,
        'projects': projects,
    }
    return render(request, 'items_working.html', context)

@login_required
def items_testing(request):
    """Items Testing page view."""
    items = Item.objects.filter(status=ItemStatus.TESTING).select_related(
        'project', 'type', 'organisation', 'requester', 'assigned_to'
    ).order_by('-updated_at')
    
    # Search filter
    q = request.GET.get('q', '')
    if q:
        items = items.filter(
            models.Q(title__icontains=q) | models.Q(description__icontains=q)
        )
    
    # Project filter
    project_id = request.GET.get('project', '')
    if project_id:
        try:
            items = items.filter(project_id=int(project_id))
        except (ValueError, TypeError):
            # Invalid project_id, ignore filter
            project_id = ''
    
    # Get all projects for filter dropdown
    projects = Project.objects.all().order_by('name')
    
    context = {
        'items': items,
        'search_query': q,
        'project_id': project_id,
        'projects': projects,
    }
    return render(request, 'items_testing.html', context)

@login_required
def items_ready(request):
    """Items Ready for Release page view."""
    items = Item.objects.filter(status=ItemStatus.READY_FOR_RELEASE).select_related(
        'project', 'type', 'organisation', 'requester', 'assigned_to'
    ).order_by('-updated_at')
    
    # Search filter
    q = request.GET.get('q', '')
    if q:
        items = items.filter(
            models.Q(title__icontains=q) | models.Q(description__icontains=q)
        )
    
    # Project filter
    project_id = request.GET.get('project', '')
    if project_id:
        try:
            items = items.filter(project_id=int(project_id))
        except (ValueError, TypeError):
            # Invalid project_id, ignore filter
            project_id = ''
    
    # Get all projects for filter dropdown
    projects = Project.objects.all().order_by('name')
    
    context = {
        'items': items,
        'search_query': q,
        'project_id': project_id,
        'projects': projects,
    }
    return render(request, 'items_ready.html', context)

def get_open_github_issues_count():
    """
    Get count of open GitHub issues linked to items with status Working or Testing.
    
    Returns:
        int: Count of open GitHub issues (excluding PRs, excluding closed issues)
    """
    return ExternalIssueMapping.objects.filter(
        item__status__in=[ItemStatus.WORKING, ItemStatus.TESTING],
        kind=ExternalIssueKind.ISSUE,
    ).exclude(
        state='closed'
    ).count()

@login_required
def items_github_open(request):
    """Open GitHub Issues page view - shows all open GitHub issues linked to Working/Testing items."""
    # Query ExternalIssueMapping directly to get all open issues from Working/Testing items
    # This avoids N+1 queries and correctly handles items with multiple mappings
    open_issue_mappings = ExternalIssueMapping.objects.filter(
        item__status__in=[ItemStatus.WORKING, ItemStatus.TESTING],
        kind=ExternalIssueKind.ISSUE,
    ).exclude(
        state='closed'
    ).select_related('item', 'item__project').prefetch_related('item__external_mappings').order_by('-number')
    
    # Build list of issue data for display
    issues_data = []
    for mapping in open_issue_mappings:
        # Get associated PRs for this item from prefetched data
        all_mappings = mapping.item.external_mappings.all()
        prs = [m for m in all_mappings if m.kind == ExternalIssueKind.PR]
        prs.sort(key=lambda x: x.number)  # Sort by PR number
        
        # Select PR according to rule: first non-merged, or first if all merged
        selected_pr = None
        if prs:
            # Try to find first non-merged PR (prefer open over closed)
            for pr in prs:
                if pr.state != 'merged':
                    selected_pr = pr
                    break
            # If all merged, use first PR
            if selected_pr is None:
                selected_pr = prs[0]
        
        # Build PR data dict
        pr_data = None
        if selected_pr:
            pr_data = {
                'number': selected_pr.number,
                'url': selected_pr.html_url,
                'state': selected_pr.state,
            }
        
        issues_data.append({
            'issue_number': mapping.number,
            'item_title': mapping.item.title,
            'item_id': mapping.item.id,
            'github_url': mapping.html_url,
            'pr': pr_data,
        })
    
    context = {
        'issues_data': issues_data,
    }
    return render(request, 'items_github_open.html', context)

@login_required
def changes(request):
    """Changes list page view."""
    changes_list = Change.objects.all().select_related(
        'project', 'created_by', 'release'
    ).prefetch_related('approvals__approver')
    
    # Server-side search filter
    q = request.GET.get('q', '')
    if q:
        changes_list = changes_list.filter(
            Q(title__icontains=q) | Q(description__icontains=q)
        )
    
    # Filter by project
    project_filter = request.GET.get('project', '')
    if project_filter:
        try:
            changes_list = changes_list.filter(project_id=int(project_filter))
        except (ValueError, TypeError):
            project_filter = ''
    
    # Filter by status
    status_filter = request.GET.get('status', '')
    if status_filter:
        changes_list = changes_list.filter(status=status_filter)
    
    # Filter by risk level
    risk_filter = request.GET.get('risk', '')
    if risk_filter:
        changes_list = changes_list.filter(risk=risk_filter)
    
    # Annotate with approval counts
    changes_list = changes_list.annotate(
        total_approvals=Count('approvals'),
        approved_count=Count('approvals', filter=Q(approvals__status=ApprovalStatus.ACCEPT)),
        pending_count=Count('approvals', filter=Q(approvals__status=ApprovalStatus.PENDING)),
        rejected_count=Count('approvals', filter=Q(approvals__status=ApprovalStatus.REJECT))
    )
    
    # Get all projects, statuses and risk levels for filter dropdowns
    projects = Project.objects.all().order_by('name')
    statuses = ChangeStatus.choices
    risk_levels = RiskLevel.choices
    
    context = {
        'changes': changes_list,
        'search_query': q,
        'projects': projects,
        'statuses': statuses,
        'risk_levels': risk_levels,
        'selected_project': project_filter,
        'selected_status': status_filter,
        'selected_risk': risk_filter,
    }
    return render(request, 'changes.html', context)

@login_required
def item_lookup(request, item_id):
    """
    Lightweight endpoint to check if an issue exists and is accessible.
    Returns JSON with status information.
    Used by the header Issue ID search feature.
    
    Note: This view currently follows the same authorization model as item_detail,
    which allows any authenticated user to access any item. If more granular
    permissions are implemented in the future, they should be added here as well.
    """
    try:
        # Use select_related to minimize queries, similar to item_detail
        item = Item.objects.select_related('project').get(id=item_id)
        return JsonResponse({
            'exists': True,
            'id': item.id,
        })
    except Item.DoesNotExist:
        return JsonResponse({
            'exists': False,
            'error': 'Es existiert kein Issue mit dieser ID.',
        }, status=404)

@login_required
def item_status(request, item_id):
    """
    HTMX endpoint: Returns the current status of an item as HTML fragment.
    Used for periodic status updates in the sidebar Recent Items.
    
    Authorization: Requires login (same as item_detail view).
    Currently, there are no additional item-level permissions in the system.
    If item-level visibility/permission logic is added in the future,
    it should be applied here as well.
    
    Returns: 
        - 200 with HTML fragment containing the status badge if item exists
        - 204 No Content if item does not exist (HTMX-friendly, no error page)
    """
    try:
        item = Item.objects.get(id=item_id)
    except Item.DoesNotExist:
        # Return 204 No Content for non-existent items
        # This is HTMX-friendly and prevents 404 error pages in the sidebar
        return HttpResponse(status=204)
    
    return render(request, 'partials/item_status_badge.html', {
        'item': item
    })

@login_required
def item_detail(request, item_id):
    """Item detail page with tabs."""
    item = get_object_or_404(
        Item.objects.select_related(
            'project', 'type', 'organisation', 'requester', 
            'assigned_to', 'responsible', 'solution_release', 'parent'
        ).prefetch_related('nodes'),
        id=item_id
    )
    
    # Get followers for this item
    followers = item.get_followers()
    
    # Get all users for the follower selection dropdown
    users = User.objects.all().order_by('name')
    
    # Get agents for responsible field
    agents = User.objects.filter(role=UserRole.AGENT).order_by('name')
    
    # Get all projects for the move modal
    projects = Project.objects.all().order_by('name')
    
    # Get releases for the inline edit (filtered by project and exclude Closed)
    releases = Release.objects.filter(
        project=item.project
    ).exclude(
        status=ReleaseStatus.CLOSED
    ).order_by('-version')
    
    # Get parent items for the inline edit (exclude closed and self)
    # Filter as per issue #352 - allow items from all projects, status != closed
    parent_items = Item.objects.exclude(
        status=ItemStatus.CLOSED
    ).exclude(
        id=item.id
    ).order_by('title')
    
    # Get requester's primary organisation short code
    requester_org_short = None
    if item.requester:
        try:
            primary_org = UserOrganisation.objects.select_related('organisation').get(
                user=item.requester,
                is_primary=True
            )
            requester_org_short = primary_org.organisation.short or None
        except UserOrganisation.DoesNotExist:
            requester_org_short = None
    
    # Get all organisations for quick-create user modal
    organisations = Organisation.objects.all().order_by('name')
    
    # Get initial tab from query parameter (default: overview)
    active_tab = request.GET.get('tab', 'overview')
    
    context = {
        'item': item,
        'followers': followers,
        'users': users,
        'agents': agents,
        'projects': projects,
        'releases': releases,
        'parent_items': parent_items,
        'requester_org_short': requester_org_short,
        'organisations': organisations,
        'active_tab': active_tab,
        'available_statuses': ItemStatus.choices,
    }
    return render(request, 'item_detail.html', context)


@login_required
@require_http_methods(["POST"])
def item_update_release(request, item_id):
    """HTMX endpoint to update item release field."""
    item = get_object_or_404(Item, id=item_id)
    
    # Track old value for activity log
    old_release = item.solution_release
    old_value = old_release.version if old_release else 'None'
    
    release_id = request.POST.get('solution_release')
    
    try:
        if release_id:
            release = get_object_or_404(Release, id=release_id)
            item.solution_release = release
            new_value = release.version
        else:
            item.solution_release = None
            new_value = 'None'
        
        item.save()
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='item.field_changed',
            target=item,
            actor=request.user,
            summary=f'Changed solution_release from {old_value} to {new_value}'
        )
        
        return HttpResponse(status=200)
    except Exception as e:
        return HttpResponse(status=400)


@login_required
@require_http_methods(["POST"])
def item_update_parent(request, item_id):
    """HTMX endpoint to update item parent field."""
    item = get_object_or_404(Item, id=item_id)
    
    # Track old value for activity log
    old_parent = item.parent
    old_value = old_parent.title if old_parent else 'None'
    
    parent_id = request.POST.get('parent_item')
    
    try:
        if parent_id:
            parent_item = get_object_or_404(Item, id=parent_id)
            
            # Validate parent item criteria
            # 1. Status must not be closed
            if parent_item.status == ItemStatus.CLOSED:
                return HttpResponse('Cannot set a closed item as parent.', status=400)
            
            # 2. Cannot be the item itself
            if parent_item.id == item.id:
                return HttpResponse('Cannot set item as its own parent.', status=400)
            
            item.parent = parent_item
            new_value = parent_item.title
        else:
            item.parent = None
            new_value = 'None'
        
        item.save()
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='item.field_changed',
            target=item,
            actor=request.user,
            summary=f'Changed parent_item from {old_value} to {new_value}'
        )
        
        return HttpResponse(status=200)
    except Exception as e:
        logger.error(f"Error updating item parent: {str(e)}", exc_info=True)
        return HttpResponse('An error occurred while updating the parent item.', status=400)


@login_required
@require_http_methods(["POST"])
def item_update_intern(request, item_id):
    """HTMX endpoint to update item intern field."""
    item = get_object_or_404(Item, id=item_id)
    
    # Track old value for activity log
    old_value = item.intern
    
    # Get the new value from POST data
    intern_value = request.POST.get('intern')
    
    try:
        # Convert to boolean (checkbox sends 'on' when checked, nothing when unchecked)
        item.intern = intern_value == 'on' or intern_value == 'true'
        new_value = item.intern
        
        item.save()
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='item.field_changed',
            target=item,
            actor=request.user,
            summary=f'Changed intern from {old_value} to {new_value}'
        )
        
        return HttpResponse(status=200)
    except Exception as e:
        return HttpResponse(status=400)


@login_required
@require_http_methods(["POST"])
def item_take_over_responsible(request, item_id):
    """
    Take over action - sets responsible to current agent user.
    Only available for users with Agent role.
    Sends email notification only if responsible actually changes.
    """
    item = get_object_or_404(Item, id=item_id)
    
    # Check if current user is an agent
    if request.user.role != UserRole.AGENT:
        return JsonResponse({
            'success': False,
            'error': 'Only users with Agent role can take over responsibility.'
        }, status=403)
    
    # Check if responsible is already set to current user
    if item.responsible == request.user:
        return JsonResponse({
            'success': True,
            'message': 'You are already the responsible person for this item.',
            'no_change': True
        })
    
    try:
        # Store old responsible for logging
        old_responsible = item.responsible
        
        # Set responsible to current user
        item.responsible = request.user
        item.save()
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='item.responsible_changed',
            target=item,
            actor=request.user,
            summary=f'Took over responsibility (was: {old_responsible.name if old_responsible else "None"})'
        )
        
        # Send email notification (only if responsible changed)
        _send_responsible_notification(item, request.user)
        
        return JsonResponse({
            'success': True,
            'message': f'Responsibility taken over by {request.user.name}',
            'responsible_name': request.user.name
        })
    except Exception as e:
        logger.error(f"Error in take over responsible: {e}")
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)


@login_required
@require_http_methods(["POST"])
def item_assign_responsible(request, item_id):
    """
    Assign action - sets responsible to selected agent user.
    Only agents can be selected.
    Sends email notification only if responsible actually changes.
    """
    item = get_object_or_404(Item, id=item_id)
    
    # Get the selected agent user ID
    agent_id = request.POST.get('agent_id')
    
    if not agent_id:
        return JsonResponse({
            'success': False,
            'error': 'No agent selected.'
        }, status=400)
    
    try:
        # Get the agent user
        agent = get_object_or_404(User, id=agent_id)
        
        # Validate that user is an agent
        if agent.role != UserRole.AGENT:
            return JsonResponse({
                'success': False,
                'error': 'Selected user must have Agent role.'
            }, status=400)
        
        # Check if responsible is already set to this agent
        if item.responsible == agent:
            return JsonResponse({
                'success': True,
                'message': f'{agent.name} is already the responsible person for this item.',
                'no_change': True
            })
        
        # Store old responsible for logging
        old_responsible = item.responsible
        
        # Set responsible to selected agent
        item.responsible = agent
        item.save()
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='item.responsible_changed',
            target=item,
            actor=request.user,
            summary=f'Assigned responsibility to {agent.name} (was: {old_responsible.name if old_responsible else "None"})'
        )
        
        # Send email notification (only if responsible changed)
        _send_responsible_notification(item, agent)
        
        return JsonResponse({
            'success': True,
            'message': f'Responsibility assigned to {agent.name}',
            'responsible_name': agent.name
        })
    except Exception as e:
        logger.error(f"Error in assign responsible: {e}")
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)


@login_required
@require_http_methods(["POST"])
def item_quick_create_user(request, item_id):
    """Quick create a user and set as requester for an item."""
    item = get_object_or_404(Item, id=item_id)
    
    # Get form data
    name = request.POST.get('name', '').strip()
    email = request.POST.get('email', '').strip()
    organization_id = request.POST.get('organization_id', '').strip()
    
    # Validate required fields
    if not name:
        return JsonResponse({'success': False, 'error': 'Name is required'}, status=400)
    if not email:
        return JsonResponse({'success': False, 'error': 'Email is required'}, status=400)
    if not organization_id:
        return JsonResponse({'success': False, 'error': 'Organization is required'}, status=400)
    
    try:
        # Get organization
        organization = get_object_or_404(Organisation, id=organization_id)
        
        # Generate username from email
        username = email.split('@')[0]
        
        # Check if username already exists, if so, append a number
        base_username = username
        counter = 1
        while User.objects.filter(username=username).exists():
            username = f"{base_username}{counter}"
            counter += 1
        
        # Check if email already exists
        if User.objects.filter(email=email).exists():
            return JsonResponse({'success': False, 'error': f'A user with email {email} already exists'}, status=400)
        
        # Create user with transaction to ensure atomicity
        with transaction.atomic():
            # Create the user
            user = User.objects.create(
                username=username,
                email=email,
                name=name,
                role=UserRole.USER,
                active=True
            )
            
            # Set a default password (user will need to reset it)
            user.set_unusable_password()
            user.save()
            
            # Create primary organization relationship
            UserOrganisation.objects.create(
                user=user,
                organisation=organization,
                role=UserRole.USER,
                is_primary=True
            )
            
            # Set the user as requester for the item
            old_requester = item.requester.name if item.requester else 'None'
            item.requester = user
            item.save()
            
            # Log activity
            activity_service = ActivityService()
            activity_service.log(
                verb='item.field_changed',
                target=item,
                actor=request.user,
                summary=f'Changed requester from {old_requester} to {user.name} (quick created)'
            )
            
            return JsonResponse({
                'success': True,
                'message': f'User {name} created successfully and set as requester',
                'user_id': user.id
            })
    except Exception as e:
        logger.error(f'Error creating user: {str(e)}')
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required
def item_comments_tab(request, item_id):
    """HTMX endpoint to load comments tab."""
    item = get_object_or_404(Item, id=item_id)
    comments = item.comments.select_related('author').order_by('-created_at')
    
    context = {
        'item': item,
        'comments': comments,
    }
    return render(request, 'partials/item_comments_tab.html', context)


@login_required
def item_attachments_tab(request, item_id):
    """HTMX endpoint to load attachments tab."""
    item = get_object_or_404(Item, id=item_id)
    
    # Get query parameters for filtering
    search_query = request.GET.get('search', '').strip()
    file_type_filter = request.GET.get('file_type', '').strip()
    
    # Get attachments linked to this item
    content_type = ContentType.objects.get_for_model(Item)
    attachment_links = AttachmentLink.objects.filter(
        target_content_type=content_type,
        target_object_id=item.id,
        role=AttachmentRole.ITEM_FILE
    ).select_related('attachment', 'attachment__created_by').order_by('-created_at')
    
    attachments = [link.attachment for link in attachment_links if not link.attachment.is_deleted]
    
    # Apply filename search filter
    if search_query:
        attachments = [a for a in attachments if search_query.lower() in a.original_name.lower()]
    
    # Apply file type filter
    if file_type_filter:
        attachments = [a for a in attachments if a.file_type == file_type_filter]
    
    # Get distinct file types for filter dropdown
    all_attachments = [link.attachment for link in attachment_links if not link.attachment.is_deleted]
    file_types = sorted(set(a.file_type for a in all_attachments if a.file_type))
    
    context = {
        'item': item,
        'attachments': attachments,
        'file_types': file_types,
        'search_query': search_query,
        'file_type_filter': file_type_filter,
    }
    return render(request, 'partials/item_attachments_tab.html', context)


@login_required
def item_activity_tab(request, item_id):
    """HTMX endpoint to load activity tab."""
    item = get_object_or_404(Item, id=item_id)
    
    # Get activities for this item
    activity_service = ActivityService()
    activities = activity_service.latest(limit=100, item=item)
    
    context = {
        'item': item,
        'activities': activities,
    }
    return render(request, 'partials/item_activity_tab.html', context)


@login_required
def item_github_tab(request, item_id):
    """HTMX endpoint to load GitHub tab."""
    from core.services.github.service import GitHubService
    
    item = get_object_or_404(Item, id=item_id)
    external_mappings = item.external_mappings.all().order_by('-last_synced_at')
    
    # Check if item can have a GitHub issue created
    github_service = GitHubService()
    can_create_issue = github_service.can_create_issue_for_item(item)
    has_existing_issue = item.external_mappings.filter(kind='Issue').exists()
    
    context = {
        'item': item,
        'external_mappings': external_mappings,
        'can_create_issue': can_create_issue,
        'has_existing_issue': has_existing_issue,
    }
    return render(request, 'partials/item_github_tab.html', context)


@login_required
def item_related_items_tab(request, item_id):
    """HTMX endpoint to load related items tab."""
    from django_tables2 import RequestConfig
    from .tables import RelatedItemsTable
    from .filters import RelatedItemsFilter
    from .models import ItemRelation, RelationType
    
    item = get_object_or_404(Item, id=item_id)
    
    # Get related items where current item is the parent (from_item)
    # and relation_type is 'Related'
    related_item_ids = ItemRelation.objects.filter(
        from_item=item,
        relation_type=RelationType.RELATED
    ).values_list('to_item_id', flat=True)
    
    # Get the actual items
    queryset = Item.objects.filter(
        id__in=related_item_ids
    ).select_related(
        'project', 'type', 'assigned_to'
    )
    
    # Apply filters
    filterset = RelatedItemsFilter(request.GET, queryset=queryset)
    
    # Create table
    table = RelatedItemsTable(filterset.qs)
    RequestConfig(request, paginate={'per_page': 25}).configure(table)
    
    # Get all items in same project for relation creation (excluding current item)
    available_items = Item.objects.filter(
        project=item.project
    ).exclude(id=item.id).order_by('title')
    
    # Get existing relations for this item (all types, not just Related)
    existing_relations = ItemRelation.objects.filter(
        from_item=item
    ).select_related('to_item').order_by('-id')
    
    context = {
        'item': item,
        'table': table,
        'filter': filterset,
        'available_items': available_items,
        'existing_relations': existing_relations,
        'relation_types': RelationType.choices,
    }
    return render(request, 'partials/item_related_items_tab.html', context)


@login_required
@require_POST
def item_relation_create(request, item_id):
    """Create a new item relation."""
    from .models import ItemRelation, RelationType
    
    item = get_object_or_404(Item, id=item_id)
    
    to_item_id = request.POST.get('to_item')
    relation_type = request.POST.get('relation_type')
    
    try:
        to_item = get_object_or_404(Item, id=to_item_id)
        
        # Validate that items are in the same project
        if to_item.project != item.project:
            return JsonResponse({
                'status': 'error',
                'message': 'Items must be in the same project'
            }, status=400)
        
        # Create the relation
        ItemRelation.objects.create(
            from_item=item,
            to_item=to_item,
            relation_type=relation_type
        )
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='item.relation.created',
            target=item,
            actor=request.user,
            summary=f'Created {relation_type} relation to {to_item.title}'
        )
        
        return JsonResponse({'status': 'ok'})
        
    except IntegrityError:
        return JsonResponse({
            'status': 'error',
            'message': 'This relation already exists'
        }, status=400)
    except Exception as e:
        logger.error(f"Error creating item relation: {str(e)}")
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)


@login_required
@require_POST
def item_relation_update(request, item_id, relation_id):
    """Update an existing item relation."""
    from .models import ItemRelation
    
    item = get_object_or_404(Item, id=item_id)
    relation = get_object_or_404(ItemRelation, id=relation_id, from_item=item)
    
    to_item_id = request.POST.get('to_item')
    relation_type = request.POST.get('relation_type')
    
    try:
        old_to_item = relation.to_item
        old_relation_type = relation.relation_type
        
        if to_item_id:
            to_item = get_object_or_404(Item, id=to_item_id)
            if to_item.project != item.project:
                return JsonResponse({
                    'status': 'error',
                    'message': 'Items must be in the same project'
                }, status=400)
            relation.to_item = to_item
        
        if relation_type:
            relation.relation_type = relation_type
        
        relation.save()
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='item.relation.updated',
            target=item,
            actor=request.user,
            summary=f'Updated relation to {relation.to_item.title}: {old_relation_type} → {relation.relation_type}'
        )
        
        return JsonResponse({'status': 'ok'})
        
    except IntegrityError:
        return JsonResponse({
            'status': 'error',
            'message': 'This relation already exists'
        }, status=400)
    except Exception as e:
        logger.error(f"Error updating item relation: {str(e)}")
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)


@login_required
@require_POST
def item_relation_delete(request, item_id, relation_id):
    """Delete an item relation."""
    from .models import ItemRelation
    
    item = get_object_or_404(Item, id=item_id)
    relation = get_object_or_404(ItemRelation, id=relation_id, from_item=item)
    
    try:
        to_item_title = relation.to_item.title
        relation_type = relation.relation_type
        
        relation.delete()
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='item.relation.deleted',
            target=item,
            actor=request.user,
            summary=f'Deleted {relation_type} relation to {to_item_title}'
        )
        
        return JsonResponse({'status': 'ok'})
        
    except Exception as e:
        logger.error(f"Error deleting item relation: {str(e)}")
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)


@login_required

@require_POST
def item_link_github(request, item_id):
    """Link a GitHub Issue or Pull Request to an item."""
    from core.services.github.service import GitHubService
    from core.services.github.client import GitHubClient
    
    item = get_object_or_404(Item, id=item_id)
    
    try:
        # Get form data
        github_type = request.POST.get('type', 'Issue')
        number = request.POST.get('number', '').strip()
        
        if not number:
            return HttpResponse("Issue/PR number is required", status=400)
        
        try:
            number = int(number)
        except ValueError:
            return HttpResponse("Issue/PR number must be a valid integer", status=400)
        
        # Initialize GitHub service
        github_service = GitHubService()
        
        if not github_service.is_enabled() or not github_service.is_configured():
            return HttpResponse("GitHub integration is not configured", status=400)
        
        # Get owner and repo from project configuration using service method
        try:
            owner, repo = github_service._get_repo_info(item)
        except ValueError as e:
            return HttpResponse(str(e), status=400)
        
        # Get GitHub client
        client = github_service._get_client()
        
        # Fetch issue or PR data from GitHub
        if github_type == 'Issue':
            github_data = client.get_issue(owner, repo, number)
            kind = ExternalIssueKind.ISSUE
        elif github_type == 'PR':
            github_data = client.get_pr(owner, repo, number)
            kind = ExternalIssueKind.PR
        else:
            return HttpResponse("Invalid type. Must be 'Issue' or 'PR'", status=400)
        
        # Check if mapping already exists
        github_id = github_data['id']
        existing_mapping = ExternalIssueMapping.objects.filter(github_id=github_id).first()
        
        if existing_mapping:
            if existing_mapping.item.id == item.id:
                return HttpResponse("This GitHub item is already linked to this item", status=400)
            else:
                return HttpResponse(f"This GitHub item is already linked to another item: {existing_mapping.item.title}", status=400)
        
        # Create new mapping
        mapping = ExternalIssueMapping.objects.create(
            item=item,
            github_id=github_id,
            number=number,
            kind=kind,
            state=github_data.get('state', 'open'),
            html_url=github_data.get('html_url', f"https://github.com/{owner}/{repo}/issues/{number}"),
        )
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='github.linked',
            target=item,
            actor=request.user if request.user.is_authenticated else None,
            summary=f"Linked GitHub {dict(ExternalIssueKind.choices)[kind]} #{number}",
        )
        
        # Return updated GitHub tab
        external_mappings = item.external_mappings.all().order_by('-last_synced_at')
        context = {
            'item': item,
            'external_mappings': external_mappings,
        }
        response = render(request, 'partials/item_github_tab.html', context)
        response['HX-Trigger'] = 'githubLinked'
        return response
        
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"GitHub linking failed for item {item_id}: {str(e)}")
        return HttpResponse(f"Failed to link GitHub item: {str(e)}", status=500)


@login_required

@require_POST
def item_create_github_issue(request, item_id):
    """Create a new GitHub issue for an item."""
    from core.services.github.service import GitHubService
    from core.services.integrations.base import IntegrationError
    from core.services.mail import check_mail_trigger, prepare_mail_preview
    
    item = get_object_or_404(Item, id=item_id)
    
    try:
        # Initialize GitHub service
        github_service = GitHubService()
        
        # Check if GitHub is enabled and configured
        if not github_service.is_enabled():
            return HttpResponse("GitHub integration is not enabled. Please enable it in GitHub Configuration.", status=400)
        
        if not github_service.is_configured():
            return HttpResponse("GitHub integration is not configured. Please add a GitHub token in GitHub Configuration.", status=400)
        
        # Check if item has valid status
        if not github_service.can_create_issue_for_item(item):
            return HttpResponse(
                f"Cannot create GitHub issue for item with status '{item.status}'. "
                f"Item must have status 'Backlog', 'Working', or 'Testing'.",
                status=400
            )
        
        # Check if this is a follow-up issue (item already has issues)
        existing_issues = item.external_mappings.filter(kind='Issue').exists()
        
        # Initialize issue_body to None (default: use full description from item)
        issue_body = None
        
        if existing_issues:
            # For follow-up issues, get the notes from the request
            notes = request.POST.get('notes', '').strip()
            
            if not notes:
                return HttpResponse("Notes are required for creating a follow-up issue.", status=400)
            
            # Check if user wants to send minimal description to GitHub
            send_minimal = request.POST.get('send_minimal_description') == 'true'
            
            # Update item description BEFORE creating GitHub issue
            # This ensures the notes and existing issue references are included in the GitHub issue body
            _append_followup_notes_to_item(item, notes)
            
            # Prepare the issue body based on user's choice
            if send_minimal:
                # Create minimal description with only references and notes
                issue_body = _create_minimal_issue_description(item, notes)
        
        # Store old status to detect changes
        old_status = item.status
        
        # Create GitHub issue (this will also change status to WORKING if applicable)
        try:
            mapping = github_service.create_issue_for_item(
                item=item,
                actor=request.user,
                body=issue_body
            )
            
            # For follow-up issues, update the references section to include the newly created issue
            if existing_issues:
                _update_issue_references(item)
            
            # Check if status changed and if mail trigger exists
            # Refresh item to get the latest status
            item.refresh_from_db()
            status_changed = (old_status != item.status)
            
            # Check for mail trigger if status changed
            if status_changed:
                mapping = check_mail_trigger(item)
                if mapping:
                    # Prepare mail preview for modal
                    mail_preview = prepare_mail_preview(item, mapping)
                    
                    # Return JSON response with mail preview and updated tab HTML
                    external_mappings = item.external_mappings.all().order_by('-last_synced_at')
                    can_create_issue = github_service.can_create_issue_for_item(item)
                    has_existing_issue = item.external_mappings.filter(kind='Issue').exists()
                    
                    context = {
                        'item': item,
                        'external_mappings': external_mappings,
                        'can_create_issue': can_create_issue,
                        'has_existing_issue': has_existing_issue,
                    }
                    
                    # Render the GitHub tab to HTML string
                    from django.template.loader import render_to_string
                    github_tab_html = render_to_string('partials/item_github_tab.html', context, request=request)
                    
                    # Return JSON with both mail preview and tab HTML
                    return JsonResponse({
                        'success': True,
                        'mail_preview': mail_preview,
                        'github_tab_html': github_tab_html,
                        'item_id': item.id
                    })
            
            # No mail trigger or no status change - return updated GitHub tab
            external_mappings = item.external_mappings.all().order_by('-last_synced_at')
            can_create_issue = github_service.can_create_issue_for_item(item)
            has_existing_issue = item.external_mappings.filter(kind='Issue').exists()
            
            context = {
                'item': item,
                'external_mappings': external_mappings,
                'can_create_issue': can_create_issue,
                'has_existing_issue': has_existing_issue,
            }
            response = render(request, 'partials/item_github_tab.html', context)
            response['HX-Trigger'] = 'githubIssueCreated'
            return response
            
        except ValueError as e:
            return HttpResponse(f"Error creating GitHub issue: {str(e)}", status=400)
        except IntegrationError as e:
            return HttpResponse(f"GitHub API error: {str(e)}", status=500)
            
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"GitHub issue creation failed for item {item_id}: {str(e)}")
        return HttpResponse(f"Failed to create GitHub issue: {str(e)}", status=500)


def _create_minimal_issue_description(item, notes):
    """
    Create a minimal issue description with only references to original issues/PRs and the notes.
    
    This is used when the user chooses to send minimal description to GitHub instead of the full item description.
    The format follows the example from the issue: "Betrifft gdsanger/Agira#345 und gdsanger/Agira#456"
    
    Args:
        item: Item instance
        notes: User-provided notes for the follow-up
        
    Returns:
        str: Minimal description formatted for GitHub
        
    Raises:
        ValueError: If project doesn't have GitHub owner/repo configured
    """
    # Get existing issue and PR mappings (excluding the one we're about to create)
    mappings = item.external_mappings.all().order_by('kind', 'number')
    
    # Get GitHub owner and repo from project (must be configured at this point)
    owner = item.project.github_owner
    repo = item.project.github_repo
    
    if not owner or not repo:
        raise ValueError(
            f"Project '{item.project.name}' does not have GitHub repository configured. "
            f"Cannot create minimal description without valid owner/repo."
        )
    
    if mappings.exists():
        # Format: "Betrifft owner/repo#123 und owner/repo#456"
        refs = [f"{owner}/{repo}#{m.number}" for m in mappings]
        if len(refs) == 1:
            ref_line = f"Betrifft {refs[0]}"
        else:
            # Join all but last with ", " and last with " und "
            ref_line = f"Betrifft {', '.join(refs[:-1])} und {refs[-1]}"
    else:
        ref_line = ""
    
    # Build minimal description
    parts = []
    if ref_line:
        parts.append(ref_line)
    parts.append(notes)
    
    # Add metadata (same as default behavior)
    parts.append(f"\n---\n**Agira Item ID:** {item.id}")
    parts.append(f"**Project:** {item.project.name}")
    parts.append(f"**Type:** {item.type.name}")
    
    return '\n\n'.join(parts)


def _append_followup_notes_to_item(item, notes):
    """
    Append follow-up notes and issue/PR references to item description.
    
    This function is called BEFORE the new GitHub issue is created to ensure
    the notes and existing issue references are included in the GitHub issue body.
    After creation, _update_issue_references() is called to add the new issue number.
    
    Args:
        item: Item instance
        notes: User-provided notes for the follow-up
    """
    from datetime import datetime
    
    # Get current date and time formatted for German locale
    current_datetime = datetime.now().strftime("%d.%m.%Y %H:%M")
    
    # Get existing issue and PR numbers
    mappings = item.external_mappings.all().order_by('kind', 'number')
    issue_pr_refs = ", ".join([f"#{m.number}" for m in mappings])
    
    # Build the addition to description
    addition_parts = []
    
    # Initialize description if it's None or empty
    if not item.description:
        item.description = ""
    
    # If description doesn't have original header, add it
    if item.description and not item.description.startswith("## "):
        # Preserve original description with header
        original_desc = item.description
        item.description = f"## Original Item Issue Text\n{original_desc}"
    
    # Add notes section with date and time
    addition_parts.append(f"\n\n## Hinweise und Änderungen {current_datetime}")
    addition_parts.append(notes)
    
    # Add issue/PR references section
    if issue_pr_refs:
        addition_parts.append("\n### Siehe folgende Issues und PRs")
        addition_parts.append(issue_pr_refs)
    
    # Append to description
    item.description += "\n".join(addition_parts)
    item.save(update_fields=['description'])


def _update_issue_references(item):
    """
    Update the issue/PR references section in the item description to include
    all current mappings (including newly created ones).
    
    This is called AFTER a new GitHub issue has been created to add its number
    to the references list.
    
    Args:
        item: Item instance
    """
    import re
    import logging
    
    logger = logging.getLogger(__name__)
    
    # Refresh item to get the latest description
    item.refresh_from_db()
    
    # Get all issue and PR numbers
    mappings = item.external_mappings.all().order_by('kind', 'number')
    issue_pr_refs = ", ".join([f"#{m.number}" for m in mappings])
    
    if not issue_pr_refs:
        logger.warning(f"No issue/PR references found for item {item.id}, but _update_issue_references was called")
        return
    
    # Find and replace the references section
    # Pattern matches the header followed by the references on the next line(s)
    # This handles cases where there might be whitespace or multiple lines of refs
    pattern = r"(### Siehe folgende Issues und PRs\s*\n)([^\n#]*)"
    replacement = rf"\g<1>{issue_pr_refs}"
    
    if re.search(pattern, item.description):
        item.description = re.sub(pattern, replacement, item.description)
        item.save(update_fields=['description'])
    else:
        logger.warning(
            f"Could not find 'Siehe folgende Issues und PRs' section in item {item.id} description. "
            f"The newly created issue may not be added to references."
        )


def _sync_answered_questions_to_description(item):
    """
    Update the item description to include answered questions at the bottom.
    
    This ensures Copilot can see the answered questions in the issue context.
    The section is dynamically regenerated each time to reflect current state.
    
    Args:
        item: Item instance
    """
    import re
    
    # Section title constant
    QUESTIONS_SECTION_TITLE = "## Offene Fragen"
    
    # Get all answered or dismissed questions
    questions = IssueOpenQuestion.objects.filter(
        issue=item
    ).exclude(
        status=OpenQuestionStatus.OPEN
    ).order_by('created_at')
    
    if not questions.exists():
        # No answered questions, remove the section if it exists
        pattern = rf'(?:\n|^){QUESTIONS_SECTION_TITLE}\n.*?(?=\n##|\Z)'
        item.description = re.sub(pattern, '', item.description, flags=re.DOTALL).strip()
        item.save(update_fields=['description'])
        return
    
    # Build the questions section
    questions_section = f"\n\n{QUESTIONS_SECTION_TITLE}\n\n"
    
    for q in questions:
        status_marker = "[x]" if q.status in [OpenQuestionStatus.ANSWERED, OpenQuestionStatus.DISMISSED] else "[ ]"
        questions_section += f"- {status_marker} {q.question}\n"
        
        if q.status == OpenQuestionStatus.ANSWERED and q.get_answer_display_text():
            # Indent the answer
            answer_text = q.get_answer_display_text().strip()
            questions_section += f"  Antwort: {answer_text}\n"
        
        questions_section += "\n"
    
    # Remove trailing newlines
    questions_section = questions_section.rstrip() + "\n"
    
    # Check if questions section already exists
    pattern = rf'(?:\n|^){QUESTIONS_SECTION_TITLE}\n.*?(?=\n##|\Z)'
    
    if re.search(pattern, item.description, flags=re.DOTALL):
        # Replace existing section
        item.description = re.sub(pattern, questions_section, item.description, flags=re.DOTALL)
    else:
        # Append new section at the end
        item.description = item.description.rstrip() + questions_section
    
    item.save(update_fields=['description'])


@login_required

@require_POST
def item_optimize_description_ai(request, item_id):
    """
    Optimize item description using AI and RAG.
    
    Uses RAG to gather context and then calls the github-issue-creation-agent
    to generate an optimized, machine-readable GitHub issue description.
    
    Only available to users with Agent role.
    """
    from core.services.rag import build_extended_context
    
    # Check user role
    if not request.user.is_authenticated or request.user.role != UserRole.AGENT:
        return JsonResponse({
            'status': 'error',
            'message': 'This feature is only available to users with Agent role'
        }, status=403)
    
    item = get_object_or_404(Item, id=item_id)
    
    try:
        # Get current description
        current_description = item.description or ""
        
        if not current_description.strip():
            return JsonResponse({
                'status': 'error',
                'message': 'Item has no description to optimize'
            }, status=400)
        
        # Build extended RAG context with question optimization
        rag_context = build_extended_context(
            query=current_description,
            project_id=str(item.project.id),
            item_id=str(item.id),
            current_item_id=str(item.id),  # Exclude current item from results (Issue #392)
            user=request.user,
            client_ip=request.META.get('REMOTE_ADDR')
        )
        
        # Build input text for agent: description + RAG context
        context_text = rag_context.to_context_text() if rag_context.all_items else "No additional context found."
        
        agent_input = f"""Original Description:
{current_description}

---
Context from similar items and related information:
{context_text}
"""
        
        # Execute the github-issue-creation-agent
        agent_service = AgentService()
        agent_response = agent_service.execute_agent(
            filename='github-issue-creation-agent.yml',
            input_text=agent_input,
            user=request.user,
            client_ip=request.META.get('REMOTE_ADDR')
        )
        
        # Parse the response - expect JSON format:
        # {"issue": {"description": "..."}, "open_questions": [...]}
        
        # Clean the response: remove markdown code fences if present
        cleaned_response = agent_response.strip()
        
        # Remove ```json or ``` code fences
        if cleaned_response.startswith('```'):
            # Find the end of the first line (```json or just ```)
            first_newline = cleaned_response.find('\n')
            if first_newline != -1:
                # Multi-line: remove first line
                cleaned_response = cleaned_response[first_newline + 1:]
            else:
                # Single line like ```json{...}``` - remove opening fence
                cleaned_response = cleaned_response[3:]  # Remove ```
                # Also strip any language identifier (json, etc)
                if cleaned_response and cleaned_response[0] not in ('{', '['):
                    # Find where JSON actually starts
                    for i, char in enumerate(cleaned_response):
                        if char in ('{', '['):
                            cleaned_response = cleaned_response[i:]
                            break
            
            # Remove trailing ```
            if cleaned_response.endswith('```'):
                cleaned_response = cleaned_response[:-3]
            
            cleaned_response = cleaned_response.strip()
        
        try:
            # Try to parse as JSON
            response_data = json.loads(cleaned_response)
            
            if isinstance(response_data, dict) and 'issue' in response_data and isinstance(response_data['issue'], dict):
                # Expected format: nested issue object with description
                optimized_description = response_data['issue'].get('description', '').strip()
                open_questions = response_data.get('open_questions', [])
            elif isinstance(response_data, dict) and 'description' in response_data:
                # Fallback format: direct description field (for backward compatibility)
                optimized_description = response_data.get('description', '').strip()
                open_questions = response_data.get('open_questions', [])
            else:
                # No recognized JSON format
                # Save cleaned response (code fences removed) as fallback
                # This preserves agent output even if format is unexpected
                optimized_description = cleaned_response
                open_questions = []
        except json.JSONDecodeError:
            # Not valid JSON after cleaning
            # Save cleaned response (code fences removed) as fallback
            optimized_description = cleaned_response
            open_questions = []
        
        # Validate we have a description
        if not optimized_description:
            raise ValueError("AI agent returned empty description")
        
        # Update item description (only the issue.description part)
        item.description = optimized_description
        item.save()
        
        # Process open questions
        questions_added = 0
        if open_questions and isinstance(open_questions, list):
            for question_text in open_questions:
                if not question_text or not isinstance(question_text, str):
                    continue
                
                question_text = question_text.strip()
                if not question_text:
                    continue
                
                # Check if an open question with identical text already exists
                existing = IssueOpenQuestion.objects.filter(
                    issue=item,
                    question=question_text,
                    status=OpenQuestionStatus.OPEN
                ).exists()
                
                if not existing:
                    # Create new open question
                    IssueOpenQuestion.objects.create(
                        issue=item,
                        question=question_text,
                        source=OpenQuestionSource.AI_AGENT,
                        sort_order=questions_added
                    )
                    questions_added += 1
        
        # Log activity - success
        activity_service = ActivityService()
        activity_service.log(
            verb='item.description.ai_optimized',
            target=item,
            actor=request.user,
            summary='Item description optimized via AI (RAG + GitHub agent)',
        )
        
        return JsonResponse({
            'status': 'ok'
        })
        
    except Exception as e:
        # Log activity - error
        activity_service = ActivityService()
        activity_service.log(
            verb='item.description.ai_error',
            target=item,
            actor=request.user,
            summary=f'AI description optimization failed: {str(e)}',
        )
        
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"AI description optimization failed for item {item_id}: {str(e)}")
        
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)


def _get_newest_pr_context(item):
    """
    Get context from the newest linked GitHub PR for an item.
    
    Returns a formatted context block with PR title and body, or None if:
    - No PRs are linked
    - GitHub integration is not configured
    - API call fails
    
    Args:
        item: Item instance
        
    Returns:
        str: Formatted PR context block or None
    """
    from core.services.github.client import GitHubClient
    from core.services.integrations.base import IntegrationError
    from core.models import GitHubConfiguration
    
    try:
        # Get all PR mappings for this item, ordered by number (descending = newest first)
        pr_mappings = item.external_mappings.filter(
            kind=ExternalIssueKind.PR
        ).order_by('-number')
        
        if not pr_mappings.exists():
            return None
        
        # Get the newest PR (highest number)
        newest_pr_mapping = pr_mappings.first()
        
        # Check if GitHub is enabled and configured
        config = GitHubConfiguration.load()
        if not config.enable_github or not config.github_token:
            logger.info(f"GitHub integration not enabled/configured, skipping PR context for item {item.id}")
            return None
        
        # Get repository info from project
        project = item.project
        if not project.github_owner or not project.github_repo:
            logger.info(f"Project {project.name} does not have GitHub repository configured")
            return None
        
        owner = project.github_owner
        repo = project.github_repo
        
        # Create GitHub client
        client = GitHubClient(
            token=config.github_token,
            base_url=config.github_api_base_url,
        )
        
        # Fetch PR details from GitHub API
        pr_data = client.get_pr(owner, repo, newest_pr_mapping.number)
        
        # Extract title and body
        pr_title = pr_data.get('title') or f'Pull Request #{newest_pr_mapping.number}'
        pr_body = pr_data.get('body', '')
        
        # Build context block
        context_parts = [
            "## GitHub PR (latest) - Description",
            f"**Title:** {pr_title}",
            f"**PR Number:** #{newest_pr_mapping.number}",
            f"**URL:** {newest_pr_mapping.html_url}",
        ]
        
        if pr_body and pr_body.strip():
            context_parts.append("\n**Description:**")
            context_parts.append(pr_body)
        else:
            context_parts.append("\n*No description provided in PR*")
        
        return "\n".join(context_parts)
        
    except IntegrationError as e:
        # GitHub integration error - log and return None (graceful degradation)
        logger.warning(f"GitHub integration error while fetching PR context for item {item.id}: {e}")
        return None
    except Exception as e:
        # Unexpected error - log and return None (graceful degradation)
        logger.warning(f"Failed to fetch PR context for item {item.id}: {e}")
        return None


@login_required
@require_POST
def item_generate_solution_ai(request, item_id):
    """
    Generate solution description using AI and RAG.
    
    Uses RAG to gather context and then calls the create-user-description agent
    to generate a solution description based on the item description.
    
    When item status is TESTING and has linked GitHub PRs, includes the PR body
    of the newest PR in the context.
    
    Only available to users with Agent role.
    """
    from core.services.rag import build_extended_context
    
    # Check user role
    if request.user.role != UserRole.AGENT:
        return JsonResponse({
            'status': 'error',
            'message': 'This feature is only available to users with Agent role'
        }, status=403)
    
    item = get_object_or_404(Item, id=item_id)
    
    try:
        # Get current description
        current_description = item.description or ""
        
        if not current_description.strip():
            return JsonResponse({
                'status': 'error',
                'message': 'Item has no description. Please provide a description first.'
            }, status=400)
        
        # Build extended RAG context with question optimization
        rag_context = build_extended_context(
            query=current_description,
            project_id=str(item.project.id),
            item_id=str(item.id),
            current_item_id=str(item.id),  # Exclude current item from results (Issue #392)
            user=request.user,
            client_ip=request.META.get('REMOTE_ADDR')
        )
        
        # Build input text for agent: description + RAG context
        context_text = rag_context.to_context_text() if rag_context.all_items else "No additional context found."
        
        agent_input = f"""Item Description:
{current_description}

---
Context from similar items and related information:
{context_text}
"""
        
        # If item status is TESTING, add PR context from newest linked PR
        if item.status == ItemStatus.TESTING:
            pr_context = _get_newest_pr_context(item)
            if pr_context:
                agent_input += f"\n\n---\n{pr_context}"
        
        # Execute the create-user-description agent
        agent_service = AgentService()
        solution_description = agent_service.execute_agent(
            filename='create-user-description.yml',
            input_text=agent_input,
            user=request.user,
            client_ip=request.META.get('REMOTE_ADDR')
        )
        
        # Update item solution_description
        item.solution_description = solution_description.strip()
        item.save(update_fields=['solution_description'])
        
        # Log activity - success
        activity_service = ActivityService()
        activity_service.log(
            verb='item.solution_description.ai_generated',
            target=item,
            actor=request.user,
            summary='Solution description generated via AI (RAG + create-user-description agent)',
        )
        
        return JsonResponse({
            'status': 'ok'
        })
        
    except Exception as e:
        # Log activity - error
        activity_service = ActivityService()
        activity_service.log(
            verb='item.solution_description.ai_error',
            target=item,
            actor=request.user,
            summary=f'AI solution description generation failed: {str(e)}',
        )
        
        logger.error(f"AI solution description generation failed for item {item_id}: {str(e)}")
        
        return JsonResponse({
            'status': 'error',
            'message': 'Failed to generate solution description. Please try again later.'
        }, status=500)


@login_required
@require_POST
def item_generate_short_description_ai(request, item_id):
    """
    Generate short description using AI based on item description.
    
    Uses the item-short-description-agent to generate an ISO-compliant
    short description (max 3-4 sentences) from the item description.
    
    Only available to users with Agent role.
    """
    # Check user role
    if request.user.role != UserRole.AGENT:
        return JsonResponse({
            'status': 'error',
            'message': 'This feature is only available to users with Agent role'
        }, status=403)
    
    item = get_object_or_404(Item, id=item_id)
    activity_service = ActivityService()
    
    try:
        # Get current description
        current_description = item.description or ""
        
        if not current_description.strip():
            return JsonResponse({
                'status': 'error',
                'message': 'Item has no description. Please provide a description first.'
            }, status=400)
        
        # Execute the item-short-description-agent
        agent_service = AgentService()
        short_description = agent_service.execute_agent(
            filename='item-short-description-agent.yml',
            input_text=current_description,
            user=request.user,
            client_ip=request.META.get('REMOTE_ADDR')
        )
        
        # Update item short_description
        item.short_description = short_description.strip()
        item.save(update_fields=['short_description'])
        
        # Log activity - success
        activity_service.log(
            verb='item.short_description.ai_generated',
            target=item,
            actor=request.user,
            summary='Short description generated via AI (item-short-description-agent)',
        )
        
        return JsonResponse({
            'status': 'ok',
            'short_description': short_description.strip()
        })
        
    except Exception as e:
        # Log activity - error
        activity_service.log(
            verb='item.short_description.ai_error',
            target=item,
            actor=request.user,
            summary=f'AI short description generation failed: {str(e)}',
        )
        
        logger.error(f"AI short description generation failed for item {item_id}: {str(e)}")
        
        return JsonResponse({
            'status': 'error',
            'message': 'Failed to generate short description. Please try again later.'
        }, status=500)


@login_required
@require_POST
def item_pre_review(request, item_id):
    """
    Generate a Pre-Review of an item using AI and RAG.
    
    Uses RAG to gather context and then calls the issue-analyse-agent
    to generate a comprehensive review with recommendations.
    
    Only available to users with Agent role.
    """
    import json
    from core.services.rag import build_extended_context
    
    # Check user role
    if not request.user.is_authenticated or request.user.role != UserRole.AGENT:
        return JsonResponse({
            'success': False,
            'error': 'This feature is only available to users with Agent role'
        }, status=403)
    
    item = get_object_or_404(Item, id=item_id)
    
    try:
        # Get current description
        current_description = item.description or ""
        
        if not current_description.strip():
            return JsonResponse({
                'success': False,
                'error': 'Item has no description to review'
            }, status=400)
        
        # Build extended RAG context with question optimization
        rag_context = build_extended_context(
            query=current_description,
            project_id=str(item.project.id),
            item_id=str(item.id),
            current_item_id=str(item.id),  # Exclude current item from results (Issue #392)
            user=request.user,
            client_ip=request.META.get('REMOTE_ADDR')
        )
        
        # Build input text for agent: description + RAG context
        context_text = rag_context.to_context_text() if rag_context.all_items else "No additional context found."
        
        agent_input = f"""Item Title: {item.title}

Item Description:
{current_description}

---
Context from similar items and related information:
{context_text}
"""
        
        # Execute the issue-analyse-agent
        agent_service = AgentService()
        review = agent_service.execute_agent(
            filename='issue-analyse-agent.yml',
            input_text=agent_input,
            user=request.user,
            client_ip=request.META.get('REMOTE_ADDR')
        )
        
        # Convert markdown to HTML for display (create new parser instance to avoid race conditions)
        md_parser = markdown.Markdown(extensions=['extra', 'fenced_code'])
        review_html = md_parser.convert(review)
        review_html = bleach.clean(review_html, tags=ALLOWED_TAGS, attributes=ALLOWED_ATTRIBUTES)
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='item.pre_review.generated',
            target=item,
            actor=request.user,
            summary='Pre-Review generated via AI (RAG + issue-analyse-agent)',
        )
        
        return JsonResponse({
            'success': True,
            'review': review,
            'review_html': review_html
        })
        
    except Exception as e:
        # Log activity - error
        activity_service = ActivityService()
        activity_service.log(
            verb='item.pre_review.error',
            target=item,
            actor=request.user,
            summary=f'Pre-Review generation failed: {str(e)}',
        )
        
        logger.error(f"Pre-Review generation failed for item {item_id}: {str(e)}")
        
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)


@login_required
@require_POST
def item_save_pre_review(request, item_id):
    """
    Save Pre-Review as an AI-generated comment on the item.
    
    The comment will be marked with kind=AI_GENERATED and synced to Weaviate.
    
    Only available to users with Agent role.
    """
    import json
    from core.models import CommentKind
    
    # Check user role
    if not request.user.is_authenticated or request.user.role != UserRole.AGENT:
        return JsonResponse({
            'success': False,
            'error': 'This feature is only available to users with Agent role'
        }, status=403)
    
    item = get_object_or_404(Item, id=item_id)
    
    try:
        # Parse request body
        data = json.loads(request.body)
        review = data.get('review', '').strip()
        
        if not review:
            return JsonResponse({
                'success': False,
                'error': 'No review content to save'
            }, status=400)
        
        # Create AI-generated comment
        comment = ItemComment.objects.create(
            item=item,
            author=request.user,
            body=review,
            kind=CommentKind.AI_GENERATED,
            subject='AI Pre-Review'
        )
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='item.comment.ai_review_saved',
            target=item,
            actor=request.user,
            summary='Pre-Review saved as AI-generated comment',
        )
        
        # The Weaviate sync will happen automatically via signals (if configured)
        
        return JsonResponse({
            'success': True,
            'comment_id': comment.id
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': 'Invalid JSON data'
        }, status=400)
    except Exception as e:
        logger.error(f"Failed to save Pre-Review for item {item_id}: {str(e)}")
        
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)


@login_required
def item_rag_retrieval_raw(request, item_id):
    """
    Execute RAG retrieval for the item and return raw results as deterministic Markdown.
    
    This feature allows users to inspect the RAG retrieval results without any
    content transformation. Results are formatted as structured Markdown for readability.
    
    Only available to users with Agent role.
    """
    from core.services.rag import build_extended_context
    from datetime import datetime
    
    # Check user role
    if not request.user.is_authenticated or request.user.role != UserRole.AGENT:
        return JsonResponse({
            'status': 'error',
            'message': 'This feature is only available to users with Agent role'
        }, status=403)
    
    item = get_object_or_404(Item, id=item_id)
    
    try:
        # Get current description to use as query
        if not (item.description or "").strip():
            return JsonResponse({
                'status': 'error',
                'message': 'Item has no description to use as RAG query'
            }, status=400)
        
        query = item.description
        
        # Record start time for duration calculation
        start_time = datetime.now()
        
        # Build extended RAG context with question optimization
        rag_context = build_extended_context(
            query=query,
            project_id=str(item.project.id),
            item_id=str(item.id),
            current_item_id=str(item.id),  # Exclude current item from results (Issue #392)
            user=request.user,
            client_ip=request.META.get('REMOTE_ADDR')
        )
        
        # Calculate duration
        duration_ms = int((datetime.now() - start_time).total_seconds() * 1000)
        
        # Format results as deterministic Markdown
        markdown_output = _format_rag_results_as_markdown(
            query=query,
            rag_context=rag_context,
            duration_ms=duration_ms
        )
        
        return JsonResponse({
            'status': 'success',
            'markdown': markdown_output
        })
        
    except Exception as e:
        logger.error(f"Failed to execute RAG retrieval for item {item_id}: {str(e)}")
        
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)


def _format_rag_results_as_markdown(query: str, rag_context, duration_ms: int) -> str:
    """
    Format RAG context results as deterministic Markdown output.
    
    This function creates a structured, deterministic Markdown representation of
    RAG retrieval results without any content transformation.
    
    Args:
        query: The original search query
        rag_context: RAGContext or ExtendedRAGContext object with search results
        duration_ms: Duration of the search in milliseconds
        
    Returns:
        Formatted Markdown string
    """
    # Constants
    MAX_QUERY_DISPLAY_LENGTH = 100
    
    lines = []
    
    # Get items - handle both RAGContext and ExtendedRAGContext
    items = getattr(rag_context, 'all_items', None) or getattr(rag_context, 'items', [])
    alpha = getattr(rag_context, 'alpha', None)
    
    # Header section
    lines.append("## RAG Retrieval (raw)")
    lines.append("")
    lines.append("### Header")
    lines.append(f"- **query:** `{query[:MAX_QUERY_DISPLAY_LENGTH]}{'...' if len(query) > MAX_QUERY_DISPLAY_LENGTH else ''}`")
    lines.append(f"- **search_type:** `hybrid`")
    if alpha is not None:
        lines.append(f"- **alpha:** `{alpha:.2f}`")
    lines.append(f"- **duration_ms:** `{duration_ms}`")
    lines.append(f"- **hits:** `{len(items)}`")
    
    # Add extended RAG stats if available
    if hasattr(rag_context, 'stats'):
        stats = rag_context.stats
        if stats.get('optimization_success'):
            lines.append(f"- **optimization:** `success`")
        if 'sem_results' in stats:
            lines.append(f"- **semantic_results:** `{stats.get('sem_results', 0)}`")
        if 'kw_results' in stats:
            lines.append(f"- **keyword_results:** `{stats.get('kw_results', 0)}`")
    
    lines.append("")
    lines.append("---")
    lines.append("")
    
    # Hits section
    lines.append("### Hits")
    lines.append("")
    
    if not items:
        lines.append("*No hits found*")
    else:
        # Sort by score if available (descending), otherwise keep order as delivered
        if items and items[0].relevance_score is not None:
            # Sort by score descending
            items = sorted(items, key=lambda x: x.relevance_score or 0, reverse=True)
        
        for idx, hit in enumerate(items, 1):
            lines.append(f"#### {idx}) Hit")
            lines.append("")
            
            # Score (if available)
            if hit.relevance_score is not None:
                lines.append(f"- **score:** `{hit.relevance_score:.4f}`")
            
            # Source/Object reference
            source_parts = []
            if hit.object_type:
                source_parts.append(f"type={hit.object_type}")
            if hit.object_id:
                source_parts.append(f"id={hit.object_id}")
            if hit.source:
                source_parts.append(f"source={hit.source}")
            if source_parts:
                lines.append(f"- **source:** `{', '.join(source_parts)}`")
            
            # Link
            if hit.link:
                lines.append(f"- **link:** `{hit.link}`")
            
            # Title
            if hit.title:
                lines.append(f"- **title:** `{hit.title}`")
            
            # Updated timestamp
            if hit.updated_at:
                lines.append(f"- **updated_at:** `{hit.updated_at}`")
            
            # Context/Content in code block
            lines.append("")
            lines.append("**context:**")
            lines.append("```")
            lines.append(hit.content or "")
            lines.append("```")
            lines.append("")
    
    return "\n".join(lines)


@login_required
@require_POST
def item_open_question_answer(request, question_id):
    """
    Answer or dismiss an open question.
    
    Expects JSON body with:
    - action: 'answer' or 'dismiss'
    - answer_type: 'free_text' or 'standard_answer' (for action='answer')
    - answer_text: Free text answer (if answer_type='free_text')
    - standard_answer_id: ID of standard answer (if answer_type='standard_answer')
    """
    from core.models import OpenQuestionAnswerType
    
    question = get_object_or_404(IssueOpenQuestion, id=question_id)
    
    # Check permissions - user must be authenticated
    if not request.user.is_authenticated:
        return JsonResponse({
            'success': False,
            'error': 'Authentication required'
        }, status=403)
    
    try:
        data = json.loads(request.body)
        action = data.get('action', '').lower()
        
        if action == 'dismiss':
            # Mark as dismissed
            question.status = OpenQuestionStatus.DISMISSED
            question.answered_at = timezone.now()
            question.answered_by = request.user
            question.save()
            
            # Update item description with answered questions
            _sync_answered_questions_to_description(question.issue)
            
            # Log activity
            activity_service = ActivityService()
            activity_service.log(
                verb='item.open_question.dismissed',
                target=question.issue,
                actor=request.user,
                summary=f'Open question dismissed: {question.question[:50]}...',
            )
            
            return JsonResponse({
                'success': True,
                'status': question.status
            })
            
        elif action == 'answer':
            # Validate answer
            answer_type = data.get('answer_type', '').lower()
            
            if answer_type == 'free_text':
                answer_text = data.get('answer_text', '').strip()
                if not answer_text:
                    return JsonResponse({
                        'success': False,
                        'error': 'Answer text is required for free text answers'
                    }, status=400)
                
                question.answer_type = OpenQuestionAnswerType.FREE_TEXT
                question.answer_text = answer_text
                question.standard_answer = None
                question.standard_answer_key = None
                
            elif answer_type == 'standard_answer':
                standard_answer_id = data.get('standard_answer_id')
                if not standard_answer_id:
                    return JsonResponse({
                        'success': False,
                        'error': 'Standard answer ID is required'
                    }, status=400)
                
                standard_answer = get_object_or_404(IssueStandardAnswer, id=standard_answer_id, is_active=True)
                
                question.answer_type = OpenQuestionAnswerType.STANDARD_ANSWER
                question.answer_text = None
                question.standard_answer = standard_answer
                question.standard_answer_key = standard_answer.key
                
            else:
                return JsonResponse({
                    'success': False,
                    'error': 'Invalid answer_type. Must be "free_text" or "standard_answer"'
                }, status=400)
            
            # Mark as answered
            question.status = OpenQuestionStatus.ANSWERED
            question.answered_at = timezone.now()
            question.answered_by = request.user
            question.save()
            
            # Update item description with answered questions
            _sync_answered_questions_to_description(question.issue)
            
            # Log activity
            activity_service = ActivityService()
            activity_service.log(
                verb='item.open_question.answered',
                target=question.issue,
                actor=request.user,
                summary=f'Open question answered: {question.question[:50]}...',
            )
            
            return JsonResponse({
                'success': True,
                'status': question.status,
                'answer': question.get_answer_display_text()
            })
            
        else:
            return JsonResponse({
                'success': False,
                'error': 'Invalid action. Must be "answer" or "dismiss"'
            }, status=400)
            
    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': 'Invalid JSON in request body'
        }, status=400)
    except Exception as e:
        logger.error(f"Error answering open question {question_id}: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)


@login_required
def item_open_questions_list(request, item_id):
    """
    Get list of open questions for an item.
    
    Returns JSON with:
    - questions: list of question objects
    - standard_answers: list of available standard answers
    - has_open: boolean indicating if there are open questions
    """
    item = get_object_or_404(Item, id=item_id)
    
    # Get all questions for this item
    questions = IssueOpenQuestion.objects.filter(issue=item).select_related(
        'standard_answer', 'answered_by'
    )
    
    # Get active standard answers
    standard_answers = IssueStandardAnswer.objects.filter(is_active=True)
    
    # Format questions
    questions_data = []
    for q in questions:
        standard_answer_data = None
        if q.standard_answer:
            standard_answer_data = {
                'id': q.standard_answer.id,
                'key': q.standard_answer.key,
                'label': q.standard_answer.label,
            }
        
        questions_data.append({
            'id': q.id,
            'question': q.question,
            'status': q.status,
            'status_display': q.get_status_display(),
            'answer_type': q.answer_type if q.answer_type != 'None' else None,
            'answer_text': q.get_answer_display_text(),
            'standard_answer': standard_answer_data,
            'source': q.source,
            'source_display': q.get_source_display(),
            'created_at': q.created_at.isoformat(),
            'answered_at': q.answered_at.isoformat() if q.answered_at else None,
            'answered_by': q.answered_by.name if q.answered_by else None,
        })
    
    # Format standard answers
    standard_answers_data = [
        {
            'id': sa.id,
            'key': sa.key,
            'label': sa.label,
            'text': sa.text,
        }
        for sa in standard_answers
    ]
    
    # Check if there are open questions
    has_open = questions.filter(status=OpenQuestionStatus.OPEN).exists()
    
    return JsonResponse({
        'success': True,
        'questions': questions_data,
        'standard_answers': standard_answers_data,
        'has_open': has_open,
    })


@login_required
@require_POST
def item_answer_question_ai(request, question_id):
    """
    Answer an open question using AI with RAG context from Weaviate.
    
    Uses the item-answer-question agent to generate a short, bullet-point
    answer based exclusively on the RAG context retrieved from Weaviate.
    
    Only available to users with Agent role.
    """
    from core.services.rag import build_extended_context
    from core.models import OpenQuestionAnswerType
    
    # Check user role
    if not request.user.is_authenticated or request.user.role != UserRole.AGENT:
        return JsonResponse({
            'status': 'error',
            'message': 'This feature is only available to users with Agent role'
        }, status=403)
    
    question = get_object_or_404(IssueOpenQuestion, id=question_id)
    
    # Only allow answering open questions
    if question.status != OpenQuestionStatus.OPEN:
        return JsonResponse({
            'status': 'error',
            'message': 'Only open questions can be answered with AI'
        }, status=400)
    
    try:
        # Get the question text
        question_text = question.question
        
        if not question_text.strip():
            return JsonResponse({
                'status': 'error',
                'message': 'Question has no text'
            }, status=400)
        
        # Build extended RAG context using the question as search query
        rag_context = build_extended_context(
            query=question_text,
            project_id=str(question.issue.project.id),
            current_item_id=str(question.issue.id),  # Exclude current item from results (Issue #392)
            user=request.user,
            client_ip=request.META.get('REMOTE_ADDR')
        )
        
        # Build input text for agent: question + RAG context
        context_text = rag_context.to_context_text() if rag_context.all_items else RAG_NO_CONTEXT_MESSAGE
        
        agent_input = f"""Question:
{question_text}

---
Relevant context from knowledge base:
{context_text}
"""
        
        # Execute the item-answer-question agent
        agent_service = AgentService()
        agent_response = agent_service.execute_agent(
            filename='item-answer-question.yml',
            input_text=agent_input,
            user=request.user,
            client_ip=request.META.get('REMOTE_ADDR')
        )
        
        # Clean the response
        answer_text = agent_response.strip()
        
        # Validate we have an answer
        if not answer_text:
            raise ValueError("AI agent returned empty answer")
        
        # Update question with AI-generated answer
        question.answer_type = OpenQuestionAnswerType.FREE_TEXT
        question.answer_text = answer_text
        question.status = OpenQuestionStatus.ANSWERED
        question.answered_at = timezone.now()
        question.answered_by = request.user
        question.save()
        
        # Update item description with answered questions
        _sync_answered_questions_to_description(question.issue)
        
        # Log activity - success
        activity_service = ActivityService()
        activity_service.log(
            verb='item.open_question.ai_answered',
            target=question.issue,
            actor=request.user,
            summary=f'Question answered by AI: {question_text[:50]}...',
        )
        
        return JsonResponse({
            'status': 'success',
            'answer': answer_text,
            'question_id': question.id
        })
        
    except Exception as e:
        logger.error(f"Error answering question {question_id} with AI: {str(e)}")
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)


@login_required
@require_http_methods(["PUT", "PATCH", "POST"])
def item_open_question_edit(request, question_id):
    """
    Edit/update an open question's text.
    
    Allows editing of question text regardless of answered status.
    The answer (if any) remains unchanged.
    
    Expects JSON body with:
    - question: The new question text
    """
    question = get_object_or_404(IssueOpenQuestion, id=question_id)
    
    try:
        data = json.loads(request.body)
        new_question_text = data.get('question', '').strip()
        
        if not new_question_text:
            return JsonResponse({
                'success': False,
                'error': 'Question text cannot be empty'
            }, status=400)
        
        # Update question text
        old_question = question.question
        question.question = new_question_text
        question.save()
        
        # Log activity
        activity_service = ActivityService()
        old_summary = f'"{old_question[:30]}"' + ('...' if len(old_question) > 30 else '')
        new_summary = f'"{new_question_text[:30]}"' + ('...' if len(new_question_text) > 30 else '')
        activity_service.log(
            verb='item.open_question.edited',
            target=question.issue,
            actor=request.user,
            summary=f'Question edited: {old_summary} → {new_summary}',
        )
        
        return JsonResponse({
            'success': True,
            'question_id': question.id,
            'question': question.question
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': 'Invalid JSON in request body'
        }, status=400)
    except Exception as e:
        logger.error(f"Error editing open question {question_id}: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)


@login_required
@require_http_methods(["DELETE", "POST"])
def item_open_question_delete(request, question_id):
    """
    Delete an open question.
    
    Allows deletion of questions regardless of answered status.
    """
    question = get_object_or_404(IssueOpenQuestion, id=question_id)
    
    try:
        # Store info for logging
        question_text = question.question
        issue = question.issue
        
        # Delete the question
        question.delete()
        
        # Log activity
        activity_service = ActivityService()
        summary_text = f'"{question_text[:50]}"' + ('...' if len(question_text) > 50 else '')
        activity_service.log(
            verb='item.open_question.deleted',
            target=issue,
            actor=request.user,
            summary=f'Question deleted: {summary_text}',
        )
        
        return JsonResponse({
            'success': True,
            'question_id': question_id
        })
        
    except Exception as e:
        logger.error(f"Error deleting open question {question_id}: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)


@login_required
@require_POST
def item_open_question_answer_edit(request, question_id):
    """
    Edit an existing answer to an open question.
    
    Allows changing the answer type (free_text ↔ standard_answer) and content.
    
    Expects JSON body with:
    - answer_type: 'free_text' or 'standard_answer'
    - answer_text: Free text answer (if answer_type='free_text')
    - standard_answer_id: ID of standard answer (if answer_type='standard_answer')
    """
    from core.models import OpenQuestionAnswerType
    
    question = get_object_or_404(IssueOpenQuestion, id=question_id)
    
    # Check if question is answered
    if question.status != OpenQuestionStatus.ANSWERED:
        return JsonResponse({
            'success': False,
            'error': 'Question must be answered to edit the answer'
        }, status=400)
    
    try:
        data = json.loads(request.body)
        answer_type = data.get('answer_type', '').lower()
        
        if answer_type == 'free_text':
            answer_text = data.get('answer_text', '').strip()
            if not answer_text:
                return JsonResponse({
                    'success': False,
                    'error': 'Answer text is required for free text answers'
                }, status=400)
            
            question.answer_type = OpenQuestionAnswerType.FREE_TEXT
            question.answer_text = answer_text
            question.standard_answer = None
            question.standard_answer_key = None
            
        elif answer_type == 'standard_answer':
            standard_answer_id = data.get('standard_answer_id')
            if not standard_answer_id:
                return JsonResponse({
                    'success': False,
                    'error': 'Standard answer ID is required'
                }, status=400)
            
            standard_answer = get_object_or_404(IssueStandardAnswer, id=standard_answer_id, is_active=True)
            
            question.answer_type = OpenQuestionAnswerType.STANDARD_ANSWER
            question.answer_text = None
            question.standard_answer = standard_answer
            question.standard_answer_key = standard_answer.key
            
        else:
            return JsonResponse({
                'success': False,
                'error': 'Invalid answer_type. Must be "free_text" or "standard_answer"'
            }, status=400)
        
        # Update answered metadata
        question.answered_at = timezone.now()
        question.answered_by = request.user
        question.save()
        
        # Update item description with answered questions
        _sync_answered_questions_to_description(question.issue)
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='item.open_question.answer_edited',
            target=question.issue,
            actor=request.user,
            summary=f'Answer edited for question: {question.question[:50]}...',
        )
        
        return JsonResponse({
            'success': True,
            'status': question.status,
            'answer': question.get_answer_display_text()
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': 'Invalid JSON in request body'
        }, status=400)
    except Exception as e:
        logger.error(f"Error editing answer for question {question_id}: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)


@login_required
@require_POST
def item_open_question_answer_delete(request, question_id):
    """
    Delete the answer to an open question, reverting it to Open status.
    
    This removes the answer but keeps the question itself.
    """
    from core.models import OpenQuestionAnswerType
    
    question = get_object_or_404(IssueOpenQuestion, id=question_id)
    
    # Check if question is answered
    if question.status != OpenQuestionStatus.ANSWERED and question.status != OpenQuestionStatus.DISMISSED:
        return JsonResponse({
            'success': False,
            'error': 'Question must be answered or dismissed to delete the answer'
        }, status=400)
    
    try:
        # Clear answer fields
        question.answer_type = OpenQuestionAnswerType.NONE
        question.answer_text = None
        question.standard_answer = None
        question.standard_answer_key = None
        question.answered_by = None
        question.answered_at = None
        
        # Revert status to Open
        question.status = OpenQuestionStatus.OPEN
        question.save()
        
        # Update item description with answered questions
        _sync_answered_questions_to_description(question.issue)
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='item.open_question.answer_deleted',
            target=question.issue,
            actor=request.user,
            summary=f'Answer deleted for question: {question.question[:50]}...',
        )
        
        return JsonResponse({
            'success': True,
            'question_id': question.id,
            'status': question.status
        })
        
    except Exception as e:
        logger.error(f"Error deleting answer for question {question_id}: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)


@login_required
@require_POST
def item_change_status(request, item_id):
    """HTMX endpoint to change item status."""
    item = get_object_or_404(Item, id=item_id)
    old_status = item.status

    # Accept status from form-data, x-www-form-urlencoded, or JSON payloads.
    # This keeps the endpoint compatible with multiple frontend callers.
    new_status = (request.POST.get('status') or request.POST.get('new_status') or '').strip()

    if not new_status and request.body:
        try:
            payload = json.loads(request.body)
            new_status = (payload.get('status') or payload.get('new_status') or '').strip()
        except (json.JSONDecodeError, UnicodeDecodeError, AttributeError):
            new_status = ''

    if not new_status:
        return JsonResponse({'success': False, 'error': 'Missing status parameter'}, status=400)
    
    try:
        guard = ItemWorkflowGuard()
        guard.transition(item, new_status, actor=request.user if request.user.is_authenticated else None)
        
        # Check for mail trigger (only if status changed)
        mapping = check_mail_trigger(item)
        if mapping and item.status != old_status:
            # Prepare mail preview for modal
            mail_preview = prepare_mail_preview(item, mapping)
            
            # Return JSON response with mail preview
            return JsonResponse({
                'success': True,
                'item_id': item.id,
                'mail_preview': mail_preview,
                'new_status': item.status,
                'new_status_display': item.get_status_display()
            })
        
        # Return updated status badge (normal HTMX swap)
        response = render(request, 'partials/item_status_badge.html', {'item': item})
        response['HX-Trigger'] = 'statusChanged'
        return response
        
    except ValidationError as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required

@require_POST
def item_add_comment(request, item_id):
    """HTMX endpoint to add a comment to an item."""
    item = get_object_or_404(Item, id=item_id)
    body = request.POST.get('body', '').strip()
    
    if not body:
        return HttpResponse("Comment body cannot be empty", status=400)
    
    # Create comment
    comment = ItemComment.objects.create(
        item=item,
        author=request.user if request.user.is_authenticated else None,
        body=body,
    )
    
    # Log activity
    activity_service = ActivityService()
    activity_service.log(
        verb='comment.added',
        target=item,
        actor=request.user if request.user.is_authenticated else None,
        summary=f"Added comment",
    )
    
    # Return updated comments list
    comments = item.comments.select_related('author').order_by('created_at')
    context = {
        'item': item,
        'comments': comments,
    }
    response = render(request, 'partials/item_comments_tab.html', context)
    response['HX-Trigger'] = 'commentAdded'
    return response


@login_required

@require_POST
def item_update_comment(request, comment_id):
    """Update a comment."""
    import json
    
    try:
        comment = get_object_or_404(ItemComment, id=comment_id)
        data = json.loads(request.body)
        new_body = data.get('body', '').strip()
        
        if not new_body:
            return JsonResponse({'success': False, 'error': 'Comment body cannot be empty'}, status=400)
        
        comment.body = new_body
        comment.save()
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='comment.updated',
            target=comment.item,
            actor=request.user if request.user.is_authenticated else None,
            summary=f"Updated comment",
        )
        
        return JsonResponse({'success': True, 'message': 'Comment updated successfully'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required

@require_POST
def item_delete_comment(request, comment_id):
    """Delete a comment."""
    try:
        comment = get_object_or_404(ItemComment, id=comment_id)
        item = comment.item
        comment.delete()
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='comment.deleted',
            target=item,
            actor=request.user if request.user.is_authenticated else None,
            summary=f"Deleted comment",
        )
        
        return JsonResponse({'success': True, 'message': 'Comment deleted successfully'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required

@require_POST
def item_upload_attachment(request, item_id):
    """HTMX endpoint to upload an attachment to an item."""
    item = get_object_or_404(Item, id=item_id)
    
    if 'file' not in request.FILES:
        return HttpResponse("No file provided", status=400)
    
    uploaded_file = request.FILES['file']
    
    try:
        # Store attachment
        storage_service = AttachmentStorageService()
        attachment = storage_service.store_attachment(
            file=uploaded_file,
            target=item,
            role=AttachmentRole.ITEM_FILE,
            created_by=request.user if request.user.is_authenticated else None,
        )
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='attachment.uploaded',
            target=item,
            actor=request.user if request.user.is_authenticated else None,
            summary=f"Uploaded file: {attachment.original_name}",
        )
        
        # Return success response (for AJAX uploads)
        # If this is an AJAX request (multi-file upload), return simple success
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest' and not request.headers.get('HX-Request'):
            return HttpResponse("Upload successful", status=200)
        
        # Return updated attachments list for HTMX requests
        content_type = ContentType.objects.get_for_model(Item)
        attachment_links = AttachmentLink.objects.filter(
            target_content_type=content_type,
            target_object_id=item.id,
            role=AttachmentRole.ITEM_FILE
        ).select_related('attachment', 'attachment__created_by').order_by('-created_at')
        
        attachments = [link.attachment for link in attachment_links if not link.attachment.is_deleted]
        
        context = {
            'item': item,
            'attachments': attachments,
        }
        response = render(request, 'partials/item_attachments_tab.html', context)
        response['HX-Trigger'] = 'attachmentUploaded'
        return response
        
    except ValidationError as e:
        return HttpResponse(f"Validation error: {str(e)}", status=400)
    except PermissionError as e:
        return HttpResponse("Permission denied", status=403)
    except Exception as e:
        # Log the error for debugging
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"Attachment upload failed for item {item_id}: {str(e)}")
        return HttpResponse("Upload failed. Please try again.", status=500)


@login_required

@require_POST
def item_delete_attachment(request, attachment_id):
    """Delete an attachment."""
    try:
        attachment = get_object_or_404(Attachment, id=attachment_id)
        
        # Get the item for activity logging
        attachment_link = AttachmentLink.objects.filter(attachment=attachment).first()
        if attachment_link and hasattr(attachment_link.target, 'id'):
            item = attachment_link.target
        else:
            item = None
        
        # Store filename for logging
        filename = attachment.original_name
        
        # Delete the attachment (hard delete: removes file from storage, DB record, and Weaviate via signal)
        storage_service = AttachmentStorageService()
        storage_service.delete_attachment(attachment, hard=True)
        
        # Log activity
        if item:
            activity_service = ActivityService()
            activity_service.log(
                verb='attachment.deleted',
                target=item,
                actor=request.user if request.user.is_authenticated else None,
                summary=f"Deleted file: {filename}",
            )
        
        return JsonResponse({'success': True, 'message': 'Attachment deleted successfully'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
def attachment_ai_summary(request, attachment_id):
    """
    Generate AI summary for an attachment using text from Weaviate.
    
    Returns HTML for modal content showing the AI-generated summary.
    """
    from core.services.weaviate.client import is_available
    from core.services.weaviate.service import fetch_object_by_type, exists_object
    from core.services.agents import AgentService
    
    attachment = get_object_or_404(Attachment, id=attachment_id)
    
    # Check if Weaviate is available
    if not is_available():
        context = {
            'attachment': attachment,
            'error': 'Weaviate service is not configured or disabled.',
            'available': False,
        }
        return render(request, 'partials/attachment_summary_modal_content.html', context)
    
    # Check if Weaviate object exists
    if not exists_object('attachment', str(attachment_id)):
        context = {
            'attachment': attachment,
            'error': 'This attachment has not been indexed in Weaviate yet.',
            'available': True,
            'exists': False,
        }
        return render(request, 'partials/attachment_summary_modal_content.html', context)
    
    try:
        # Fetch Weaviate object to get text
        obj_data = fetch_object_by_type('attachment', str(attachment_id))
        
        if not obj_data or 'text' not in obj_data:
            context = {
                'attachment': attachment,
                'error': 'No text content found in Weaviate object.',
                'available': True,
                'exists': True,
            }
            return render(request, 'partials/attachment_summary_modal_content.html', context)
        
        text_content = obj_data['text']
        
        if not text_content or len(text_content.strip()) < 10:
            context = {
                'attachment': attachment,
                'error': 'Text content is too short to summarize.',
                'available': True,
                'exists': True,
            }
            return render(request, 'partials/attachment_summary_modal_content.html', context)
        
        # Execute AI agent to generate summary
        agent_service = AgentService()
        summary = agent_service.execute_agent(
            filename='summarize-text-agent.yml',
            input_text=text_content,
            user=request.user if request.user.is_authenticated else None,
            client_ip=request.META.get('REMOTE_ADDR')
        )
        
        context = {
            'attachment': attachment,
            'summary': summary,
            'available': True,
            'exists': True,
            'success': True,
        }
        return render(request, 'partials/attachment_summary_modal_content.html', context)
        
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"Error generating summary for attachment {attachment_id}: {e}")
        
        context = {
            'attachment': attachment,
            'error': f'Error generating summary: {str(e)}',
            'available': True,
            'exists': True,
        }
        return render(request, 'partials/attachment_summary_modal_content.html', context)


@login_required
@require_POST
def item_upload_transcript(request, item_id):
    """Upload a .docx meeting transcript and extract summary and tasks."""
    logger.info(f"Transcript upload started for item {item_id} by user {request.user}")
    
    item = get_object_or_404(Item, id=item_id)
    
    # Verify this is a meeting item
    if item.type.key.lower() != 'meeting':
        logger.warning(f"Transcript upload rejected for item {item_id}: Not a meeting item (type={item.type.key})")
        return JsonResponse({
            'success': False,
            'error': 'This feature is only available for Meeting items.'
        }, status=400)
    
    # Check for uploaded file
    if 'file' not in request.FILES:
        logger.warning(f"Transcript upload rejected for item {item_id}: No file provided in request")
        return JsonResponse({
            'success': False,
            'error': 'No file provided.'
        }, status=400)
    
    uploaded_file = request.FILES['file']
    file_size_mb = uploaded_file.size / (1024 * 1024)
    logger.info(f"Transcript upload for item {item_id}: File '{uploaded_file.name}' ({file_size_mb:.2f} MB)")
    
    # Validate file extension
    if not uploaded_file.name.lower().endswith('.docx'):
        logger.warning(f"Transcript upload rejected for item {item_id}: Invalid file type '{uploaded_file.name}'")
        return JsonResponse({
            'success': False,
            'error': 'Only .docx files are supported. Please upload a Word document.'
        }, status=400)
    
    try:
        # Store attachment with 50 MB limit for transcripts
        logger.debug(f"Storing transcript attachment for item {item_id}...")
        storage_service = AttachmentStorageService(max_size_mb=50)
        attachment = storage_service.store_attachment(
            file=uploaded_file,
            target=item,
            role=AttachmentRole.TRANSKRIPT,
            created_by=request.user if request.user.is_authenticated else None,
        )
        logger.info(f"Transcript attachment stored successfully for item {item_id}: {attachment.id}")
        
        # Log attachment upload
        activity_service = ActivityService()
        activity_service.log(
            verb='attachment.uploaded',
            target=item,
            actor=request.user if request.user.is_authenticated else None,
            summary=f"Uploaded transcript: {attachment.original_name}",
        )
        
        # Extract text from DOCX
        logger.debug(f"Extracting text from DOCX for item {item_id}...")
        from docx import Document
        import io
        
        file_content = storage_service.read_attachment(attachment)
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        # Extract all paragraphs
        transcript_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
        transcript_length = len(transcript_text)
        logger.info(f"Extracted {transcript_length} characters from transcript for item {item_id}")
        
        if not transcript_text.strip():
            logger.warning(f"Transcript processing failed for item {item_id}: Document is empty")
            return JsonResponse({
                'success': False,
                'error': 'The uploaded document appears to be empty.'
            }, status=400)
        
        # Execute AI agent
        logger.info(f"Executing AI agent for transcript processing (item {item_id})...")
        agent_service = AgentService()
        agent_response = agent_service.execute_agent(
            filename='get-meeting-details.yml',
            input_text=transcript_text,
            user=request.user if request.user.is_authenticated else None,
            client_ip=request.META.get('REMOTE_ADDR')
        )
        logger.debug(f"AI agent response received for item {item_id} (length: {len(agent_response)} chars)")
        
        # Parse JSON response
        try:
            import json
            logger.debug(f"Parsing AI agent JSON response for item {item_id}...")
            # Remove markdown code blocks if present
            clean_response = agent_response.strip()
            if clean_response.startswith('```'):
                # Extract JSON from markdown code block
                lines = clean_response.split('\n')
                clean_response = '\n'.join([line for line in lines if not line.startswith('```')])
            
            result = json.loads(clean_response)
            logger.info(f"Successfully parsed AI response for item {item_id}: Summary field present, {len(result.get('Tasks', []))} tasks found")
            
            # Validate required fields
            if 'Summary' not in result:
                raise ValueError("Agent response missing 'Summary' field")
            if 'Tasks' not in result or not isinstance(result['Tasks'], list):
                # Allow missing or empty Tasks array
                logger.warning(f"AI response for item {item_id} missing Tasks field, defaulting to empty array")
                result['Tasks'] = []
                
        except (json.JSONDecodeError, ValueError) as e:
            logger.error(f"Failed to parse agent response for item {item_id}: {e}\nResponse preview: {agent_response[:200]}...")
            return JsonResponse({
                'success': False,
                'error': 'Failed to process the transcript. The AI response was invalid.'
            }, status=500)
        
        # Update meeting description with summary
        logger.debug(f"Updating item {item_id} with meeting summary and tasks...")
        with transaction.atomic():
            item.description = result['Summary']
            item.save()
            
            # Log description update
            activity_service.log(
                verb='item.updated',
                target=item,
                actor=request.user if request.user.is_authenticated else None,
                summary="Updated description from meeting transcript",
            )
            
            # Create task items
            tasks_created = 0
            for task_data in result['Tasks']:
                if 'Title' not in task_data:
                    logger.warning(f"Skipping task without Title: {task_data}")
                    continue
                
                # Get Task item type
                try:
                    task_type = ItemType.objects.get(key='task')
                except ItemType.DoesNotExist:
                    logger.error("ItemType 'task' not found in database")
                    continue
                
                # Create child task item
                task_item = Item.objects.create(
                    project=item.project,
                    parent=item,
                    type=task_type,
                    title=task_data['Title'],
                    description=task_data.get('Description', ''),
                    status=ItemStatus.INBOX,
                    assigned_to=request.user if request.user.is_authenticated else None,
                    requester=None,
                )
                
                # Log task creation
                activity_service.log(
                    verb='item.created',
                    target=task_item,
                    actor=request.user if request.user.is_authenticated else None,
                    summary=f"Created from meeting transcript",
                )
                
                tasks_created += 1
        
        logger.info(f"Transcript processing completed for item {item_id}: Created {tasks_created} task(s)")
        return JsonResponse({
            'success': True,
            'message': f'Transcript processed successfully. Created {tasks_created} task(s).',
            'summary': result['Summary'],
            'tasks_created': tasks_created
        })
    
    except AttachmentTooLarge as e:
        logger.warning(f"Transcript upload failed for item {item_id}: File too large - {str(e)} (file size: {file_size_mb:.2f} MB)")
        # Extract file size info from exception or provide default message
        error_msg = str(e)
        # Provide user-friendly German error message
        if 'MB' in error_msg and 'exceeds' in error_msg:
            # Parse the sizes from the exception message
            sizes = re.findall(r'(\d+\.?\d*)MB', error_msg)
            if len(sizes) >= 2:
                actual_size = sizes[0]
                max_size = sizes[1]
                error_msg = f'Datei zu groß ({actual_size} MB). Maximum: {max_size} MB'
            else:
                error_msg = 'Datei zu groß. Maximum: 50 MB'
        else:
            error_msg = 'Datei zu groß. Maximum: 50 MB'
        
        return JsonResponse({
            'success': False,
            'error': error_msg
        }, status=413)  # 413 Payload Too Large
        
    except Exception as e:
        logger.error(f"Transcript processing failed for item {item_id}: {type(e).__name__}: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': f'Failed to process transcript: {str(e)}'
        }, status=500)


@login_required
def item_view_attachment(request, attachment_id):
    """View an attachment (for viewable file types)."""
    try:
        attachment = get_object_or_404(Attachment, id=attachment_id)
        
        if attachment.is_deleted:
            return JsonResponse({'success': False, 'error': 'Attachment not found'}, status=404)
        
        # Get file extension using os.path.splitext for reliability
        import os
        _, extension = os.path.splitext(attachment.original_name.lower())
        extension = extension.lstrip('.')  # Remove leading dot
        
        # Read file content
        storage_service = AttachmentStorageService()
        file_content = storage_service.read_attachment(attachment)
        
        # Process based on file type
        if extension == 'md':
            # Render markdown to HTML - create parser instance per request for thread safety
            md_parser = markdown.Markdown(extensions=['extra', 'fenced_code'])
            html_content = md_parser.convert(file_content.decode('utf-8', errors='replace'))
            # Sanitize HTML
            clean_html = bleach.clean(
                html_content,
                tags=ALLOWED_TAGS,
                attributes=ALLOWED_ATTRIBUTES,
                strip=True
            )
            return JsonResponse({'success': True, 'content_html': clean_html})
        elif extension == 'pdf':
            # Return base64 encoded PDF
            import base64
            pdf_base64 = base64.b64encode(file_content).decode('utf-8')
            return JsonResponse({'success': True, 'content_base64': pdf_base64})
        elif extension in ['html', 'htm']:
            # Return sanitized HTML content (will be displayed in iframe)
            html_content = file_content.decode('utf-8', errors='replace')
            # Sanitize HTML before returning
            clean_html = bleach.clean(
                html_content,
                tags=ALLOWED_TAGS + ['html', 'head', 'body', 'meta', 'title', 'style'],
                attributes=ALLOWED_ATTRIBUTES,
                strip=True
            )
            return JsonResponse({'success': True, 'content': clean_html})
        else:
            # Plain text
            return JsonResponse({'success': True, 'content': file_content.decode('utf-8', errors='replace')})
            
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
def item_download_attachment(request, attachment_id):
    """Download an attachment."""
    try:
        attachment = get_object_or_404(Attachment, id=attachment_id)
        
        if attachment.is_deleted:
            return HttpResponse("Attachment not found", status=404)
        
        # Read file content
        storage_service = AttachmentStorageService()
        file_content = storage_service.read_attachment(attachment)
        
        # Create response with file
        response = HttpResponse(file_content, content_type=attachment.content_type or 'application/octet-stream')
        response['Content-Disposition'] = f'attachment; filename="{attachment.original_name}"'
        response['Content-Length'] = len(file_content)
        
        return response
    except Exception as e:
        return HttpResponse(f"Download failed: {str(e)}", status=500)


@login_required
def attachment_view(request, attachment_id):
    """
    Generic attachment view handler.
    This view determines the parent of an attachment and redirects to the appropriate
    item or project attachment view, or serves the attachment directly.
    """
    try:
        attachment = get_object_or_404(Attachment, id=attachment_id)
        
        if attachment.is_deleted:
            return HttpResponse("Attachment not found", status=404)
        
        # Try to determine parent via AttachmentLink
        first_link = attachment.links.first()
        if first_link:
            # Check the parent type by model name to avoid ContentType lookups
            target = first_link.target
            if target:
                model_name = target.__class__.__name__
                if model_name == 'Item':
                    # Redirect to item detail
                    return redirect('item-detail', item_id=first_link.target_object_id)
                elif model_name == 'Project':
                    # Redirect to project detail
                    return redirect('project-detail', id=first_link.target_object_id)
        
        # If no parent found or parent is not item/project, serve attachment directly
        # Get file extension
        _, extension = os.path.splitext(attachment.original_name.lower())
        extension = extension.lstrip('.')
        
        # Read file content
        storage_service = AttachmentStorageService()
        file_content = storage_service.read_attachment(attachment)
        
        # Determine content type and response
        if extension in ['md', 'txt', 'html', 'htm', 'pdf']:
            # For viewable types, return content for display
            if extension == 'md':
                md_parser = markdown.Markdown(extensions=['extra', 'fenced_code'])
                html_content = md_parser.convert(file_content.decode('utf-8', errors='replace'))
                clean_html = bleach.clean(
                    html_content,
                    tags=ALLOWED_TAGS,
                    attributes=ALLOWED_ATTRIBUTES,
                    strip=True
                )
                # Return as HTML page
                return HttpResponse(clean_html, content_type='text/html')
            elif extension == 'pdf':
                return HttpResponse(file_content, content_type='application/pdf')
            elif extension in ['html', 'htm']:
                html_content = file_content.decode('utf-8', errors='replace')
                clean_html = bleach.clean(
                    html_content,
                    tags=ALLOWED_TAGS + ['html', 'head', 'body', 'meta', 'title', 'style'],
                    attributes=ALLOWED_ATTRIBUTES,
                    strip=True
                )
                return HttpResponse(clean_html, content_type='text/html')
            else:  # txt
                return HttpResponse(file_content.decode('utf-8', errors='replace'), content_type='text/plain')
        else:
            # For other file types, trigger download
            response = HttpResponse(file_content, content_type=attachment.content_type or 'application/octet-stream')
            response['Content-Disposition'] = f'attachment; filename="{attachment.original_name}"'
            response['Content-Length'] = len(file_content)
            return response
            
    except Http404:
        # Re-raise Http404 to let Django handle it properly
        raise
    except Exception as e:
        logger.error(f"Error viewing attachment {attachment_id}: {str(e)}")
        return HttpResponse(f"Error viewing attachment: {str(e)}", status=500)


@login_required
@require_POST
def item_classify(request, item_id):
    """
    HTMX endpoint to classify an inbox item.
    
    Expects POST parameter 'action' with value 'backlog' or 'start'.
    Returns updated table row HTML.
    """
    item = get_object_or_404(Item, id=item_id)
    action = request.POST.get('action')
    
    if not action:
        return HttpResponse(
            "Missing 'action' parameter. Valid actions: 'backlog', 'start'", 
            status=400
        )
    
    try:
        # Use workflow guard to transition
        guard = ItemWorkflowGuard()
        guard.classify_inbox(item, action, actor=request.user if request.user.is_authenticated else None)
        
        # Return empty response with HX-Trigger to remove row
        # The row will be hidden by HTMX swap
        response = HttpResponse("")
        response['HX-Trigger'] = 'itemClassified'
        return response
        
    except ValidationError as e:
        return HttpResponse(str(e), status=400)


def _get_user_primary_organisation(user):
    """
    Helper function to get the primary organisation of a user.
    
    Args:
        user: User object
        
    Returns:
        Organisation object if user has a primary organisation, None otherwise
    """
    if not user or not user.is_authenticated:
        return None
    
    primary_org = UserOrganisation.objects.filter(
        user=user, 
        is_primary=True
    ).select_related('organisation').first()
    
    return primary_org.organisation if primary_org else None


def _auto_generate_title_from_description(description, user=None):
    """
    Generate a title from description using AI agent if description is provided.
    Returns the generated title or empty string if generation fails or description is empty.
    """
    if not description or not description.strip():
        return ''
    
    try:
        agent_service = AgentService()
        title = agent_service.execute_agent(
            filename='text-to-title-generator.yml',
            input_text=description,
            user=user,
            client_ip=None
        )
        # Clean up the title (remove quotes, newlines, etc.)
        title = title.strip().strip('"').strip("'").replace('\n', ' ').replace('\r', '')
        return title
    except Exception as e:
        logger.error(f"Failed to auto-generate title: {str(e)}")
        return ''


def _send_responsible_notification(item, new_responsible):
    """
    Send email notification to new responsible user.
    Uses mail template with key 'resp'.
    
    Args:
        item: Item instance
        new_responsible: User instance who is now responsible
    """
    try:
        from .services.graph.mail_service import send_email
        
        # Get the mail template
        template = MailTemplate.objects.filter(key='resp', is_active=True).first()
        if not template:
            logger.warning("Mail template 'resp' not found or inactive")
            return
        
        # Get base URL from GlobalSettings
        settings = GlobalSettings.get_instance()
        base_url = settings.base_url if settings and settings.base_url else 'http://localhost:8000'
        
        # Build absolute link to item detail
        item_link = f"{base_url.rstrip('/')}/items/{item.id}/"
        
        # Prepare template context
        context = {
            'issue': {
                'title': item.title,
                'type': item.type.name if item.type else '—',
                'project': item.project.name if item.project else '—',
                'responsible': new_responsible.name if new_responsible else '—',
                'assigned_to': item.assigned_to.name if item.assigned_to else '',
                'requester': item.requester.name if item.requester else '—',
                'status': item.get_status_display(),
                'link': item_link,
            }
        }
        
        # Render template
        from django.template import Context, Template
        subject_template = Template(template.subject)
        message_template = Template(template.message)
        
        subject = subject_template.render(Context(context))
        message = message_template.render(Context(context))
        
        # Send email
        send_email(
            subject=subject,
            body=message,
            to=[new_responsible.email],
            body_is_html=True
        )
        
        logger.info(f"Sent responsible notification to {new_responsible.email} for item {item.id}")
    except Exception as e:
        logger.error(f"Failed to send responsible notification: {e}")


def _update_item_followers(item, follower_ids):
    """
    Update followers for an item.
    
    Args:
        item: Item instance to update followers for
        follower_ids: List of user IDs to set as followers
        
    This function atomically replaces all followers with the provided list.
    Validates that all user IDs are valid and prevents duplicates.
    """
    from core.models import ItemFollower
    
    if not follower_ids:
        follower_ids = []
    
    with transaction.atomic():
        # Remove existing followers
        ItemFollower.objects.filter(item=item).delete()
        
        # Add new followers
        seen_user_ids = set()
        for user_id in follower_ids:
            if not user_id or user_id in seen_user_ids:
                continue
                
            try:
                # Validate user ID is an integer
                user_id_int = int(user_id)
                # Try to get the user
                user = User.objects.filter(id=user_id_int).first()
                if user:
                    ItemFollower.objects.create(item=item, user=user)
                    seen_user_ids.add(user_id)
                else:
                    logger.warning(f"Invalid user ID {user_id} for item {item.id} followers")
            except (ValueError, TypeError) as e:
                logger.warning(f"Invalid user ID format {user_id} for item {item.id} followers: {e}")


@login_required
def item_create(request):
    """Item create page view."""
    if request.method == 'GET':
        # Show the create form
        from .models import IssueBlueprint
        
        projects = Project.objects.all().order_by('name')
        item_types = ItemType.objects.filter(is_active=True).order_by('name')
        organisations = Organisation.objects.all().order_by('name')
        # Prefetch user organizations for efficient org short code lookup in template
        users = User.objects.prefetch_related(
            'user_organisations__organisation'
        ).all().order_by('name')
        # Filter agents for responsible field
        agents = User.objects.filter(role=UserRole.AGENT).order_by('name')
        statuses = ItemStatus.choices
        
        # Get active blueprints for the create form
        blueprints = IssueBlueprint.objects.filter(is_active=True).select_related('category').order_by('-updated_at')
        
        # Auto-populate default values for requester and organisation
        default_requester = None
        default_organisation = None
        if request.user.is_authenticated:
            default_requester = request.user
            default_organisation = _get_user_primary_organisation(request.user)
        
        # Support pre-selecting project via query parameter
        default_project = None
        project_id = request.GET.get('project')
        nodes = []
        if project_id:
            try:
                # Validate and convert project_id to integer
                project_id_int = int(project_id)
                default_project = Project.objects.get(id=project_id_int)
                # Get nodes for the default project
                nodes = Node.objects.filter(project=default_project).order_by('name')
            except (ValueError, Project.DoesNotExist, TypeError):
                # Invalid or non-existent project ID, ignore it gracefully
                pass
        
        context = {
            'item': None,
            'projects': projects,
            'item_types': item_types,
            'organisations': organisations,
            'users': users,
            'agents': agents,
            'statuses': statuses,
            'default_requester': default_requester,
            'default_organisation': default_organisation,
            'default_project': default_project,
            'nodes': nodes,
            'blueprints': blueprints,
        }
        return render(request, 'item_form.html', context)
    
    # Handle POST request (HTMX form submission)
    try:
        from .models import IssueBlueprint
        from .utils.blueprint_variables import replace_variables
        import json
        
        project_id = request.POST.get('project')
        project = get_object_or_404(Project, id=project_id)
        
        type_id = request.POST.get('type')
        item_type = get_object_or_404(ItemType, id=type_id)
        
        # Get title and description
        title = request.POST.get('title', '').strip()
        description = request.POST.get('description', '')
        user_input = request.POST.get('user_input', '')
        
        # Apply blueprint if selected
        blueprint_id = request.POST.get('blueprint')
        if blueprint_id:
            try:
                blueprint = IssueBlueprint.objects.get(id=blueprint_id, is_active=True)
                
                # Only process variables if we'll use the blueprint content
                if not title or not description:
                    # Parse blueprint variables
                    blueprint_variables_json = request.POST.get('blueprint_variables', '{}')
                    try:
                        blueprint_variables = json.loads(blueprint_variables_json)
                    except json.JSONDecodeError:
                        blueprint_variables = {}
                    
                    # Replace variables in blueprint content
                    if not title:  # Only use blueprint title if no title provided
                        title = replace_variables(blueprint.title, blueprint_variables)
                    if not description:  # Only use blueprint description if no description provided
                        description = replace_variables(blueprint.description_md, blueprint_variables)
                
            except IssueBlueprint.DoesNotExist:
                pass  # Ignore if blueprint doesn't exist or is inactive
        
        # Auto-generate title from description if title is empty
        if not title and description:
            title = _auto_generate_title_from_description(description, request.user if request.user.is_authenticated else None)
        
        # Create the item
        item = Item(
            project=project,
            title=title,
            description=description,
            user_input=user_input,
            solution_description=request.POST.get('solution_description', ''),
            type=item_type,
            status=request.POST.get('status', ItemStatus.INBOX),
        )
        
        # Set optional fields with automatic pre-population
        # Auto-populate requester with current user if not explicitly provided
        requester_id = request.POST.get('requester')
        if requester_id:
            item.requester = get_object_or_404(User, id=requester_id)
        elif request.user.is_authenticated:
            # Auto-set requester to current user if not provided
            item.requester = request.user
        
        # Auto-populate organisation with user's primary organisation if not explicitly provided
        org_id = request.POST.get('organisation')
        if org_id:
            item.organisation = get_object_or_404(Organisation, id=org_id)
        elif request.user.is_authenticated:
            item.organisation = _get_user_primary_organisation(request.user)
        
        assigned_to_id = request.POST.get('assigned_to')
        if assigned_to_id:
            item.assigned_to = get_object_or_404(User, id=assigned_to_id)
        
        responsible_id = request.POST.get('responsible')
        if responsible_id:
            responsible_user = get_object_or_404(User, id=responsible_id)
            # Validate that user has Agent role
            if responsible_user.role != UserRole.AGENT:
                raise ValidationError({'responsible': 'Responsible user must have role "Agent".'})
            item.responsible = responsible_user
        
        parent_id = request.POST.get('parent')
        if parent_id:
            item.parent = get_object_or_404(Item, id=parent_id)
        
        solution_release_id = request.POST.get('solution_release')
        if solution_release_id:
            item.solution_release = get_object_or_404(Release, id=solution_release_id)
        
        # Handle node selection and update description before saving
        node_id = request.POST.get('node')
        if node_id:
            node = get_object_or_404(Node, id=node_id, project=item.project)
            # Update description with breadcrumb before saving
            item.update_description_with_breadcrumb(node)
        
        # Save item once with all changes
        item.save()
        
        # Update nodes relationship (ManyToMany, requires item to be saved first)
        if node_id:
            node = get_object_or_404(Node, id=node_id, project=item.project)
            item.nodes.add(node)
            # Validate nodes belong to project
            item.validate_nodes()
        
        # Handle followers (ManyToMany, requires item to be saved first)
        follower_ids = request.POST.getlist('follower_ids')
        _update_item_followers(item, follower_ids)
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='item.created',
            target=item,
            actor=request.user if request.user.is_authenticated else None,
            summary=f"Created item: {item.title}",
        )
        
        # Check for mail trigger
        response_data = {
            'success': True,
            'message': 'Item created successfully',
            'redirect': f'/items/{item.id}/',
            'item_id': item.id,
            'followers': list(item.get_followers().values('id', 'username', 'email', 'name'))
        }
        
        mapping = check_mail_trigger(item)
        if mapping:
            # Prepare mail preview for modal
            mail_preview = prepare_mail_preview(item, mapping)
            response_data['mail_preview'] = mail_preview
        
        return JsonResponse(response_data)
    except ValidationError as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)
    except Exception as e:
        # Log the full error for debugging
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"Item creation failed: {str(e)}")
        return JsonResponse({'success': False, 'error': 'Failed to create item. Please check your input.'}, status=400)


@login_required
def item_edit(request, item_id):
    """Item edit page view."""
    item = get_object_or_404(
        Item.objects.select_related(
            'project', 'type', 'organisation', 'requester', 
            'assigned_to', 'solution_release', 'parent'
        ).prefetch_related('nodes'),
        id=item_id
    )
    
    if request.method == 'GET':
        # Show the edit form
        projects = Project.objects.all().order_by('name')
        item_types = ItemType.objects.filter(is_active=True).order_by('name')
        organisations = Organisation.objects.all().order_by('name')
        # Prefetch user organizations for efficient org short code lookup in template
        users = User.objects.prefetch_related(
            'user_organisations__organisation'
        ).all().order_by('name')
        # Filter agents for responsible field
        agents = User.objects.filter(role=UserRole.AGENT).order_by('name')
        statuses = ItemStatus.choices
        
        # Get releases for the current project
        releases = Release.objects.filter(project=item.project).order_by('-version')
        
        # Get potential parent items, exclude closed and self
        # Filter as per issue #352 - allow items from all projects, status != closed
        parent_items = Item.objects.exclude(status=ItemStatus.CLOSED).exclude(id=item.id).order_by('title')
        
        # Get nodes from the current project
        nodes = Node.objects.filter(project=item.project).order_by('name')
        
        context = {
            'item': item,
            'projects': projects,
            'item_types': item_types,
            'organisations': organisations,
            'users': users,
            'agents': agents,
            'statuses': statuses,
            'releases': releases,
            'parent_items': parent_items,
            'nodes': nodes,
        }
        return render(request, 'item_form.html', context)
    
    # Handle POST request (HTMX form submission) - handled by item_update
    return redirect('item-update', item_id=item_id)


@login_required

@require_http_methods(["POST"])
def item_update(request, item_id):
    """Update item details."""
    item = get_object_or_404(Item, id=item_id)
    
    # Capture old status to detect changes
    old_status = item.status
    
    try:
        # Check if node is being updated
        node_id = request.POST.get('node')
        node_changed = node_id is not None
        
        # Get title and description
        title = request.POST.get('title', '').strip()
        description = request.POST.get('description', item.description)
        user_input = request.POST.get('user_input', item.user_input)
        
        # Auto-generate title from description if title is empty
        if not title and description:
            title = _auto_generate_title_from_description(description, request.user if request.user.is_authenticated else None)
        
        # Update basic fields
        item.title = title if title else item.title
        # Don't update description from POST if we're changing node - we'll handle it separately
        if not node_changed:
            item.description = description
        else:
            # User might have edited description in form - we need to extract non-breadcrumb content
            # and then re-add the new breadcrumb
            posted_description = description
            # Temporarily set description to extract content without breadcrumb
            item.description = posted_description
        
        item.user_input = user_input
        item.solution_description = request.POST.get('solution_description', item.solution_description)
        item.short_description = request.POST.get('short_description', item.short_description)
        item.status = request.POST.get('status', item.status)
        
        # Update boolean fields
        item.intern = request.POST.get('intern') == 'on' or request.POST.get('intern') == 'true'
        
        # Update foreign key fields
        project_id = request.POST.get('project')
        if project_id:
            item.project = get_object_or_404(Project, id=project_id)
        
        type_id = request.POST.get('type')
        if type_id:
            item.type = get_object_or_404(ItemType, id=type_id)
        
        # Update optional foreign key fields
        org_id = request.POST.get('organisation')
        if org_id:
            item.organisation = get_object_or_404(Organisation, id=org_id)
        elif org_id == '':  # Empty string means clear the field
            item.organisation = None
        
        requester_id = request.POST.get('requester')
        if requester_id:
            item.requester = get_object_or_404(User, id=requester_id)
        elif requester_id == '':
            item.requester = None
        
        assigned_to_id = request.POST.get('assigned_to')
        if assigned_to_id:
            item.assigned_to = get_object_or_404(User, id=assigned_to_id)
        elif assigned_to_id == '':
            item.assigned_to = None
        
        responsible_id = request.POST.get('responsible')
        if responsible_id:
            responsible_user = get_object_or_404(User, id=responsible_id)
            # Validate that user has Agent role
            if responsible_user.role != UserRole.AGENT:
                raise ValidationError({'responsible': 'Responsible user must have role "Agent".'})
            item.responsible = responsible_user
        elif responsible_id == '':
            item.responsible = None
        
        parent_id = request.POST.get('parent')
        if parent_id:
            item.parent = get_object_or_404(Item, id=parent_id)
        elif parent_id == '':
            item.parent = None
        
        solution_release_id = request.POST.get('solution_release')
        if solution_release_id:
            item.solution_release = get_object_or_404(Release, id=solution_release_id)
        elif solution_release_id == '':
            item.solution_release = None
        
        # Handle node selection before first save
        if node_changed:
            if node_id:
                node = get_object_or_404(Node, id=node_id, project=item.project)
                # Update description with breadcrumb before saving
                item.update_description_with_breadcrumb(node)
            else:
                # Clear breadcrumb
                item.update_description_with_breadcrumb(None)
        
        # Save item once with all changes
        item.save()
        
        # Update nodes relationship (ManyToMany, requires item to be saved first)
        if node_changed:
            if node_id:
                node = get_object_or_404(Node, id=node_id, project=item.project)
                item.nodes.clear()
                item.nodes.add(node)
                # Validate nodes belong to project
                item.validate_nodes()
            else:
                item.nodes.clear()
        
        # Handle followers (ManyToMany, requires item to be saved first)
        # Only update if follower_ids was provided in the request
        if 'follower_ids' in request.POST:
            follower_ids = request.POST.getlist('follower_ids')
            _update_item_followers(item, follower_ids)
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='item.updated',
            target=item,
            actor=request.user if request.user.is_authenticated else None,
            summary=f"Updated item: {item.title}",
        )
        
        # Check for mail trigger (only if status changed)
        response_data = {
            'success': True,
            'message': 'Item updated successfully',
            'item_id': item.id,
            'redirect': f'/items/{item.id}/',
            'followers': list(item.get_followers().values('id', 'username', 'email', 'name'))
        }
        
        # Only check for mail trigger if status changed
        if item.status != old_status:
            mapping = check_mail_trigger(item)
            if mapping:
                # Prepare mail preview for modal
                mail_preview = prepare_mail_preview(item, mapping)
                response_data['mail_preview'] = mail_preview
        
        return JsonResponse(response_data)
    except ValidationError as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)
    except Exception as e:
        # Log the full error for debugging
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"Item update failed for item {item_id}: {str(e)}")
        return JsonResponse({'success': False, 'error': 'Failed to update item. Please check your input.'}, status=400)


@login_required

@require_http_methods(["POST"])
def item_delete(request, item_id):
    """Delete an item."""
    item = get_object_or_404(Item, id=item_id)
    project_id = item.project.id
    
    try:
        item.delete()
        return JsonResponse({
            'success': True,
            'message': 'Item deleted successfully',
            'redirect': f'/projects/{project_id}/'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required
@require_POST
def item_send_status_mail(request, item_id):
    """Send status change email for an item."""
    from .services.graph.mail_service import send_email
    from .services.mail import get_notification_recipients_for_item
    from core.utils.html_sanitization import sanitize_html
    
    item = get_object_or_404(Item, id=item_id)
    
    # Authorization check: user must be requester, assignee, or have permission to edit item
    # For now, we allow anyone who can view the item (authenticated users)
    # In a more restrictive scenario, check project membership or specific permissions
    if not request.user.is_authenticated:
        return JsonResponse({'success': False, 'error': 'Authentication required'}, status=401)
    
    try:
        data = json.loads(request.body)
        
        # Extract and validate mail data from request
        subject = data.get('subject', '').strip()
        message = data.get('message', '').strip()
        to_address = data.get('to', '').strip()
        from_address = data.get('from_address', '').strip()
        cc_address = data.get('cc_address', '').strip()
        
        if not subject or not message:
            return JsonResponse({'success': False, 'error': 'Subject and message are required'}, status=400)
        
        # Sanitize HTML message to prevent XSS attacks while preserving formatting
        # Uses centralized sanitization with support for inline styles and CSS properties
        message = sanitize_html(message)
        
        if not to_address:
            # Try to get recipient from requester
            if item.requester and item.requester.email:
                to_address = item.requester.email
            else:
                return JsonResponse({'success': False, 'error': 'No recipient email available'}, status=400)
        
        # Get follower emails for CC using the utility function
        recipients = get_notification_recipients_for_item(item)
        follower_emails = recipients['cc']
        
        # Prepare recipient list
        to = [to_address] if isinstance(to_address, str) else to_address
        
        # Merge follower emails with any existing CC addresses from the request
        cc_list = []
        if cc_address:
            if isinstance(cc_address, str):
                cc_list.append(cc_address)
            else:
                cc_list.extend(cc_address)
        
        # Add follower emails to CC, avoiding duplicates
        seen = set(cc_list)
        for email in follower_emails:
            if email and email not in seen:
                cc_list.append(email)
                seen.add(email)
        
        cc = cc_list if cc_list else None
        
        # Send email using graph service
        result = send_email(
            subject=subject,
            body=message,
            to=to,
            body_is_html=True,
            cc=cc,
            sender=from_address if from_address else None,
            item=item,
            author=request.user if request.user.is_authenticated else None,
            visibility='Internal',
            client_ip=request.META.get('REMOTE_ADDR')
        )
        
        if result.success:
            return JsonResponse({
                'success': True,
                'message': 'Email sent successfully'
            })
        else:
            return JsonResponse({
                'success': False,
                'error': result.error or 'Failed to send email'
            }, status=500)
            
    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'error': 'Invalid JSON data'}, status=400)
    except Exception as e:
        logger.error(f"Failed to send status mail for item {item_id}: {str(e)}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
@require_POST
def item_send_status_update(request, item_id):
    """Manually send a status-update email for an item using the 'featurenew' mail template."""
    from .services.graph.mail_service import send_email
    from .services.mail import get_notification_recipients_for_item
    from .services.mail.template_processor import process_template

    item = get_object_or_404(Item, id=item_id)

    try:
        # Load the 'featurenew' mail template (must exist and be active)
        try:
            template = MailTemplate.objects.get(key='featurenew', is_active=True)
        except MailTemplate.DoesNotExist:
            return JsonResponse(
                {'ok': False, 'error': "Mail template 'featurenew' not found or inactive"},
                status=404,
            )

        # Process template variables
        processed = process_template(template, item)
        subject = processed['subject']
        message = processed['message']

        # Determine recipients
        recipients = get_notification_recipients_for_item(item)
        to_address = recipients.get('to')
        cc_list = recipients.get('cc') or []

        if not to_address:
            return JsonResponse(
                {'ok': False, 'error': 'No recipient email available (requester has no email)'},
                status=400,
            )

        # Send email via Graph mail service
        result = send_email(
            subject=subject,
            body=message,
            to=[to_address],
            body_is_html=True,
            cc=cc_list if cc_list else None,
            sender=template.from_address if template.from_address else None,
            item=item,
            author=request.user,
            visibility='Internal',
            client_ip=request.META.get('REMOTE_ADDR'),
        )

        if not result.success:
            return JsonResponse(
                {'ok': False, 'error': result.error or 'Failed to send email'},
                status=500,
            )

        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='item.manual_status_update_sent',
            target=item,
            actor=request.user,
            summary='Manual status update sent',
        )

        return JsonResponse({'ok': True, 'message': 'Status update sent successfully'})

    except Exception as e:
        logger.error(f"Failed to send status update for item {item_id}: {str(e)}")
        return JsonResponse({'ok': False, 'error': str(e)}, status=500)


@login_required
@require_POST
def item_move_project(request, item_id):
    """Move item to a different project with optional email notification."""
    from .services.graph.mail_service import send_email
    from .services.mail import get_notification_recipients_for_item
    from .services.mail.template_processor import process_template
    
    item = get_object_or_404(Item, id=item_id)
    
    try:
        data = json.loads(request.body)
        
        # Extract parameters
        target_project_id = data.get('target_project_id')
        send_mail_to_requester = data.get('send_mail_to_requester', True)
        
        if not target_project_id:
            return JsonResponse({'success': False, 'error': 'Target project is required'}, status=400)
        
        # Get target project
        target_project = get_object_or_404(Project, id=target_project_id)
        
        # Check if the project is actually changing
        if item.project.id == target_project.id:
            return JsonResponse({'success': False, 'error': 'Item is already in the target project'}, status=400)
        
        # Store old project for logging
        old_project = item.project
        
        with transaction.atomic():
            # Update item project
            item.project = target_project
            
            # Clear project-dependent fields that may not be valid in new project
            # Clear nodes (they are project-specific)
            item.nodes.clear()
            
            # Clear parent if it's in a different project
            if item.parent and item.parent.project != target_project:
                item.parent = None
            
            # Clear solution_release if it belongs to different project
            if item.solution_release and item.solution_release.project != target_project:
                item.solution_release = None
            
            # Clear organisation if not a client of the new project
            if item.organisation:
                project_clients = list(target_project.clients.all())
                # Clear organisation if project has clients and organisation is not one of them,
                # or if project has no clients at all
                if not project_clients or item.organisation not in project_clients:
                    item.organisation = None
            
            # Save the item
            item.save()
            
            # Log activity
            activity_service = ActivityService()
            activity_service.log(
                verb='item.moved',
                target=item,
                actor=request.user if request.user.is_authenticated else None,
                summary=f"Moved from {old_project.name} to {target_project.name}",
            )
        
        # Send email notification if requested
        mail_sent = False
        mail_error = None
        
        if send_mail_to_requester:
            try:
                # Get the mail template with key 'moved'
                template = MailTemplate.objects.filter(key='moved', is_active=True).first()
                
                if template and item.requester and item.requester.email:
                    # Reload item with all necessary relations for template processing
                    item = Item.objects.select_related(
                        'project', 'requester', 'assigned_to', 'type', 'solution_release'
                    ).prefetch_related(
                        'requester__user_organisations__organisation'
                    ).get(id=item.id)
                    
                    # Process template with updated item data
                    processed = process_template(template, item)
                    
                    # Get recipients
                    recipients = get_notification_recipients_for_item(item)
                    
                    if recipients['to']:
                        # Send email
                        result = send_email(
                            subject=processed['subject'],
                            body=processed['message'],
                            to=[recipients['to']],
                            body_is_html=True,
                            cc=recipients['cc'] if recipients['cc'] else None,
                            sender=template.from_address if template.from_address else None,
                            item=item,
                            author=request.user,
                            visibility='Internal',
                            client_ip=request.META.get('REMOTE_ADDR')
                        )
                        
                        if result.success:
                            mail_sent = True
                            logger.info(
                                f"Move notification email sent for item {item.id} "
                                f"from project '{old_project.name}' to '{target_project.name}'"
                            )
                        else:
                            mail_error = result.error
                            logger.error(
                                f"Failed to send move notification for item {item.id}: "
                                f"Template: 'moved', Requester: {item.requester.email}, Error: {result.error}"
                            )
                    else:
                        mail_error = "No recipient email available"
                        logger.warning(f"No recipient email for move notification of item {item.id}")
                elif not template:
                    mail_error = "Mail template 'moved' not found or inactive"
                    logger.warning("Mail template with key 'moved' not found or inactive")
                elif not item.requester:
                    mail_error = "Item has no requester"
                elif not item.requester.email:
                    mail_error = "Requester has no email address"
                    
            except Exception as e:
                mail_error = str(e)
                logger.error(
                    f"Exception while sending move notification for item {item.id}: {str(e)}"
                )
        
        # Prepare response
        response_data = {
            'success': True,
            'message': f'Item moved to {target_project.name}',
            'item_id': item.id,
            'new_project_id': target_project.id,
            'new_project_name': target_project.name,
            'mail_sent': mail_sent,
        }
        
        if mail_error:
            response_data['mail_error'] = mail_error
        
        return JsonResponse(response_data)
        
    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'error': 'Invalid JSON data'}, status=400)
    except Http404:
        return JsonResponse({'success': False, 'error': 'Target project not found'}, status=404)
    except ValidationError as e:
        logger.warning(f"Validation error moving item {item_id}: {str(e)}")
        return JsonResponse({'success': False, 'error': str(e)}, status=400)
    except IntegrityError as e:
        logger.error(f"Database integrity error moving item {item_id}: {str(e)}")
        return JsonResponse({
            'success': False, 
            'error': 'Failed to move item due to database constraints. Please contact support.'
        }, status=500)
    except Exception as e:
        logger.error(f"Failed to move item {item_id}: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False, 
            'error': 'An unexpected error occurred while moving the item. Please try again or contact support.'
        }, status=500)


@login_required

@require_POST
def ai_generate_title(request):
    """Generate a title from description using AI agent."""
    import json
    
    try:
        data = json.loads(request.body)
        text = data.get('text', '')
        
        if not text:
            return JsonResponse({'success': False, 'error': 'No text provided'}, status=400)
        
        # Use AgentService to execute the text-to-title-generator agent
        agent_service = AgentService()
        title = agent_service.execute_agent(
            filename='text-to-title-generator.yml',
            input_text=text,
            user=request.user if request.user.is_authenticated else None,
            client_ip=request.META.get('REMOTE_ADDR')
        )
        
        # Clean up the title (remove quotes, newlines, etc.)
        title = title.strip().strip('"').strip("'").replace('\n', ' ').replace('\r', '')
        
        return JsonResponse({'success': True, 'title': title})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required

@require_POST
def ai_optimize_text(request):
    """Optimize text using AI agent."""
    import json
    
    try:
        data = json.loads(request.body)
        text = data.get('text', '')
        
        if not text:
            return JsonResponse({'success': False, 'error': 'No text provided'}, status=400)
        
        # Use AgentService to execute the text-optimization-agent
        agent_service = AgentService()
        optimized_text = agent_service.execute_agent(
            filename='text-optimization-agent.yml',
            input_text=text,
            user=request.user if request.user.is_authenticated else None,
            client_ip=request.META.get('REMOTE_ADDR')
        )
        
        return JsonResponse({'success': True, 'text': optimized_text})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


# Project CRUD operations
@login_required

@require_http_methods(["POST"])
def project_update(request, id):
    """Update project details."""
    project = get_object_or_404(Project, id=id)
    
    try:
        project.name = request.POST.get('name', project.name)
        project.description = request.POST.get('description', project.description)
        project.status = request.POST.get('status', project.status)
        project.github_owner = request.POST.get('github_owner', project.github_owner)
        project.github_repo = request.POST.get('github_repo', project.github_repo)
        project.save()
        return JsonResponse({'success': True, 'message': 'Project updated successfully'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required

@require_http_methods(["POST"])
def project_delete(request, id):
    """Delete a project."""
    project = get_object_or_404(Project, id=id)
    
    try:
        project.delete()
        return JsonResponse({'success': True, 'redirect': '/projects/'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required

@require_http_methods(["POST"])
def project_add_client(request, id):
    """Add a client (organisation) to a project."""
    project = get_object_or_404(Project, id=id)
    org_id = request.POST.get('organisation_id')
    
    if not org_id:
        return JsonResponse({'success': False, 'error': 'Organisation ID required'}, status=400)
    
    try:
        organisation = get_object_or_404(Organisation, id=org_id)
        project.clients.add(organisation)
        return JsonResponse({'success': True, 'message': 'Client added successfully'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required

@require_http_methods(["POST"])
def project_remove_client(request, id):
    """Remove a client (organisation) from a project."""
    project = get_object_or_404(Project, id=id)
    org_id = request.POST.get('organisation_id')
    
    if not org_id:
        return JsonResponse({'success': False, 'error': 'Organisation ID required'}, status=400)
    
    try:
        organisation = get_object_or_404(Organisation, id=org_id)
        project.clients.remove(organisation)
        return JsonResponse({'success': True, 'message': 'Client removed successfully'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required

@require_http_methods(["POST"])
def project_add_item(request, id):
    """Add a new item to a project."""
    project = get_object_or_404(Project, id=id)
    
    try:
        title = request.POST.get('title')
        description = request.POST.get('description', '')
        type_id = request.POST.get('type_id')
        
        if not title or not type_id:
            return JsonResponse({'success': False, 'error': 'Title and Type are required'}, status=400)
        
        item_type = get_object_or_404(ItemType, id=type_id)
        
        item = Item.objects.create(
            project=project,
            title=title,
            description=description,
            type=item_type,
            status=ItemStatus.INBOX
        )
        
        return JsonResponse({'success': True, 'message': 'Item created successfully', 'item_id': item.id})
    except ValidationError as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required

@require_http_methods(["POST"])
def project_add_node(request, id):
    """Add a new node to a project."""
    project = get_object_or_404(Project, id=id)
    
    try:
        name = request.POST.get('name')
        node_type = request.POST.get('type')
        description = request.POST.get('description', '')
        parent_node_id = request.POST.get('parent_node_id')
        
        if not name or not node_type:
            return JsonResponse({'success': False, 'error': 'Name and Type are required'}, status=400)
        
        # Handle parent node if specified
        parent_node = None
        if parent_node_id:
            try:
                parent_node = Node.objects.get(id=parent_node_id, project=project)
            except Node.DoesNotExist:
                return JsonResponse({'success': False, 'error': 'Invalid parent node'}, status=400)
        
        node = Node.objects.create(
            project=project,
            name=name,
            type=node_type,
            description=description,
            parent_node=parent_node
        )
        
        return JsonResponse({'success': True, 'message': 'Node created successfully', 'node_id': node.id})
    except ValidationError as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required
def project_node_detail(request, project_id, node_id):
    """Get details of a specific node as JSON."""
    project = get_object_or_404(Project, id=project_id)
    node = get_object_or_404(Node, id=node_id, project=project)
    
    # Get child nodes
    children = []
    for child in node.child_nodes.all():
        children.append({
            'id': child.id,
            'name': child.name,
            'type': child.type,
            'breadcrumb': child.get_breadcrumb()
        })
    
    data = {
        'id': node.id,
        'name': node.name,
        'type': node.type,
        'description': node.description,
        'parent_node_id': node.parent_node.id if node.parent_node else None,
        'parent_node_name': node.parent_node.name if node.parent_node else None,
        'breadcrumb': node.get_breadcrumb(),
        'children': children
    }
    
    return JsonResponse(data)


@login_required
@require_http_methods(["POST"])
def project_node_update(request, project_id, node_id):
    """Update a node."""
    project = get_object_or_404(Project, id=project_id)
    node = get_object_or_404(Node, id=node_id, project=project)
    
    try:
        name = request.POST.get('name')
        node_type = request.POST.get('type')
        description = request.POST.get('description', '')
        parent_node_id = request.POST.get('parent_node_id')
        
        if not name or not node_type:
            return JsonResponse({'success': False, 'error': 'Name and Type are required'}, status=400)
        
        # Handle parent node if specified
        parent_node = None
        if parent_node_id:
            try:
                parent_node = Node.objects.get(id=parent_node_id, project=project)
                
                # Check for circular references
                if node.would_create_cycle(parent_node):
                    return JsonResponse({
                        'success': False, 
                        'error': 'Cannot set parent: would create circular reference'
                    }, status=400)
                    
            except Node.DoesNotExist:
                return JsonResponse({'success': False, 'error': 'Invalid parent node'}, status=400)
        
        # Update node
        node.name = name
        node.type = node_type
        node.description = description
        node.parent_node = parent_node
        node.save()
        
        return JsonResponse({'success': True, 'message': 'Node updated successfully'})
    except ValidationError as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required
def project_nodes_tree(request, id):
    """Get the hierarchical tree structure of all nodes in a project."""
    project = get_object_or_404(Project, id=id)
    
    # Get all root nodes with optimized prefetching of children
    root_nodes = Node.objects.filter(
        project=project, 
        parent_node=None
    ).prefetch_related('child_nodes')
    
    # Build the tree using the model method
    tree = [root.get_tree_structure() for root in root_nodes.order_by('name')]
    
    return JsonResponse({'tree': tree})


@login_required

@require_http_methods(["POST"])
def project_add_release(request, id):
    """Add a new release to a project."""
    project = get_object_or_404(Project, id=id)
    
    try:
        name = request.POST.get('name')
        version = request.POST.get('version')
        release_type = request.POST.get('type')
        status = request.POST.get('status', ReleaseStatus.PLANNED)
        planned_date_str = request.POST.get('planned_date')
        
        if not name or not version:
            return JsonResponse({'success': False, 'error': 'Name and Version are required'}, status=400)
        
        # Validate release type if provided
        if release_type and release_type not in ReleaseType.values:
            return JsonResponse({'success': False, 'error': f'Invalid release type. Must be one of: {", ".join(ReleaseType.values)}'}, status=400)
        
        # Validate status
        if status and status not in ReleaseStatus.values:
            return JsonResponse({'success': False, 'error': f'Invalid status. Must be one of: {", ".join(ReleaseStatus.values)}'}, status=400)
        
        # Parse planned_date if provided
        planned_date = None
        if planned_date_str:
            try:
                from datetime import datetime
                planned_date = datetime.strptime(planned_date_str, '%Y-%m-%d').date()
            except ValueError:
                return JsonResponse({'success': False, 'error': 'Invalid date format. Use YYYY-MM-DD'}, status=400)
        
        release = Release.objects.create(
            project=project,
            name=name,
            version=version,
            type=release_type if release_type else None,
            status=status,
            planned_date=planned_date
        )
        
        return JsonResponse({'success': True, 'message': 'Release created successfully', 'release_id': release.id})
    except ValidationError as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required
def project_update_release(request, id, release_id):
    """Update a release."""
    project = get_object_or_404(Project, id=id)
    release = get_object_or_404(Release, id=release_id, project=project)
    
    try:
        name = request.POST.get('name')
        version = request.POST.get('version')
        release_type = request.POST.get('type')
        status = request.POST.get('status')
        planned_date_str = request.POST.get('planned_date')
        
        if not name or not version:
            return JsonResponse({'success': False, 'error': 'Name and Version are required'}, status=400)
        
        # Validate release type if provided
        if release_type and release_type not in ReleaseType.values:
            return JsonResponse({'success': False, 'error': f'Invalid release type. Must be one of: {", ".join(ReleaseType.values)}'}, status=400)
        
        # Validate status
        if status and status not in ReleaseStatus.values:
            return JsonResponse({'success': False, 'error': f'Invalid status. Must be one of: {", ".join(ReleaseStatus.values)}'}, status=400)
        
        # Parse planned_date if provided
        if planned_date_str:
            try:
                from datetime import datetime
                planned_date = datetime.strptime(planned_date_str, '%Y-%m-%d').date()
                release.planned_date = planned_date
            except ValueError:
                return JsonResponse({'success': False, 'error': 'Invalid date format. Use YYYY-MM-DD'}, status=400)
        elif planned_date_str == '':
            # Empty string means clear the date
            release.planned_date = None
        
        # Update release fields
        release.name = name
        release.version = version
        release.type = release_type if release_type else None
        release.status = status if status else release.status
        release.save()
        
        return JsonResponse({'success': True, 'message': 'Release updated successfully'})
    except ValidationError as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required
def project_delete_release(request, id, release_id):
    """Delete a release."""
    project = get_object_or_404(Project, id=id)
    release = get_object_or_404(Release, id=release_id, project=project)
    
    try:
        release.delete()
        return JsonResponse({'success': True, 'message': 'Release deleted successfully'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required
def project_close_release(request, id, release_id):
    """Close a release by setting status to CLOSED and recording closed_at/closed_by."""
    project = get_object_or_404(Project, id=id)
    release = get_object_or_404(Release, id=release_id, project=project)
    
    try:
        # Check if already closed
        if release.status == ReleaseStatus.CLOSED:
            return JsonResponse({'success': False, 'error': 'Release is already closed'}, status=400)
        
        # Update release to closed status
        release.status = ReleaseStatus.CLOSED
        release.closed_at = timezone.now()
        release.closed_by = request.user
        release.save()
        
        return JsonResponse({'success': True, 'message': 'Release closed successfully'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required
def release_detail_modal(request, release_id):
    """Return the release detail modal content with items table."""
    from django_tables2 import RequestConfig
    from .tables import ReleaseItemsTable
    from .filters import ReleaseItemsFilter
    
    release = get_object_or_404(Release, id=release_id)
    
    # Get items for this release
    items_queryset = Item.objects.filter(solution_release=release).select_related(
        'type', 'organisation', 'assigned_to'
    )
    
    # Apply filters
    item_filter = ReleaseItemsFilter(request.GET, queryset=items_queryset)
    
    # Create table
    table = ReleaseItemsTable(item_filter.qs)
    RequestConfig(request, paginate={'per_page': 25}).configure(table)
    
    # Get the primary change for this release
    primary_change = release.get_primary_change()
    
    context = {
        'release': release,
        'table': table,
        'filter': item_filter,
        'items_count': items_queryset.count(),
        'primary_change': primary_change,
    }
    
    return render(request, 'partials/release_detail_modal_content.html', context)


@login_required
def release_create_change(request, release_id):
    """Create a new Change from a Release."""
    release = get_object_or_404(Release, id=release_id)
    
    # Check if a change already exists for this release
    existing_change = release.get_primary_change()
    if existing_change:
        return JsonResponse({
            'success': False, 
            'error': 'A change already exists for this release',
            'change_id': existing_change.id
        }, status=400)
    
    try:
        # Create the change
        change = Change.objects.create(
            project=release.project,
            title=f'Change für {release.name}',
            description=f'Automatically created change for release {release.version}',
            release=release,
            planned_date=release.planned_date,  # Date-only field
            status=ChangeStatus.DRAFT
        )
        
        return JsonResponse({
            'success': True, 
            'message': 'Change created successfully',
            'change_id': change.id
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required
def project_attachments_tab(request, id):
    """Return the project attachments tab content."""
    project = get_object_or_404(Project, id=id)
    
    # Get query parameters for filtering
    search_query = request.GET.get('search', '').strip()
    file_type_filter = request.GET.get('file_type', '').strip()
    
    # Get attachments linked to this project
    content_type = ContentType.objects.get_for_model(Project)
    attachment_links = AttachmentLink.objects.filter(
        target_content_type=content_type,
        target_object_id=project.id,
        role=AttachmentRole.PROJECT_FILE
    ).select_related('attachment', 'attachment__created_by').order_by('-created_at')
    
    attachments = [link.attachment for link in attachment_links if not link.attachment.is_deleted]
    
    # Apply filename search filter
    if search_query:
        attachments = [a for a in attachments if search_query.lower() in a.original_name.lower()]
    
    # Apply file type filter
    if file_type_filter:
        attachments = [a for a in attachments if a.file_type == file_type_filter]
    
    # Get distinct file types for filter dropdown
    all_attachments = [link.attachment for link in attachment_links if not link.attachment.is_deleted]
    file_types = sorted(set(a.file_type for a in all_attachments if a.file_type))
    
    # Pagination - 10 items per page
    paginator = Paginator(attachments, 10)
    page_number = request.GET.get('page', 1)
    page_obj = paginator.get_page(page_number)
    
    context = {
        'project': project,
        'page_obj': page_obj,
        'file_types': file_types,
        'search_query': search_query,
        'file_type_filter': file_type_filter,
    }
    return render(request, 'partials/project_attachments_tab.html', context)


@login_required

@require_POST
def project_upload_attachment(request, id):
    """Upload an attachment to a project."""
    project = get_object_or_404(Project, id=id)
    
    if 'file' not in request.FILES:
        return HttpResponse("No file provided", status=400)
    
    uploaded_file = request.FILES['file']
    
    try:
        # Store attachment
        storage_service = AttachmentStorageService()
        attachment = storage_service.store_attachment(
            file=uploaded_file,
            target=project,
            role=AttachmentRole.PROJECT_FILE,
            created_by=request.user if request.user.is_authenticated else None,
        )
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='attachment.uploaded',
            target=project,
            actor=request.user if request.user.is_authenticated else None,
            summary=f"Uploaded file: {attachment.original_name}",
        )
        
        # Return success response (for AJAX uploads)
        # If this is an AJAX request (multi-file upload), return simple success
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest' and not request.headers.get('HX-Request'):
            return HttpResponse("Upload successful", status=200)
        
        # Return updated attachments list for HTMX requests
        content_type = ContentType.objects.get_for_model(Project)
        attachment_links = AttachmentLink.objects.filter(
            target_content_type=content_type,
            target_object_id=project.id,
            role=AttachmentRole.PROJECT_FILE
        ).select_related('attachment', 'attachment__created_by').order_by('-created_at')
        
        attachments = [link.attachment for link in attachment_links if not link.attachment.is_deleted]
        
        context = {
            'project': project,
            'attachments': attachments,
        }
        response = render(request, 'partials/project_attachments_tab.html', context)
        response['HX-Trigger'] = 'attachmentUploaded'
        return response
        
    except ValidationError as e:
        return HttpResponse(f"Validation error: {str(e)}", status=400)
    except PermissionError as e:
        return HttpResponse("Permission denied", status=403)
    except Exception as e:
        # Log the error for debugging
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"Attachment upload failed for project {id}: {str(e)}")
        return HttpResponse("Upload failed. Please try again.", status=500)


@login_required

@require_POST
def project_delete_attachment(request, attachment_id):
    """Delete a project attachment."""
    try:
        attachment = get_object_or_404(Attachment, id=attachment_id)
        
        # Get the project for activity logging
        attachment_link = AttachmentLink.objects.filter(
            attachment=attachment,
            role=AttachmentRole.PROJECT_FILE
        ).first()
        
        if attachment_link and hasattr(attachment_link.target, 'id'):
            project = attachment_link.target
        else:
            project = None
        
        # Store filename for logging
        filename = attachment.original_name
        
        # Delete the attachment (hard delete: removes file from storage, DB record, and Weaviate via signal)
        storage_service = AttachmentStorageService()
        storage_service.delete_attachment(attachment, hard=True)
        
        # Log activity if we have a project
        if project:
            activity_service = ActivityService()
            activity_service.log(
                verb='attachment.deleted',
                target=project,
                actor=request.user if request.user.is_authenticated else None,
                summary=f"Deleted file: {filename}",
            )
        
        return JsonResponse({'success': True, 'message': 'Attachment deleted successfully'})
        
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"Attachment deletion failed: {str(e)}")
        return JsonResponse({'success': False, 'error': 'Failed to delete attachment'}, status=500)


@login_required
def project_view_attachment(request, attachment_id):
    """View a project attachment (for viewable file types)."""
    try:
        attachment = get_object_or_404(Attachment, id=attachment_id)
        
        if attachment.is_deleted:
            return JsonResponse({'success': False, 'error': 'Attachment not found'}, status=404)
        
        # Get file extension using os.path.splitext for reliability
        import os
        _, extension = os.path.splitext(attachment.original_name.lower())
        extension = extension.lstrip('.')  # Remove leading dot
        
        # Read file content
        storage_service = AttachmentStorageService()
        file_content = storage_service.read_attachment(attachment)
        
        # Process based on file type
        if extension == 'md':
            # Render markdown to HTML - create parser instance per request for thread safety
            md_parser = markdown.Markdown(extensions=['extra', 'fenced_code'])
            html_content = md_parser.convert(file_content.decode('utf-8', errors='replace'))
            # Sanitize HTML
            clean_html = bleach.clean(
                html_content,
                tags=ALLOWED_TAGS,
                attributes=ALLOWED_ATTRIBUTES,
                strip=True
            )
            return JsonResponse({'success': True, 'content_html': clean_html})
        elif extension == 'pdf':
            # Return base64 encoded PDF
            import base64
            pdf_base64 = base64.b64encode(file_content).decode('utf-8')
            return JsonResponse({'success': True, 'content_base64': pdf_base64})
        elif extension in ['html', 'htm']:
            # Return sanitized HTML content (will be displayed in iframe)
            html_content = file_content.decode('utf-8', errors='replace')
            # Sanitize HTML before returning
            clean_html = bleach.clean(
                html_content,
                tags=ALLOWED_TAGS + ['html', 'head', 'body', 'meta', 'title', 'style'],
                attributes=ALLOWED_ATTRIBUTES,
                strip=True
            )
            return JsonResponse({'success': True, 'content': clean_html})
        else:
            # Plain text
            return JsonResponse({'success': True, 'content': file_content.decode('utf-8', errors='replace')})
            
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
def project_download_attachment(request, attachment_id):
    """Download a project attachment."""
    try:
        attachment = get_object_or_404(Attachment, id=attachment_id)
        
        if attachment.is_deleted:
            return HttpResponse("Attachment not found", status=404)
        
        # Read file content
        storage_service = AttachmentStorageService()
        file_content = storage_service.read_attachment(attachment)
        
        # Create response with file
        response = HttpResponse(file_content, content_type=attachment.content_type or 'application/octet-stream')
        response['Content-Disposition'] = f'attachment; filename="{attachment.original_name}"'
        response['Content-Length'] = len(file_content)
        
        return response
    except Exception as e:
        return HttpResponse(f"Download failed: {str(e)}", status=500)


@login_required

@require_POST
def project_import_github_issues(request, id):
    """Import closed GitHub issues for a project."""
    from core.services.github.service import GitHubService
    from core.services.integrations.base import IntegrationDisabled, IntegrationNotConfigured
    
    project = get_object_or_404(Project, id=id)
    
    # Check if project has GitHub repo configured
    if not project.github_owner or not project.github_repo:
        return JsonResponse({
            'success': False,
            'error': 'Project does not have a GitHub repository configured'
        }, status=400)
    
    try:
        # Initialize GitHub service
        github_service = GitHubService()
        
        # Import closed issues
        stats = github_service.import_closed_issues_for_project(
            project=project,
            actor=request.user if request.user.is_authenticated else None,
        )
        
        # Prepare response message
        if stats['issues_imported'] > 0:
            message = (
                f"Successfully imported {stats['issues_imported']} closed issue(s). "
                f"Found {stats['issues_found']} total closed issues. "
                f"Linked {stats['prs_linked']} PR(s)."
            )
        elif stats['issues_found'] > 0:
            message = (
                f"Found {stats['issues_found']} closed issue(s), but all were already imported. "
                f"Linked {stats['prs_linked']} PR(s)."
            )
        else:
            message = "No closed issues found in the repository."
        
        # Add error info if there were errors
        if stats['errors']:
            message += f" Note: {len(stats['errors'])} error(s) occurred during import."
        
        return JsonResponse({
            'success': True,
            'message': message,
            'stats': stats,
        })
    
    except IntegrationDisabled:
        return JsonResponse({
            'success': False,
            'error': 'GitHub integration is not enabled'
        }, status=400)
    
    except IntegrationNotConfigured:
        return JsonResponse({
            'success': False,
            'error': 'GitHub integration is not configured. Please configure GitHub token in admin.'
        }, status=400)
    
    except ValueError as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=400)
    
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"Error importing GitHub issues: {e}", exc_info=True)
        
        return JsonResponse({
            'success': False,
            'error': 'An unexpected error occurred while importing issues'
        }, status=500)


@login_required
@require_POST
def project_sync_markdown(request, id):
    """Synchronize markdown files from GitHub repository for a project."""
    from core.services.github.service import GitHubService
    from core.services.github_sync.markdown_sync import MarkdownSyncService
    from core.services.integrations.base import IntegrationDisabled, IntegrationNotConfigured
    
    project = get_object_or_404(Project, id=id)
    
    # Check if project has GitHub repo configured
    if not project.github_owner or not project.github_repo:
        return JsonResponse({
            'success': False,
            'error': 'Project does not have a GitHub repository configured'
        }, status=400)
    
    try:
        # Initialize GitHub service
        github_service = GitHubService()
        
        # Get GitHub client
        client = github_service._get_client()
        
        # Initialize markdown sync service
        markdown_service = MarkdownSyncService(github_client=client)
        
        # Sync markdown files for this project
        stats = markdown_service.sync_project_markdown_files(project)
        
        # Prepare response message
        if stats['files_created'] > 0 or stats['files_updated'] > 0:
            message = (
                f"Successfully synced {stats['files_found']} markdown file(s). "
                f"Created: {stats['files_created']}, "
                f"Updated: {stats['files_updated']}, "
                f"Skipped: {stats['files_skipped']}"
            )
        elif stats['files_found'] > 0:
            message = (
                f"Found {stats['files_found']} markdown file(s), but all were up to date. "
                f"No changes needed."
            )
        else:
            message = "No markdown files found in the repository."
        
        # Add error info if there were errors
        if stats['errors']:
            message += f" Note: {len(stats['errors'])} error(s) occurred during sync."
        
        return JsonResponse({
            'success': True,
            'message': message,
            'stats': stats,
        })
    
    except IntegrationDisabled:
        return JsonResponse({
            'success': False,
            'error': 'GitHub integration is not enabled'
        }, status=400)
    
    except IntegrationNotConfigured:
        return JsonResponse({
            'success': False,
            'error': 'GitHub integration is not configured. Please configure GitHub token in admin.'
        }, status=400)
    
    except ValueError as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=400)
    
    except Exception as e:
        import logging
        logger = logging.getLogger(__name__)
        logger.error(f"Error syncing markdown files: {e}", exc_info=True)
        
        return JsonResponse({
            'success': False,
            'error': 'An unexpected error occurred while syncing markdown files'
        }, status=500)


# ============================================================================
# Organisation CRUD Views
# ============================================================================

@login_required
def organisations(request):
    """Organisations list view with filtering."""
    orgs = Organisation.objects.all()
    
    # Annotate with user and project counts
    orgs = orgs.annotate(
        user_count=Count('user_organisations', distinct=True),
        project_count=Count('projects', distinct=True)
    )
    
    # Server-side search filter
    q = request.GET.get('q', '')
    if q:
        orgs = orgs.filter(name__icontains=q)
    
    context = {
        'organisations': orgs,
        'search_query': q,
    }
    return render(request, 'organisations.html', context)


@login_required
def organisation_create(request):
    """Organisation create page view."""
    if request.method == 'GET':
        # Show the create form
        context = {
            'organisation': None,
        }
        return render(request, 'organisation_form.html', context)
    
    # Handle POST request (HTMX form submission)
    try:
        short = request.POST.get('short', '').strip()
        
        # Validate short field length
        if short and len(short) > 10:
            return JsonResponse({
                'success': False, 
                'error': 'Short code must be 10 characters or less'
            }, status=400)
        
        organisation = Organisation.objects.create(
            name=request.POST.get('name'),
            short=short
        )
        return JsonResponse({
            'success': True,
            'message': 'Organisation created successfully',
            'organisation_id': organisation.id,
            'redirect': f'/organisations/{organisation.id}/'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required
def organisation_edit(request, id):
    """Organisation edit page view."""
    organisation = get_object_or_404(Organisation, id=id)
    
    if request.method == 'GET':
        # Show the edit form
        context = {
            'organisation': organisation,
        }
        return render(request, 'organisation_form.html', context)


@login_required
def organisation_update(request, id):
    """Organisation update endpoint."""
    organisation = get_object_or_404(Organisation, id=id)
    
    # Handle POST request (HTMX form submission)
    try:
        short = request.POST.get('short', '').strip()
        
        # Validate short field length
        if short and len(short) > 10:
            return JsonResponse({
                'success': False, 
                'error': 'Short code must be 10 characters or less'
            }, status=400)
        
        organisation.name = request.POST.get('name', organisation.name)
        organisation.short = short
        organisation.save()
        
        return JsonResponse({
            'success': True,
            'message': 'Organisation updated successfully',
            'redirect': f'/organisations/{organisation.id}/'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required
def organisation_detail(request, id):
    """Organisation detail page view."""
    organisation = get_object_or_404(
        Organisation.objects.annotate(
            user_count=Count('user_organisations', distinct=True),
            project_count=Count('projects', distinct=True)
        ),
        id=id
    )
    
    # Get all users for the user management
    all_users = User.objects.filter(active=True).order_by('name')
    
    # Get all projects for the project linking
    all_projects = Project.objects.all().order_by('name')
    
    # Get users in this organisation
    user_organisations = organisation.user_organisations.select_related('user').order_by('user__name')
    
    # Get projects linked to this organisation
    projects = organisation.projects.all().order_by('name')
    
    context = {
        'organisation': organisation,
        'all_users': all_users,
        'all_projects': all_projects,
        'user_organisations': user_organisations,
        'projects': projects,
        'user_roles': UserRole.choices,
        'default_role': UserRole.USER,
    }
    return render(request, 'organisation_detail.html', context)


@login_required
def organisation_delete(request, id):
    """Delete an organisation."""
    organisation = get_object_or_404(Organisation, id=id)
    
    try:
        organisation.delete()
        return JsonResponse({'success': True, 'redirect': '/organisations/'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required

@require_http_methods(["POST"])
def organisation_add_user(request, id):
    """Add a user to an organisation."""
    organisation = get_object_or_404(Organisation, id=id)
    user_id = request.POST.get('user_id')
    is_primary = request.POST.get('is_primary', 'false') == 'true'
    role = request.POST.get('role', UserRole.USER)
    
    if not user_id:
        return JsonResponse({'success': False, 'error': 'User ID required'}, status=400)
    
    try:
        user = get_object_or_404(User, id=user_id)
        
        # Check if user is already in this organisation
        if organisation.user_organisations.filter(user=user).exists():
            return JsonResponse({'success': False, 'error': 'User already in this organisation'}, status=400)
        
        # Create the UserOrganisation relationship
        UserOrganisation.objects.create(
            organisation=organisation,
            user=user,
            role=role,
            is_primary=is_primary
        )
        
        return JsonResponse({'success': True, 'message': 'User added successfully'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required

@require_http_methods(["POST"])
def organisation_remove_user(request, id):
    """Remove a user from an organisation."""
    organisation = get_object_or_404(Organisation, id=id)
    user_id = request.POST.get('user_id')
    
    if not user_id:
        return JsonResponse({'success': False, 'error': 'User ID required'}, status=400)
    
    try:
        user = get_object_or_404(User, id=user_id)
        organisation.user_organisations.filter(user=user).delete()
        return JsonResponse({'success': True, 'message': 'User removed successfully'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required

@require_http_methods(["POST"])
def organisation_update_user(request, id):
    """Update a user's role and primary status in an organisation."""
    organisation = get_object_or_404(Organisation, id=id)
    user_id = request.POST.get('user_id')
    role = request.POST.get('role')
    is_primary = request.POST.get('is_primary', 'false') == 'true'
    
    if not user_id:
        return JsonResponse({'success': False, 'error': 'User ID required'}, status=400)
    
    if not role:
        return JsonResponse({'success': False, 'error': 'Role required'}, status=400)
    
    try:
        user = get_object_or_404(User, id=user_id)
        user_org = get_object_or_404(UserOrganisation, organisation=organisation, user=user)
        
        # Update the role and primary status
        user_org.role = role
        user_org.is_primary = is_primary
        user_org.save()
        
        return JsonResponse({'success': True, 'message': 'User updated successfully'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required
@require_http_methods(["POST"])
def organisation_create_user(request, id):
    """Create a new user and add them to the organisation as primary member."""
    organisation = get_object_or_404(Organisation, id=id)

    username = request.POST.get('username', '').strip()
    email = request.POST.get('email', '').strip().lower()
    name = request.POST.get('name', '').strip()
    role = request.POST.get('role', UserRole.USER)

    # Duplicate-email check (case-insensitive via lowercasing)
    if User.objects.filter(email__iexact=email).exists():
        return JsonResponse(
            {'success': False, 'global_error': 'A user with this email address already exists.'},
            status=400
        )

    try:
        with transaction.atomic():
            user = User(username=username, email=email, name=name, role=role, active=True)
            user.set_unusable_password()
            user.full_clean()
            user.save()

            UserOrganisation.objects.create(
                organisation=organisation,
                user=user,
                role=role,
                is_primary=True,
            )

        return JsonResponse({
            'success': True,
            'message': f'User {name} created successfully.',
            'user_id': user.id,
        })
    except ValidationError as e:
        errors = {}
        if hasattr(e, 'message_dict'):
            errors = {field: list(msgs) for field, msgs in e.message_dict.items()}
        else:
            errors['__all__'] = list(e.messages)
        return JsonResponse({'success': False, 'errors': errors}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required

@require_http_methods(["POST"])
def organisation_link_project(request, id):
    """Link a project to an organisation."""
    organisation = get_object_or_404(Organisation, id=id)
    project_id = request.POST.get('project_id')
    
    if not project_id:
        return JsonResponse({'success': False, 'error': 'Project ID required'}, status=400)
    
    try:
        project = get_object_or_404(Project, id=project_id)
        project.clients.add(organisation)
        return JsonResponse({'success': True, 'message': 'Project linked successfully'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required

@require_http_methods(["POST"])
def organisation_unlink_project(request, id):
    """Unlink a project from an organisation."""
    organisation = get_object_or_404(Organisation, id=id)
    project_id = request.POST.get('project_id')
    
    if not project_id:
        return JsonResponse({'success': False, 'error': 'Project ID required'}, status=400)
    
    try:
        project = get_object_or_404(Project, id=project_id)
        project.clients.remove(organisation)
        return JsonResponse({'success': True, 'message': 'Project unlinked successfully'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


# ============================================================================
# AI Provider CRUD Views
# ============================================================================

@login_required
def ai_providers(request):
    """AI Providers list view with filtering."""
    providers = AIProvider.objects.all()
    
    # Annotate with model count
    providers = providers.annotate(
        model_count=Count('models')
    )
    
    # Search filter
    q = request.GET.get('q', '')
    if q:
        providers = providers.filter(
            Q(name__icontains=q) | Q(provider_type__icontains=q)
        )
    
    # Provider type filter
    provider_type_filter = request.GET.get('provider_type', '')
    if provider_type_filter:
        providers = providers.filter(provider_type=provider_type_filter)
    
    # Active filter
    active_filter = request.GET.get('active', '')
    if active_filter:
        providers = providers.filter(active=(active_filter == 'true'))
    
    # Get provider types for filter dropdown
    provider_types = AIProviderType.choices
    
    context = {
        'providers': providers,
        'search_query': q,
        'provider_types': provider_types,
        'selected_provider_type': provider_type_filter,
        'selected_active': active_filter,
    }
    return render(request, 'ai_providers.html', context)


@login_required
def ai_provider_detail(request, id):
    """AI Provider detail view with models."""
    provider = get_object_or_404(AIProvider, id=id)
    
    # Get all models for this provider
    models = provider.models.all().order_by('-is_default', 'name')
    
    # Get provider types for dropdown
    provider_types = AIProviderType.choices
    
    context = {
        'provider': provider,
        'models': models,
        'provider_types': provider_types,
    }
    return render(request, 'ai_provider_detail.html', context)


@login_required
def ai_provider_create(request):
    """Create a new AI Provider."""
    if request.method == 'GET':
        provider_types = AIProviderType.choices
        context = {
            'provider': None,
            'provider_types': provider_types,
        }
        return render(request, 'ai_provider_form.html', context)
    
    # Handle POST request
    try:
        provider = AIProvider.objects.create(
            name=request.POST.get('name'),
            provider_type=request.POST.get('provider_type'),
            api_key=request.POST.get('api_key'),
            organization_id=request.POST.get('organization_id', ''),
            active=request.POST.get('active') == 'on'
        )
        
        # Return HTMX response with redirect
        response = HttpResponse()
        response['HX-Redirect'] = f'/ai-providers/{provider.id}/'
        return response
        
    except Exception as e:
        return HttpResponse(f"Error creating provider: {str(e)}", status=400)


@login_required

@require_http_methods(["POST"])
def ai_provider_update(request, id):
    """Update AI Provider."""
    provider = get_object_or_404(AIProvider, id=id)
    
    try:
        provider.name = request.POST.get('name', provider.name)
        provider.provider_type = request.POST.get('provider_type', provider.provider_type)
        provider.organization_id = request.POST.get('organization_id', provider.organization_id)
        provider.active = request.POST.get('active') == 'on'
        
        # Only update api_key if a new one is provided (not masked)
        api_key = request.POST.get('api_key', '').strip()
        # Skip if empty or if it's only asterisks (masked placeholder)
        # Define masked pattern check for clarity
        is_masked = api_key and set(api_key) == {'*'}
        if api_key and not is_masked:
            # API key contains non-asterisk characters, so it's a real key
            provider.api_key = api_key
        
        provider.save()
        
        # Return success toast trigger
        response = HttpResponse()
        response['HX-Trigger'] = 'showToast'
        return response
        
    except Exception as e:
        return HttpResponse(f"Error updating provider: {str(e)}", status=400)


@login_required

@require_http_methods(["POST"])
def ai_provider_delete(request, id):
    """Delete AI Provider."""
    provider = get_object_or_404(AIProvider, id=id)
    
    try:
        provider.delete()
        
        # Return redirect to list
        response = HttpResponse()
        response['HX-Redirect'] = '/ai-providers/'
        return response
        
    except Exception as e:
        return HttpResponse(f"Error deleting provider: {str(e)}", status=400)


@require_http_methods(["GET"])

@login_required
def ai_provider_get_api_key(request, id):
    """Get decrypted API key for copying. Requires authentication."""
    # Check if user is authenticated
    if not request.user.is_authenticated:
        return JsonResponse({'error': 'Authentication required'}, status=401)
    
    provider = get_object_or_404(AIProvider, id=id)
    
    # Return the actual API key for clipboard copy
    return JsonResponse({
        'api_key': provider.api_key
    })


@login_required

@require_http_methods(["POST"])
def ai_provider_fetch_models(request, id):
    """Fetch available models from the provider API and save them to the database."""
    provider = get_object_or_404(AIProvider, id=id)
    
    try:
        models_data = []
        
        if provider.provider_type == 'OpenAI':
            # Validate API key exists
            if not provider.api_key:
                raise ValueError("OpenAI API key is not configured. Please add your API key in the provider settings.")
            
            # Use OpenAI API to list models
            openai_client = openai.OpenAI(api_key=provider.api_key)
            models_list = openai_client.models.list()
            
            # Include all GPT models without restrictive filtering
            # This includes gpt-3.5, gpt-4, gpt-4o, gpt-5, o1, o3, and all variants
            for model in models_list.data:
                model_id_lower = model.id.lower()
                # Include all GPT models (gpt-*, o1-*, o3-*) but exclude embeddings and other non-chat models
                if (model_id_lower.startswith('gpt-') or 
                    model_id_lower.startswith('o1-') or 
                    model_id_lower.startswith('o3-')):
                    models_data.append({
                        'name': model.id,
                        'model_id': model.id
                    })
        
        elif provider.provider_type == 'Gemini':
            # Validate API key exists
            if not provider.api_key:
                raise ValueError("Gemini API key is not configured. Please add your API key in the provider settings.")
            
            # Use Gemini API to list all available models
            gemini_client = genai.Client(api_key=provider.api_key)
            models_list = gemini_client.models.list()
            
            # Filter to only generative models (exclude embedding models)
            for model in models_list:
                # Only include models that support generateContent
                if (hasattr(model, 'supported_generation_methods') and 
                    model.supported_generation_methods and 
                    'generateContent' in model.supported_generation_methods):
                    # Use the model name without 'models/' prefix if present
                    model_id = getattr(model, 'name', str(model)).replace('models/', '')
                    # Use display_name if available, otherwise use the model_id
                    model_name = getattr(model, 'display_name', None) or model_id
                    models_data.append({
                        'name': model_name,
                        'model_id': model_id
                    })
        
        elif provider.provider_type == 'Claude':
            # For Claude, use predefined list (updated as of Jan 2026)
            models_data = [
                {'name': 'Claude 3.5 Sonnet', 'model_id': 'claude-3-5-sonnet-20241022'},
                {'name': 'Claude 3 Opus', 'model_id': 'claude-3-opus-20240229'},
                {'name': 'Claude 3 Sonnet', 'model_id': 'claude-3-sonnet-20240229'},
                {'name': 'Claude 3 Haiku', 'model_id': 'claude-3-haiku-20240307'},
            ]
        
        # Save fetched models to the database
        created_count = 0
        existing_count = 0
        updated_count = 0
        
        # Track which model_ids were fetched from the API
        fetched_model_ids = set()
        
        for model_data in models_data:
            model_id = model_data['model_id']
            fetched_model_ids.add(model_id)
            
            model, created = AIModel.objects.get_or_create(
                provider=provider,
                model_id=model_id,
                defaults={
                    'name': model_data['name'],
                    'active': False,  # New models are inactive by default
                    'is_default': False,
                    'input_price_per_1m_tokens': None,
                    'output_price_per_1m_tokens': None,
                }
            )
            if created:
                created_count += 1
            else:
                # Update the name if it changed in the remote API
                if model.name != model_data['name']:
                    model.name = model_data['name']
                    model.save()
                    updated_count += 1
                existing_count += 1
        
        # Deactivate models that are no longer available in the remote API
        deactivated_models = []
        all_local_models = provider.models.all()
        for local_model in all_local_models:
            if local_model.model_id not in fetched_model_ids and local_model.active:
                local_model.active = False
                local_model.save()
                deactivated_models.append(local_model.name)
        
        # Return updated models list using the partial template
        models = provider.models.all().order_by('-is_default', 'name')
        context = {
            'provider': provider,
            'models': models,
            'sync_info': {
                'created': created_count,
                'existing': existing_count,
                'updated': updated_count,
                'deactivated': deactivated_models,
            }
        }
        return render(request, 'partials/ai_models_list.html', context)
        
    except Exception as e:
        # Return HTML error message for HTMX to display properly
        logger.error(f"Failed to fetch models for provider {provider.id} ({provider.provider_type}): {str(e)}")
        models = provider.models.all().order_by('-is_default', 'name')
        context = {
            'provider': provider,
            'models': models,
            'error_message': f'Failed to fetch models from {provider.provider_type} API: {str(e)}',
        }
        return render(request, 'partials/ai_models_list.html', context)


@login_required

@require_http_methods(["POST"])
def ai_model_create(request, provider_id):
    """Create a new AI Model for a provider."""
    provider = get_object_or_404(AIProvider, id=provider_id)
    
    try:
        # Get input values and convert empty strings to None for decimal fields
        input_price = request.POST.get('input_price_per_1m_tokens', '').strip()
        output_price = request.POST.get('output_price_per_1m_tokens', '').strip()
        
        model = AIModel.objects.create(
            provider=provider,
            name=request.POST.get('name'),
            model_id=request.POST.get('model_id'),
            input_price_per_1m_tokens=input_price if input_price else None,
            output_price_per_1m_tokens=output_price if output_price else None,
            active=request.POST.get('active') == 'on',
            is_default=request.POST.get('is_default') == 'on'
        )
        
        # If this is set as default, unset other defaults
        if model.is_default:
            AIModel.objects.filter(provider=provider).exclude(id=model.id).update(is_default=False)
        
        # Return updated models list
        models = provider.models.all().order_by('-is_default', 'name')
        context = {
            'provider': provider,
            'models': models,
        }
        return render(request, 'partials/ai_models_list.html', context)
        
    except Exception as e:
        return HttpResponse(f"Error creating model: {str(e)}", status=400)


@login_required

@require_http_methods(["POST"])
def ai_model_update(request, provider_id, model_id):
    """Update an AI Model."""
    provider = get_object_or_404(AIProvider, id=provider_id)
    model = get_object_or_404(AIModel, id=model_id, provider=provider)
    
    try:
        model.name = request.POST.get('name', model.name)
        model.model_id = request.POST.get('model_id', model.model_id)
        
        # Handle decimal fields - convert empty strings to None
        input_price = request.POST.get('input_price_per_1m_tokens', '').strip()
        output_price = request.POST.get('output_price_per_1m_tokens', '').strip()
        model.input_price_per_1m_tokens = input_price if input_price else None
        model.output_price_per_1m_tokens = output_price if output_price else None
        
        model.active = request.POST.get('active') == 'on'
        model.is_default = request.POST.get('is_default') == 'on'
        model.save()
        
        # If this is set as default, unset other defaults
        if model.is_default:
            AIModel.objects.filter(provider=provider).exclude(id=model.id).update(is_default=False)
        
        # Return updated models list
        models = provider.models.all().order_by('-is_default', 'name')
        context = {
            'provider': provider,
            'models': models,
        }
        return render(request, 'partials/ai_models_list.html', context)
        
    except Exception as e:
        return HttpResponse(f"Error updating model: {str(e)}", status=400)


@login_required

@require_http_methods(["POST"])
def ai_model_delete(request, provider_id, model_id):
    """Delete an AI Model."""
    provider = get_object_or_404(AIProvider, id=provider_id)
    model = get_object_or_404(AIModel, id=model_id, provider=provider)
    
    try:
        model.delete()
        
        # Return updated models list
        models = provider.models.all().order_by('-is_default', 'name')
        context = {
            'provider': provider,
            'models': models,
        }
        return render(request, 'partials/ai_models_list.html', context)
        
    except Exception as e:
        return HttpResponse(f"Error deleting model: {str(e)}", status=400)


@login_required

@require_http_methods(["POST"])
def ai_model_update_field(request, provider_id, model_id):
    """Update a single field of an AI Model via HTMX."""
    provider = get_object_or_404(AIProvider, id=provider_id)
    model = get_object_or_404(AIModel, id=model_id, provider=provider)
    
    try:
        field = request.POST.get('field')
        value = request.POST.get('value', '').strip()
        
        if field in ['input_price_per_1m_tokens', 'output_price_per_1m_tokens']:
            # Validate and convert to Decimal if value is provided
            if value:
                try:
                    decimal_value = Decimal(value)
                    if decimal_value < 0:
                        return HttpResponse("Price cannot be negative", status=400)
                except (InvalidOperation, ValueError):
                    return HttpResponse("Invalid price value", status=400)
                
                if field == 'input_price_per_1m_tokens':
                    model.input_price_per_1m_tokens = decimal_value
                else:
                    model.output_price_per_1m_tokens = decimal_value
            else:
                # Empty value sets to None
                if field == 'input_price_per_1m_tokens':
                    model.input_price_per_1m_tokens = None
                else:
                    model.output_price_per_1m_tokens = None
        else:
            return HttpResponse("Invalid field", status=400)
        
        model.save()
        return HttpResponse(status=204)  # No content, but success
        
    except Exception as e:
        return HttpResponse(f"Error updating field: {str(e)}", status=400)


@login_required

@require_http_methods(["POST"])
def ai_model_toggle_active(request, provider_id, model_id):
    """Toggle the active status of an AI Model via HTMX."""
    provider = get_object_or_404(AIProvider, id=provider_id)
    model = get_object_or_404(AIModel, id=model_id, provider=provider)
    
    try:
        model.active = not model.active
        model.save()
        
        # Return just the updated table row
        context = {
            'provider': provider,
            'model': model,
        }
        return render(request, 'partials/ai_model_row.html', context)
        
    except Exception as e:
        return HttpResponse(f"Error toggling active status: {str(e)}", status=400)


# ==================== Agent Views ====================

@login_required
def agents(request):
    """Agent list page view."""
    agent_service = AgentService()
    agents_list = agent_service.list_agents()
    
    # Server-side search filter
    q = request.GET.get('q', '')
    if q:
        agents_list = [a for a in agents_list if q.lower() in a.get('name', '').lower() or q.lower() in a.get('description', '').lower()]
    
    # Filter by provider
    provider_filter = request.GET.get('provider', '')
    if provider_filter:
        agents_list = [a for a in agents_list if a.get('provider', '').lower() == provider_filter.lower()]
    
    # Get unique providers for filter dropdown
    providers = sorted(list(set(a.get('provider', 'openai') for a in agent_service.list_agents())))
    
    context = {
        'agents': agents_list,
        'search_query': q,
        'providers': providers,
        'selected_provider': provider_filter,
    }
    return render(request, 'agents.html', context)


@login_required
def agent_detail(request, filename):
    """Agent detail/edit page view."""
    agent_service = AgentService()
    agent = agent_service.get_agent(filename)
    
    if not agent:
        return HttpResponse("Agent not found", status=404)
    
    # Ensure cache config exists with defaults
    if 'cache' not in agent:
        agent['cache'] = {
            'enabled': False,
            'ttl_seconds': 7776000,
            'key_strategy': 'content_hash',
            'agent_version': 1,
        }
    
    # Get all active providers and their models
    providers = AIProvider.objects.filter(active=True).prefetch_related('models')
    provider_types = AIProviderType.choices
    
    context = {
        'agent': agent,
        'is_new': False,
        'providers': providers,
        'provider_types': provider_types,
    }
    return render(request, 'agent_detail.html', context)


@login_required
def agent_create(request):
    """Agent create page view."""
    # Get all active providers and their models
    providers = AIProvider.objects.filter(active=True).prefetch_related('models')
    provider_types = AIProviderType.choices
    
    context = {
        'agent': {
            'name': '',
            'description': '',
            'provider': 'openai',
            'model': 'gpt-3.5-turbo',
            'role': '',
            'task': '',
            'parameters': {},
            'cache': {
                'enabled': False,
                'ttl_seconds': 7776000,
                'key_strategy': 'content_hash',
                'agent_version': 1,
            },
        },
        'is_new': True,
        'providers': providers,
        'provider_types': provider_types,
    }
    return render(request, 'agent_detail.html', context)


@login_required

@require_http_methods(["POST"])
def agent_save(request, filename):
    """Save agent (update existing)."""
    agent_service = AgentService()
    
    try:
        # Load existing agent to check for task changes
        existing_agent = agent_service.get_agent(filename)
        
        # Get form data
        name = request.POST.get('name', '').strip()
        description = request.POST.get('description', '').strip()
        provider = request.POST.get('provider', 'openai').strip()
        model = request.POST.get('model', 'gpt-3.5-turbo').strip()
        role = request.POST.get('role', '').strip()
        task = request.POST.get('task', '').strip()
        
        # Validate required fields
        if not name:
            return HttpResponse("Agent name is required", status=400)
        
        # Build agent data
        agent_data = {
            'name': name,
            'description': description,
            'provider': provider,
            'model': model,
            'role': role,
            'task': task,
        }
        
        # Parse parameters from form
        parameters = {}
        param_keys = request.POST.getlist('param_key[]')
        param_types = request.POST.getlist('param_type[]')
        param_descriptions = request.POST.getlist('param_description[]')
        param_required = request.POST.getlist('param_required[]')
        
        for i, key in enumerate(param_keys):
            if key.strip():
                param_def = {}
                if i < len(param_types) and param_types[i]:
                    param_def['type'] = param_types[i]
                if i < len(param_descriptions) and param_descriptions[i]:
                    param_def['description'] = param_descriptions[i]
                if i < len(param_required) and param_required[i] == 'true':
                    param_def['required'] = True
                else:
                    param_def['required'] = False
                    
                parameters[key] = param_def
        
        if parameters:
            agent_data['parameters'] = parameters
        
        # Parse cache configuration
        cache_enabled = request.POST.get('cache_enabled') == 'true'
        
        if cache_enabled:
            # Validate ttl_seconds
            ttl_seconds_str = request.POST.get('cache_ttl_seconds', '').strip()
            if not ttl_seconds_str:
                return HttpResponse("TTL seconds is required when cache is enabled", status=400)
            
            try:
                ttl_seconds = int(ttl_seconds_str)
                if ttl_seconds <= 0:
                    return HttpResponse("TTL seconds must be a positive integer", status=400)
            except ValueError:
                return HttpResponse("TTL seconds must be a valid integer", status=400)
            
            key_strategy = request.POST.get('cache_key_strategy', 'content_hash').strip()
            
            # Validate key_strategy
            if key_strategy not in ['content_hash']:
                return HttpResponse("Invalid key strategy. Only 'content_hash' is supported", status=400)
            
            # Get agent_version from form (user may have edited it, but we'll override if task changed)
            agent_version_str = request.POST.get('cache_agent_version', '1').strip()
            try:
                agent_version = int(agent_version_str)
            except ValueError:
                agent_version = 1
            
            # Check if task has changed - if so, increment agent_version
            if existing_agent:
                existing_task = existing_agent.get('task', '')
                if existing_task != task:
                    # Task changed - increment version
                    existing_cache = existing_agent.get('cache', {})
                    existing_version = existing_cache.get('agent_version', 0)
                    agent_version = existing_version + 1
            
            agent_data['cache'] = {
                'enabled': True,
                'ttl_seconds': ttl_seconds,
                'key_strategy': key_strategy,
                'agent_version': agent_version,
            }
        else:
            # Cache disabled - check if task changed to still increment version if cache exists
            if existing_agent:
                existing_task = existing_agent.get('task', '')
                existing_cache = existing_agent.get('cache', {})
                
                if existing_cache:
                    # Preserve cache block even if disabled, but update agent_version if task changed
                    agent_version = existing_cache.get('agent_version', 1)
                    if existing_task != task:
                        agent_version = agent_version + 1
                    
                    agent_data['cache'] = {
                        'enabled': False,
                        'ttl_seconds': existing_cache.get('ttl_seconds', 7776000),
                        'key_strategy': existing_cache.get('key_strategy', 'content_hash'),
                        'agent_version': agent_version,
                    }
        
        # Save agent
        agent_service.save_agent(filename, agent_data)
        
        # Redirect to detail view
        return redirect('agent-detail', filename=filename)
        
    except Exception as e:
        return HttpResponse(f"Error saving agent: {str(e)}", status=400)


@login_required

@require_http_methods(["POST"])
def agent_create_save(request):
    """Save agent (create new)."""
    agent_service = AgentService()
    
    try:
        # Get form data
        name = request.POST.get('name', '').strip()
        description = request.POST.get('description', '').strip()
        provider = request.POST.get('provider', 'openai').strip()
        model = request.POST.get('model', 'gpt-3.5-turbo').strip()
        role = request.POST.get('role', '').strip()
        task = request.POST.get('task', '').strip()
        
        # Validate required fields
        if not name:
            return HttpResponse("Agent name is required", status=400)
        
        # Build agent data
        agent_data = {
            'name': name,
            'description': description,
            'provider': provider,
            'model': model,
            'role': role,
            'task': task,
        }
        
        # Parse parameters from form
        parameters = {}
        param_keys = request.POST.getlist('param_key[]')
        param_types = request.POST.getlist('param_type[]')
        param_descriptions = request.POST.getlist('param_description[]')
        param_required = request.POST.getlist('param_required[]')
        
        for i, key in enumerate(param_keys):
            if key.strip():
                param_def = {}
                if i < len(param_types) and param_types[i]:
                    param_def['type'] = param_types[i]
                if i < len(param_descriptions) and param_descriptions[i]:
                    param_def['description'] = param_descriptions[i]
                if i < len(param_required) and param_required[i] == 'true':
                    param_def['required'] = True
                else:
                    param_def['required'] = False
                    
                parameters[key] = param_def
        
        if parameters:
            agent_data['parameters'] = parameters
        
        # Parse cache configuration
        cache_enabled = request.POST.get('cache_enabled') == 'true'
        
        if cache_enabled:
            # Validate ttl_seconds
            ttl_seconds_str = request.POST.get('cache_ttl_seconds', '').strip()
            if not ttl_seconds_str:
                return HttpResponse("TTL seconds is required when cache is enabled", status=400)
            
            try:
                ttl_seconds = int(ttl_seconds_str)
                if ttl_seconds <= 0:
                    return HttpResponse("TTL seconds must be a positive integer", status=400)
            except ValueError:
                return HttpResponse("TTL seconds must be a valid integer", status=400)
            
            key_strategy = request.POST.get('cache_key_strategy', 'content_hash').strip()
            
            # Validate key_strategy
            if key_strategy not in ['content_hash']:
                return HttpResponse("Invalid key strategy. Only 'content_hash' is supported", status=400)
            
            # For new agents, start with agent_version = 1
            agent_data['cache'] = {
                'enabled': True,
                'ttl_seconds': ttl_seconds,
                'key_strategy': key_strategy,
                'agent_version': 1,
            }
        
        # Creating new agent - generate filename from name
        safe_name = name.lower().replace(' ', '-').replace('_', '-')
        # Remove any non-alphanumeric characters except hyphens
        safe_name = ''.join(c for c in safe_name if c.isalnum() or c == '-')
        filename = f"{safe_name}.yml"
        
        # Check if file already exists and add suffix if needed
        counter = 1
        original_filename = filename
        while agent_service.get_agent(filename) is not None:
            name_part = original_filename.replace('.yml', '')
            filename = f"{name_part}-{counter}.yml"
            counter += 1
        
        # Save agent
        agent_service.save_agent(filename, agent_data)
        
        # Redirect to detail view
        return redirect('agent-detail', filename=filename)
        
    except Exception as e:
        return HttpResponse(f"Error saving agent: {str(e)}", status=400)


@login_required

@require_http_methods(["POST"])
def agent_delete(request, filename):
    """Delete an agent."""
    agent_service = AgentService()
    
    try:
        success = agent_service.delete_agent(filename)
        if not success:
            return HttpResponse("Agent not found", status=404)
        
        # Redirect to agent list
        return redirect('agents')
        
    except Exception as e:
        return HttpResponse(f"Error deleting agent: {str(e)}", status=400)


@login_required

@require_http_methods(["POST"])
def agent_test(request, filename):
    """Test an agent with input text."""
    agent_service = AgentService()
    
    try:
        # Get input text and parameters from request
        input_text = request.POST.get('input_text', '').strip()
        
        if not input_text:
            return JsonResponse({'error': 'Input text is required'}, status=400)
        
        # Parse parameters if provided
        parameters = {}
        param_keys = request.POST.getlist('test_param_key[]')
        param_values = request.POST.getlist('test_param_value[]')
        
        for i, key in enumerate(param_keys):
            if key.strip() and i < len(param_values):
                parameters[key] = param_values[i]
        
        # Get user and client IP
        user = request.user if request.user.is_authenticated else None
        client_ip = request.META.get('REMOTE_ADDR')
        
        # Execute agent
        result = agent_service.execute_agent(
            filename=filename,
            input_text=input_text,
            user=user,
            client_ip=client_ip,
            parameters=parameters if parameters else None
        )
        
        return JsonResponse({'result': result})
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
def ai_jobs_history(request):
    """AI Jobs History list view with filtering and pagination."""
    jobs = AIJobsHistory.objects.select_related('user', 'provider', 'model').all()
    
    # Order by timestamp descending (newest first)
    jobs = jobs.order_by('-timestamp')
    
    # Search filter (searches in agent name and error message)
    q = request.GET.get('q', '')
    if q:
        jobs = jobs.filter(
            Q(agent__icontains=q) | Q(error_message__icontains=q)
        )
    
    # Status filter
    status_filter = request.GET.get('status', '')
    if status_filter:
        jobs = jobs.filter(status=status_filter)
    
    # Provider filter
    provider_filter = request.GET.get('provider', '')
    if provider_filter:
        try:
            jobs = jobs.filter(provider_id=int(provider_filter))
        except (ValueError, TypeError):
            provider_filter = ''
    
    # Model filter
    model_filter = request.GET.get('model', '')
    if model_filter:
        try:
            jobs = jobs.filter(model_id=int(model_filter))
        except (ValueError, TypeError):
            model_filter = ''
    
    # Pagination - 25 items per page
    paginator = Paginator(jobs, 25)
    page_number = request.GET.get('page', 1)
    page_obj = paginator.get_page(page_number)
    
    # Get all providers and models for filter dropdowns
    providers = AIProvider.objects.all().order_by('name')
    models = AIModel.objects.all().order_by('name')
    
    # Get status choices from model
    from .models import AIJobStatus
    status_choices = AIJobStatus.choices
    
    context = {
        'page_obj': page_obj,
        'search_query': q,
        'providers': providers,
        'models': models,
        'status_choices': status_choices,
        'selected_status': status_filter,
        'selected_provider': provider_filter,
        'selected_model': model_filter,
    }
    return render(request, 'ai_jobs_history.html', context)


@login_required
def ai_job_statistics(request):
    """AI Job Statistics dashboard with aggregated KPIs, tables, and charts."""
    from datetime import timedelta
    from django.db.models.functions import TruncDate
    from django.db.models import Sum, Avg, Count
    from .models import AIJobStatus

    def _start_of_day(d):
        """Return timezone-aware start-of-day datetime for a date object."""
        return timezone.make_aware(datetime.combine(d, time.min))

    now = timezone.now()
    today = timezone.localdate()

    # KPI: Costs today
    costs_today = AIJobsHistory.objects.filter(
        timestamp__gte=_start_of_day(today)
    ).aggregate(total=Sum('costs'))['total'] or Decimal('0')

    # KPI: Costs current calendar week (Mon–Sun)
    start_of_week = today - timedelta(days=today.weekday())  # Monday
    costs_week = AIJobsHistory.objects.filter(
        timestamp__gte=_start_of_day(start_of_week)
    ).aggregate(total=Sum('costs'))['total'] or Decimal('0')

    # KPI: Costs current calendar month
    start_of_month = today.replace(day=1)
    costs_month = AIJobsHistory.objects.filter(
        timestamp__gte=_start_of_day(start_of_month)
    ).aggregate(total=Sum('costs'))['total'] or Decimal('0')

    # KPI: Errors last 7 days (today + 6 preceding days = 7-day window)
    start_of_7d_window = today - timedelta(days=6)
    errors_7d = AIJobsHistory.objects.filter(
        timestamp__gte=_start_of_day(start_of_7d_window),
        status=AIJobStatus.ERROR
    ).count()

    # Table: per Agent
    by_agent = (
        AIJobsHistory.objects.values('agent')
        .annotate(requests=Count('id'), total_costs=Sum('costs'))
        .order_by('-requests')
    )

    # Table: per Model
    by_model = (
        AIJobsHistory.objects.filter(model__isnull=False)
        .values('model__name')
        .annotate(requests=Count('id'), total_costs=Sum('costs'))
        .order_by('-requests')
    )

    # Table: per User
    by_user = (
        AIJobsHistory.objects.filter(user__isnull=False)
        .values('user__username')
        .annotate(requests=Count('id'), total_costs=Sum('costs'))
        .order_by('-requests')
    )

    # Timeseries last 7 days: requests per day
    requests_by_day_qs = (
        AIJobsHistory.objects.filter(timestamp__gte=_start_of_day(start_of_7d_window))
        .annotate(day=TruncDate('timestamp'))
        .values('day')
        .annotate(count=Count('id'))
        .order_by('day')
    )
    requests_by_day_dict = {item['day']: item['count'] for item in requests_by_day_qs}

    # Timeseries last 7 days: avg duration per day per agent
    duration_by_day_agent_qs = (
        AIJobsHistory.objects.filter(
            timestamp__gte=_start_of_day(start_of_7d_window),
            duration_ms__isnull=False
        )
        .annotate(day=TruncDate('timestamp'))
        .values('day', 'agent')
        .annotate(avg_duration=Avg('duration_ms'))
        .order_by('day', 'agent')
    )

    # Build chart data structures
    date_labels = []
    requests_series = []
    for i in range(6, -1, -1):
        day = today - timedelta(days=i)
        date_labels.append(day.strftime('%d.%m.'))
        requests_series.append(requests_by_day_dict.get(day, 0))

    # Build duration chart: one series per agent
    agent_names = sorted({item['agent'] for item in duration_by_day_agent_qs})
    duration_data = {}  # agent -> {day: avg_duration}
    for item in duration_by_day_agent_qs:
        agent = item['agent']
        if agent not in duration_data:
            duration_data[agent] = {}
        duration_data[agent][item['day']] = float(item['avg_duration'] or 0)

    duration_datasets = []
    colors = [
        'rgba(13, 110, 253, 0.7)',
        'rgba(25, 135, 84, 0.7)',
        'rgba(220, 53, 69, 0.7)',
        'rgba(255, 193, 7, 0.7)',
        'rgba(13, 202, 240, 0.7)',
        'rgba(111, 66, 193, 0.7)',
    ]
    for idx, agent in enumerate(agent_names):
        series = []
        for i in range(6, -1, -1):
            day = today - timedelta(days=i)
            series.append(duration_data.get(agent, {}).get(day, None))
        duration_datasets.append({
            'label': agent,
            'data': series,
            'borderColor': colors[idx % len(colors)],
            'backgroundColor': colors[idx % len(colors)].replace('0.7', '0.1'),
            'tension': 0.3,
            'spanGaps': True,
        })

    context = {
        'costs_today': costs_today,
        'costs_week': costs_week,
        'costs_month': costs_month,
        'errors_7d': errors_7d,
        'by_agent': list(by_agent),
        'by_model': list(by_model),
        'by_user': list(by_user),
        'requests_chart_json': json.dumps({
            'labels': date_labels,
            'data': requests_series,
        }),
        'duration_chart_json': json.dumps({
            'labels': date_labels,
            'datasets': duration_datasets,
        }),
    }
    return render(request, 'ai_job_statistics.html', context)


# ============================================================================
# Weaviate Sync Views
# ============================================================================

@login_required
def weaviate_status(request, object_type, object_id):
    """
    Check Weaviate sync status for an object.
    
    Returns HTML for the Weaviate status button (green if exists, red if not).
    This is an HTMX endpoint that can be called to refresh the button state.
    """
    from core.services.weaviate.client import is_available
    from core.services.weaviate.service import exists_object, is_excluded_from_sync
    
    # Check if Weaviate is available
    if not is_available():
        return render(request, 'partials/weaviate_button.html', {
            'object_type': object_type,
            'object_id': object_id,
            'exists': False,
            'available': False,
        })
    
    # Check if object is excluded from sync
    is_excluded, exclusion_reason = is_excluded_from_sync(object_type, object_id)
    if is_excluded:
        return render(request, 'partials/weaviate_button.html', {
            'object_type': object_type,
            'object_id': object_id,
            'exists': False,
            'available': True,
            'excluded': True,
        })
    
    # Check if object exists in Weaviate
    try:
        exists = exists_object(object_type, object_id)
    except Exception as e:
        exists = False
    
    return render(request, 'partials/weaviate_button.html', {
        'object_type': object_type,
        'object_id': object_id,
        'exists': exists,
        'available': True,
    })


@login_required
def weaviate_object(request, object_type, object_id):
    """
    Fetch Weaviate object data and display in modal content.
    
    Returns HTML for the modal body showing either:
    - The Weaviate object as formatted JSON (if exists)
    - A message with a "Push to Weaviate" button (if not exists)
    """
    from core.services.weaviate.client import is_available
    from core.services.weaviate.service import fetch_object_by_type
    import json
    
    # Check if Weaviate is available
    if not is_available():
        return render(request, 'partials/weaviate_modal_content.html', {
            'object_type': object_type,
            'object_id': object_id,
            'available': False,
            'error': 'Weaviate service is not configured or disabled.',
        })
    
    # Fetch object from Weaviate
    try:
        obj_data = fetch_object_by_type(object_type, object_id)
        
        if obj_data:
            # Format as pretty JSON
            json_str = json.dumps(obj_data, indent=2, default=str)
            
            return render(request, 'partials/weaviate_modal_content.html', {
                'object_type': object_type,
                'object_id': object_id,
                'available': True,
                'exists': True,
                'json_data': json_str,
            })
        else:
            return render(request, 'partials/weaviate_modal_content.html', {
                'object_type': object_type,
                'object_id': object_id,
                'available': True,
                'exists': False,
            })
            
    except Exception as e:
        return render(request, 'partials/weaviate_modal_content.html', {
            'object_type': object_type,
            'object_id': object_id,
            'available': True,
            'error': str(e),
        })


@login_required

@require_POST
def weaviate_push(request, object_type, object_id):
    """
    Manually push an object to Weaviate.
    
    This endpoint performs a manual sync of the object to Weaviate.
    Returns updated modal content showing the synced object.
    """
    from core.services.weaviate.client import is_available
    from core.services.weaviate.service import upsert_object, is_excluded_from_sync
    import json
    
    # Check if Weaviate is available
    if not is_available():
        return render(request, 'partials/weaviate_modal_content.html', {
            'object_type': object_type,
            'object_id': object_id,
            'available': False,
            'error': 'Weaviate service is not configured or disabled.',
        })
    
    # Check if object is excluded from sync
    is_excluded, exclusion_reason = is_excluded_from_sync(object_type, object_id)
    if is_excluded:
        return render(request, 'partials/weaviate_modal_content.html', {
            'object_type': object_type,
            'object_id': object_id,
            'available': True,
            'excluded': True,
            'info_message': exclusion_reason or 'This object is excluded from Weaviate indexing.',
        })
    
    # Push object to Weaviate
    try:
        uuid_str = upsert_object(object_type, object_id)
        
        if uuid_str:
            # Fetch the newly created object to show it
            from core.services.weaviate.service import fetch_object_by_type
            obj_data = fetch_object_by_type(object_type, object_id)
            
            if obj_data:
                json_str = json.dumps(obj_data, indent=2, default=str)
                
                return render(request, 'partials/weaviate_modal_content.html', {
                    'object_type': object_type,
                    'object_id': object_id,
                    'available': True,
                    'exists': True,
                    'json_data': json_str,
                    'success_message': 'Successfully pushed to Weaviate!',
                })
            else:
                return render(request, 'partials/weaviate_modal_content.html', {
                    'object_type': object_type,
                    'object_id': object_id,
                    'available': True,
                    'exists': True,
                    'success_message': 'Successfully pushed to Weaviate!',
                })
        else:
            return render(request, 'partials/weaviate_modal_content.html', {
                'object_type': object_type,
                'object_id': object_id,
                'available': True,
                'error': 'Could not push object to Weaviate. Object type may not be supported.',
            })
            
    except Exception as e:
        return render(request, 'partials/weaviate_modal_content.html', {
            'object_type': object_type,
            'object_id': object_id,
            'available': True,
            'error': f'Error pushing to Weaviate: {str(e)}',
        })


# Change Management Views

@login_required
def change_detail(request, id):
    """Change detail page view."""
    change = get_object_or_404(
        Change.objects.select_related(
            'project', 'created_by', 'release'
        ).prefetch_related('approvals__approver', 'organisations'),
        id=id
    )
    
    # Render markdown fields to HTML with sanitization
    description_html = None
    if change.description:
        MARKDOWN_PARSER.reset()
        html = MARKDOWN_PARSER.convert(change.description)
        description_html = mark_safe(bleach.clean(html, tags=ALLOWED_TAGS, attributes=ALLOWED_ATTRIBUTES, strip=True))
    
    risk_description_html = None
    if change.risk_description:
        MARKDOWN_PARSER.reset()
        html = MARKDOWN_PARSER.convert(change.risk_description)
        risk_description_html = mark_safe(bleach.clean(html, tags=ALLOWED_TAGS, attributes=ALLOWED_ATTRIBUTES, strip=True))
    
    mitigation_html = None
    if change.mitigation:
        MARKDOWN_PARSER.reset()
        html = MARKDOWN_PARSER.convert(change.mitigation)
        mitigation_html = mark_safe(bleach.clean(html, tags=ALLOWED_TAGS, attributes=ALLOWED_ATTRIBUTES, strip=True))
    
    rollback_plan_html = None
    if change.rollback_plan:
        MARKDOWN_PARSER.reset()
        html = MARKDOWN_PARSER.convert(change.rollback_plan)
        rollback_plan_html = mark_safe(bleach.clean(html, tags=ALLOWED_TAGS, attributes=ALLOWED_ATTRIBUTES, strip=True))
    
    communication_plan_html = None
    if change.communication_plan:
        MARKDOWN_PARSER.reset()
        html = MARKDOWN_PARSER.convert(change.communication_plan)
        communication_plan_html = mark_safe(bleach.clean(html, tags=ALLOWED_TAGS, attributes=ALLOWED_ATTRIBUTES, strip=True))
    
    # Get users for approver selection - filter by organisations assigned to change
    # Only show users with role != USER who are members of the change's organisations
    change_org_ids = list(change.organisations.values_list('id', flat=True))
    if change_org_ids:
        # Get users who belong to any of the change's organisations and have role != USER
        all_users = User.objects.filter(
            active=True,
            user_organisations__organisation_id__in=change_org_ids
        ).exclude(
            user_organisations__role=UserRole.USER
        ).distinct().order_by('name')
    else:
        # If no organisations assigned, show all active users with role != USER
        all_users = User.objects.filter(active=True).exclude(
            user_organisations__role=UserRole.USER
        ).distinct().order_by('name')
    
    # Load attachments for each approval
    approval_ct = ContentType.objects.get_for_model(ChangeApproval)
    for approval in change.approvals.all():
        approval.attachments = AttachmentLink.objects.filter(
            target_content_type=approval_ct,
            target_object_id=approval.id,
            role=AttachmentRole.APPROVER_ATTACHMENT
        ).select_related('attachment')
    
    # Get items associated with this change (direct + release items, deduplicated)
    items = change.get_associated_items()
    
    # Get today's date for the decision date default
    from datetime import date
    today = date.today()
    
    context = {
        'change': change,
        'description_html': description_html,
        'risk_description_html': risk_description_html,
        'mitigation_html': mitigation_html,
        'rollback_plan_html': rollback_plan_html,
        'communication_plan_html': communication_plan_html,
        'all_users': all_users,
        'items': items,
        'today': today,
    }
    return render(request, 'change_detail.html', context)


@login_required
def change_create(request):
    """Change create page view."""
    if request.method == 'GET':
        # Show the create form
        projects = Project.objects.all().order_by('name')
        statuses = ChangeStatus.choices
        risk_levels = RiskLevel.choices
        releases = Release.objects.all().select_related('project').order_by('-update_date')
        organisations = Organisation.objects.all().order_by('name')
        
        context = {
            'change': None,
            'projects': projects,
            'statuses': statuses,
            'risk_levels': risk_levels,
            'releases': releases,
            'organisations': organisations,
        }
        return render(request, 'change_form.html', context)
    
    # Handle POST request
    try:
        project_id = request.POST.get('project')
        if not project_id:
            if request.headers.get('HX-Request'):
                return JsonResponse({'success': False, 'error': 'Project is required'}, status=400)
            else:
                # Regular form submission - re-render form with error
                projects = Project.objects.all().order_by('name')
                statuses = ChangeStatus.choices
                risk_levels = RiskLevel.choices
                releases = Release.objects.all().select_related('project').order_by('-update_date')
                organisations = Organisation.objects.all().order_by('name')
                context = {
                    'change': None,
                    'projects': projects,
                    'statuses': statuses,
                    'risk_levels': risk_levels,
                    'releases': releases,
                    'organisations': organisations,
                    'error': 'Project is required'
                }
                return render(request, 'change_form.html', context)
        
        project = get_object_or_404(Project, id=project_id)
        
        # Get release if provided
        release = None
        release_id = request.POST.get('release')
        if release_id:
            release = get_object_or_404(Release, id=release_id)
        
        # Parse datetime fields
        planned_start = request.POST.get('planned_start')
        planned_end = request.POST.get('planned_end')
        executed_at = request.POST.get('executed_at')
        
        # Get safety relevant flag
        is_safety_relevant = request.POST.get('is_safety_relevant') == 'true'
        
        change = Change.objects.create(
            project=project,
            title=request.POST.get('title'),
            description=request.POST.get('description', ''),
            planned_start=planned_start if planned_start else None,
            planned_end=planned_end if planned_end else None,
            executed_at=executed_at if executed_at else None,
            status=request.POST.get('status', ChangeStatus.DRAFT),
            risk=request.POST.get('risk', RiskLevel.NORMAL),
            risk_description=request.POST.get('risk_description', ''),
            mitigation=request.POST.get('mitigation', ''),
            rollback_plan=request.POST.get('rollback_plan', ''),
            communication_plan=request.POST.get('communication_plan', ''),
            release=release,
            is_safety_relevant=is_safety_relevant,
            created_by=request.user if request.user.is_authenticated else None,
        )
        
        # Set organisations (many-to-many field, must be set after object creation)
        organisation_ids = request.POST.getlist('organisations')
        if organisation_ids:
            change.organisations.set(organisation_ids)
        
        # Sync approvers based on policy
        try:
            sync_result = ChangePolicyService.sync_change_approvers(change)
            logger.info(f"Synced approvers for change {change.id}: {sync_result}")
        except Exception as e:
            logger.error(f"Error syncing approvers for change {change.id}: {e}", exc_info=True)
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='change.created',
            target=change,
            actor=request.user if request.user.is_authenticated else None,
            summary=f'Change "{change.title}" was created'
        )
        
        # Return JSON for HTMX requests, redirect for regular requests
        if request.headers.get('HX-Request'):
            return JsonResponse({
                'success': True,
                'message': 'Change created successfully',
                'change_id': change.id,
                'redirect': f'/changes/{change.id}/'
            })
        else:
            return redirect('change-detail', id=change.id)
    except Exception as e:
        if request.headers.get('HX-Request'):
            return JsonResponse({'success': False, 'error': str(e)}, status=400)
        else:
            projects = Project.objects.all().order_by('name')
            statuses = ChangeStatus.choices
            risk_levels = RiskLevel.choices
            releases = Release.objects.all().select_related('project').order_by('-update_date')
            organisations = Organisation.objects.all().order_by('name')
            context = {
                'change': None,
                'projects': projects,
                'statuses': statuses,
                'risk_levels': risk_levels,
                'releases': releases,
                'organisations': organisations,
                'error': str(e)
            }
            return render(request, 'change_form.html', context)


@login_required
def change_edit(request, id):
    """Change edit page view."""
    change = get_object_or_404(Change, id=id)
    
    if request.method == 'GET':
        # Show the edit form
        projects = Project.objects.all().order_by('name')
        statuses = ChangeStatus.choices
        risk_levels = RiskLevel.choices
        releases = Release.objects.filter(project=change.project).order_by('-update_date')
        organisations = Organisation.objects.all().order_by('name')
        
        context = {
            'change': change,
            'projects': projects,
            'statuses': statuses,
            'risk_levels': risk_levels,
            'releases': releases,
            'organisations': organisations,
        }
        return render(request, 'change_form.html', context)


@login_required

@require_http_methods(["POST"])
def change_update(request, id):
    """Update change details."""
    change = get_object_or_404(Change, id=id)
    
    try:
        # Update basic fields
        change.title = request.POST.get('title', change.title)
        change.description = request.POST.get('description', change.description)
        change.status = request.POST.get('status', change.status)
        change.risk = request.POST.get('risk', change.risk)
        change.risk_description = request.POST.get('risk_description', change.risk_description)
        change.mitigation = request.POST.get('mitigation', change.mitigation)
        change.rollback_plan = request.POST.get('rollback_plan', change.rollback_plan)
        change.communication_plan = request.POST.get('communication_plan', change.communication_plan)
        
        # Update safety flag
        change.is_safety_relevant = request.POST.get('is_safety_relevant') == 'true'
        
        # Update project if changed
        project_id = request.POST.get('project')
        if project_id and int(project_id) != change.project.id:
            change.project = get_object_or_404(Project, id=project_id)
        
        # Update release if provided
        release_id = request.POST.get('release')
        if release_id:
            change.release = get_object_or_404(Release, id=release_id)
        else:
            change.release = None
        
        # Parse datetime fields
        planned_start = request.POST.get('planned_start')
        planned_end = request.POST.get('planned_end')
        executed_at = request.POST.get('executed_at')
        
        change.planned_start = planned_start if planned_start else None
        change.planned_end = planned_end if planned_end else None
        change.executed_at = executed_at if executed_at else None
        
        change.save()
        
        # Update organisations (many-to-many field)
        organisation_ids = request.POST.getlist('organisations')
        if organisation_ids:
            change.organisations.set(organisation_ids)
        else:
            change.organisations.clear()
        
        # Sync approvers based on policy
        try:
            sync_result = ChangePolicyService.sync_change_approvers(change)
            logger.info(f"Synced approvers for change {change.id}: {sync_result}")
        except Exception as e:
            logger.error(f"Error syncing approvers for change {change.id}: {e}", exc_info=True)
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='change.updated',
            target=change,
            actor=request.user if request.user.is_authenticated else None,
            summary=f'Change "{change.title}" was updated'
        )
        
        return JsonResponse({'success': True, 'message': 'Change updated successfully'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required

@require_http_methods(["POST"])
def change_delete(request, id):
    """Delete a change."""
    change = get_object_or_404(Change, id=id)
    
    try:
        title = change.title
        change.delete()
        
        # Log activity - note: can't log to deleted object, so we skip this
        
        return JsonResponse({'success': True, 'redirect': '/changes/'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required
def change_print(request, id):
    """Generate and return PDF report for a change using Weasyprint."""
    from datetime import datetime
    from core.printing import PdfRenderService
    
    # Get the change with all related data
    change = get_object_or_404(
        Change.objects.select_related(
            'project', 'created_by', 'release'
        ).prefetch_related('approvals__approver', 'organisations'),
        id=id
    )
    
    # Get associated items
    items = change.get_associated_items()
    
    # Get approvals
    approvals = change.approvals.select_related('approver').all()
    
    # Get organisations
    organisations = change.organisations.all()
    
    # Get system settings for header/footer
    system_setting = SystemSetting.get_instance()
    
    # Generate filename following pattern: {Change-Referenz}_Change.pdf
    # Change-Referenz format: YYYYMMDD-ID (e.g., 20260209-324)
    change_reference = f"{change.created_at.strftime('%Y%m%d')}-{change.id}"
    filename = f"{change_reference}_Change.pdf"
    
    # Prepare context for template
    context = {
        'change': change,
        'items': items,
        'approvals': approvals,
        'organisations': organisations,
        'now': datetime.now(),
        'system_setting': system_setting,
        'change_reference': change_reference,
    }
    
    # Render PDF using Weasyprint
    service = PdfRenderService()
    result = service.render(
        template_name='printing/change_report.html',
        context=context,
        base_url=request.build_absolute_uri('/'),
        filename=filename
    )
    
    # Create response with PDF
    response = HttpResponse(result.pdf_bytes, content_type=result.content_type)
    response['Content-Disposition'] = f'inline; filename="{result.filename}"'
    
    return response


@login_required
@require_http_methods(["POST"])
def change_add_approver(request, id):
    """Add an approver to a change."""
    change = get_object_or_404(Change, id=id)
    user_id = request.POST.get('user_id')
    is_required = request.POST.get('is_required', 'true').lower() == 'true'
    
    if not user_id:
        return JsonResponse({'success': False, 'error': 'User ID required'}, status=400)
    
    try:
        user = get_object_or_404(User, id=user_id)
        
        # Check if approver already exists
        if ChangeApproval.objects.filter(change=change, approver=user).exists():
            return JsonResponse({'success': False, 'error': 'Approver already exists'}, status=400)
        
        approval = ChangeApproval.objects.create(
            change=change,
            approver=user,
            is_required=is_required,
            status=ApprovalStatus.PENDING
        )
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='change.approver_added',
            target=change,
            actor=request.user if request.user.is_authenticated else None,
            summary=f'{user.name} was added as an approver'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Approver added successfully',
            'reload': True
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required

@require_http_methods(["POST"])
def change_remove_approver(request, id, approval_id):
    """Remove an approver from a change."""
    change = get_object_or_404(Change, id=id)
    approval = get_object_or_404(ChangeApproval, id=approval_id, change=change)
    
    try:
        approver_name = approval.approver.name
        approval.delete()
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='change.approver_removed',
            target=change,
            actor=request.user if request.user.is_authenticated else None,
            summary=f'{approver_name} was removed as an approver'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Approver removed successfully',
            'reload': True
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required
@require_http_methods(["POST"])
def change_approve(request, id, approval_id):
    """Approve a change - Sets status to Accept and requires decision date."""
    change = get_object_or_404(Change, id=id)
    approval = get_object_or_404(ChangeApproval, id=approval_id, change=change)
    
    try:
        # Get decision_date from request - it's required
        decision_date_str = request.POST.get('decision_date', '').strip()
        
        if not decision_date_str:
            return JsonResponse({
                'success': False,
                'error': 'Decision date is required for approval'
            }, status=400)
        
        # Parse the date
        try:
            # Parse ISO format date (YYYY-MM-DD)
            decision_date = datetime.strptime(decision_date_str, '%Y-%m-%d').date()
        except ValueError:
            return JsonResponse({
                'success': False,
                'error': 'Invalid date format. Expected YYYY-MM-DD'
            }, status=400)
        
        # Store decision date in approved_at as date-only (datetime at midnight, timezone-aware)
        decision_datetime = timezone.make_aware(
            datetime.combine(decision_date, time.min)
        )
        approval.approved_at = decision_datetime
        # Set status to Accept
        approval.status = ApprovalStatus.ACCEPT
        # Also set decision_at for compatibility
        approval.decision_at = decision_datetime
        
        approval.save()
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='change.approved',
            target=change,
            actor=request.user if request.user.is_authenticated else None,
            summary=f'{approval.approver.name} approved the change on {decision_date.isoformat()}'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Change approved successfully',
            'reload': True
        })
    except Exception as e:
        logger.exception(f"Error approving change: {e}")
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required
@require_http_methods(["POST"])
def change_reject(request, id, approval_id):
    """Reject a change - Sets status to Reject and requires decision date and comment."""
    change = get_object_or_404(Change, id=id)
    approval = get_object_or_404(ChangeApproval, id=approval_id, change=change)
    
    try:
        # Get comment from request - it's required for reject
        comment = request.POST.get('comment', '').strip()
        
        if not comment:
            return JsonResponse({
                'success': False,
                'error': 'Comment is required for rejection'
            }, status=400)
        
        # Get decision_date from request - it's required
        decision_date_str = request.POST.get('decision_date', '').strip()
        
        if not decision_date_str:
            return JsonResponse({
                'success': False,
                'error': 'Decision date is required for rejection'
            }, status=400)
        
        # Parse the date
        try:
            # Parse ISO format date (YYYY-MM-DD)
            decision_date = datetime.strptime(decision_date_str, '%Y-%m-%d').date()
        except ValueError:
            return JsonResponse({
                'success': False,
                'error': 'Invalid date format. Expected YYYY-MM-DD'
            }, status=400)
        
        # Set status to Reject
        approval.status = ApprovalStatus.REJECT
        # Save the comment
        approval.comment = comment
        # Store decision date in approved_at as date-only (datetime at midnight, timezone-aware)
        decision_datetime = timezone.make_aware(
            datetime.combine(decision_date, time.min)
        )
        approval.approved_at = decision_datetime
        # Also set decision_at for compatibility
        approval.decision_at = decision_datetime
        
        approval.save()
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='change.rejected',
            target=change,
            actor=request.user if request.user.is_authenticated else None,
            summary=f'{approval.approver.name} rejected the change on {decision_date.isoformat()}'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Change rejected',
            'reload': True
        })
    except Exception as e:
        logger.exception(f"Error rejecting change: {e}")
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required
@require_http_methods(["POST"])
def change_abstain(request, id, approval_id):
    """Abstain from a change - Sets status to Abstained and requires decision date."""
    change = get_object_or_404(Change, id=id)
    approval = get_object_or_404(ChangeApproval, id=approval_id, change=change)
    
    try:
        # Get comment from request - it's optional for abstained
        comment = request.POST.get('comment', '').strip()
        
        # Get decision_date from request - it's required
        decision_date_str = request.POST.get('decision_date', '').strip()
        
        if not decision_date_str:
            return JsonResponse({
                'success': False,
                'error': 'Decision date is required for abstention'
            }, status=400)
        
        # Parse the date
        try:
            # Parse ISO format date (YYYY-MM-DD)
            decision_date = datetime.strptime(decision_date_str, '%Y-%m-%d').date()
        except ValueError:
            return JsonResponse({
                'success': False,
                'error': 'Invalid date format. Expected YYYY-MM-DD'
            }, status=400)
        
        # Set status to Abstained
        approval.status = ApprovalStatus.ABSTAINED
        # Save the comment if provided
        if comment:
            approval.comment = comment
        # Store decision date in approved_at as date-only (datetime at midnight, timezone-aware)
        decision_datetime = timezone.make_aware(
            datetime.combine(decision_date, time.min)
        )
        approval.approved_at = decision_datetime
        # Also set decision_at for compatibility
        approval.decision_at = decision_datetime
        
        approval.save()
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='change.abstained',
            target=change,
            actor=request.user if request.user.is_authenticated else None,
            summary=f'{approval.approver.name} abstained from approving the change on {decision_date.isoformat()}'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Abstention recorded',
            'reload': True
        })
    except Exception as e:
        logger.exception(f"Error recording abstention: {e}")
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@require_http_methods(["GET"])
def change_approval_decision(request):
    """
    Public endpoint for email-based approval decision (approve/reject).
    
    Query parameters:
    - token: Decision token (required)
    - change_id: Change ID (required)
    - decision: 'approve' or 'reject' (required)
    
    Returns HTML response with decision confirmation.
    """
    # Get query parameters
    token = request.GET.get('token', '').strip()
    change_id_str = request.GET.get('change_id', '').strip()
    decision = request.GET.get('decision', '').strip().lower()
    
    # Validate parameters
    if not token or not change_id_str or not decision:
        return HttpResponse(
            "<html><body><h1>Error</h1><p>Missing required parameters (token, change_id, or decision).</p></body></html>",
            status=400
        )
    
    if decision not in ['approve', 'reject']:
        return HttpResponse(
            "<html><body><h1>Error</h1><p>Invalid decision. Must be 'approve' or 'reject'.</p></body></html>",
            status=400
        )
    
    # Parse change_id
    try:
        change_id = int(change_id_str)
    except ValueError:
        return HttpResponse(
            "<html><body><h1>Error</h1><p>Invalid change_id format.</p></body></html>",
            status=400
        )
    
    # Lookup approval by token and change_id
    try:
        approval = ChangeApproval.objects.select_related('change', 'approver').get(
            change_id=change_id,
            decision_token=token
        )
    except ChangeApproval.DoesNotExist:
        return HttpResponse(
            "<html><body><h1>Error</h1><p>Invalid or expired approval link.</p></body></html>",
            status=403
        )
    
    # Check if already decided (idempotency guard)
    if approval.status != ApprovalStatus.PENDING:
        return HttpResponse(
            f"""<html><body>
            <h1>Already Decided</h1>
            <p>This approval request has already been processed.</p>
            <p>Current status: <strong>{approval.status}</strong></p>
            <p>Decision made at: <strong>{approval.decision_at}</strong></p>
            </body></html>""",
            status=400
        )
    
    # Apply decision
    now = timezone.now()
    
    if decision == 'approve':
        approval.status = ApprovalStatus.ACCEPT
        approval.decision_at = now
        approval.approved_at = now
        approval.save(update_fields=['status', 'decision_at', 'approved_at'])
        decision_text = "approved"
        decision_emoji = "✅"
    else:  # reject
        approval.status = ApprovalStatus.REJECT
        approval.decision_at = now
        approval.save(update_fields=['status', 'decision_at'])
        decision_text = "rejected"
        decision_emoji = "❌"
    
    # Log activity
    activity_service = ActivityService()
    activity_service.log(
        verb=f'change.{decision_text}_via_email',
        target=approval.change,
        actor=approval.approver,
        summary=f'{approval.approver.name} {decision_text} the change via email link'
    )
    
    # Build response HTML
    response_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Decision Recorded</title>
        <meta charset="utf-8">
        <meta name="viewport" content="width=device-width, initial-scale=1">
        <style>
            body {{
                font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
                max-width: 600px;
                margin: 50px auto;
                padding: 20px;
                background-color: #f5f5f5;
            }}
            .card {{
                background: white;
                border-radius: 8px;
                padding: 30px;
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            }}
            h1 {{
                color: #333;
                margin-top: 0;
            }}
            .emoji {{
                font-size: 48px;
                margin: 20px 0;
            }}
            .decision {{
                color: {'#28a745' if decision == 'approve' else '#dc3545'};
                font-weight: bold;
                font-size: 20px;
            }}
            .info {{
                margin: 20px 0;
                color: #666;
            }}
            .thank-you {{
                margin-top: 30px;
                padding-top: 20px;
                border-top: 1px solid #eee;
                color: #999;
            }}
        </style>
    </head>
    <body>
        <div class="card">
            <div class="emoji">{decision_emoji}</div>
            <h1>Decision Recorded</h1>
            <p class="decision">You have {decision_text} this change.</p>
            <div class="info">
                <p><strong>Change:</strong> {approval.change.title}</p>
                <p><strong>Your decision:</strong> {decision_text.capitalize()}</p>
                <p><strong>Recorded at:</strong> {now.strftime('%Y-%m-%d %H:%M:%S')}</p>
            </div>
            {"<p><strong>Note:</strong> Since you rejected this change, the responsible team will be contacted to discuss next steps.</p>" if decision == 'reject' else ""}
            <div class="thank-you">
                <p>Thank you for your timely response!</p>
            </div>
        </div>
    </body>
    </html>
    """
    
    return HttpResponse(response_html)


@login_required
@require_http_methods(["POST"])
def change_send_approval_requests(request, id):
    """
    Trigger sending approval request emails to all approvers for a Change.
    
    This endpoint is called when the "Get Approvals" button is clicked.
    """
    from core.services.changes.approval_mailer import send_change_approval_request_emails
    from core.services.exceptions import ServiceError, ServiceDisabled, ServiceNotConfigured
    from core.services.config import get_graph_config
    
    change = get_object_or_404(Change, id=id)
    
    # Validate Graph API service is configured before attempting to send emails
    try:
        config = get_graph_config()
        if config is None:
            # Configuration object doesn't exist at all
            return JsonResponse({
                'success': False,
                'error': 'Email service is not configured. Please configure Microsoft Graph API in system settings.'
            }, status=400)
        
        if not config.enabled:
            # Configuration exists but service is disabled
            return JsonResponse({
                'success': False,
                'error': 'Email service is not enabled. Please enable Microsoft Graph API in system settings.'
            }, status=400)
        
        if not config.tenant_id or not config.client_id or not config.client_secret:
            # Service enabled but missing required credentials
            return JsonResponse({
                'success': False,
                'error': 'Email service is not properly configured. Please configure Microsoft Graph API credentials in system settings.'
            }, status=400)
    except Exception as e:
        # Configuration check failures are critical and should be investigated
        logger.error(
            f"Critical error checking Graph API configuration for Change {change.id}: {str(e)}",
            exc_info=True
        )
        return JsonResponse({
            'success': False,
            'error': 'Unable to verify email service configuration. Please contact system administrator.'
        }, status=500)
    
    try:
        # Build base URL for PDF generation and email links
        request_base_url = request.build_absolute_uri('/')
        
        # Send approval request emails
        result = send_change_approval_request_emails(change, request_base_url)
        
        if result['success']:
            # Log activity
            activity_service = ActivityService()
            activity_service.log(
                verb='change.approval_requests_sent',
                target=change,
                actor=request.user,
                summary=f'Approval request emails sent to {result["sent_count"]} approvers'
            )
            
            return JsonResponse({
                'success': True,
                'message': f'Approval request emails sent to {result["sent_count"]} approvers',
                'sent_count': result['sent_count']
            })
        else:
            # Partial failure
            error_msg = f'Sent {result["sent_count"]} emails, but {result["failed_count"]} failed'
            if result['errors']:
                error_msg += f': {"; ".join(result["errors"][:3])}'  # Show first 3 errors
            
            logger.warning(f"Partial failure sending approval requests for Change {change.id}: {error_msg}")
            
            return JsonResponse({
                'success': False,
                'error': error_msg,
                'sent_count': result['sent_count'],
                'failed_count': result['failed_count']
            }, status=500)
    
    except ServiceError as e:
        logger.error(f"Service error sending approval requests for Change {change.id}: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)
    
    except Exception as e:
        logger.exception(f"Unexpected error sending approval requests for Change {change.id}: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': f'Unexpected error: {str(e)}'
        }, status=500)


@login_required
@require_http_methods(["POST"])
def change_send_approval_reminders(request, id):
    """
    Trigger sending reminder emails to all PENDING approvers for a Change.

    This endpoint is called when the "Send Reminder" button is clicked.
    """
    from core.services.changes.approval_mailer import send_change_approval_reminder_emails
    from core.services.exceptions import ServiceError
    from core.services.config import get_graph_config

    change = get_object_or_404(Change, id=id)

    # Validate Graph API service is configured before attempting to send emails
    try:
        config = get_graph_config()
        if config is None:
            return JsonResponse({
                'success': False,
                'error': 'Email service is not configured. Please configure Microsoft Graph API in system settings.'
            }, status=400)

        if not config.enabled:
            return JsonResponse({
                'success': False,
                'error': 'Email service is not enabled. Please enable Microsoft Graph API in system settings.'
            }, status=400)

        if not config.tenant_id or not config.client_id or not config.client_secret:
            return JsonResponse({
                'success': False,
                'error': 'Email service is not properly configured. Please configure Microsoft Graph API credentials in system settings.'
            }, status=400)
    except Exception as e:
        logger.error(
            f"Critical error checking Graph API configuration for Change {change.id}: {str(e)}",
            exc_info=True
        )
        return JsonResponse({
            'success': False,
            'error': 'Unable to verify email service configuration. Please contact system administrator.'
        }, status=500)

    try:
        request_base_url = request.build_absolute_uri('/')

        result = send_change_approval_reminder_emails(change, request_base_url)

        if result['success']:
            activity_service = ActivityService()
            activity_service.log(
                verb='change.approval_reminders_sent',
                target=change,
                actor=request.user,
                summary=f'Approval reminder emails sent to {result["sent_count"]} approvers'
            )

            return JsonResponse({
                'success': True,
                'message': f'Approval reminder emails sent to {result["sent_count"]} approvers',
                'sent_count': result['sent_count']
            })
        else:
            error_msg = f'Sent {result["sent_count"]} emails, but {result["failed_count"]} failed'
            if result['errors']:
                error_msg += f': {"; ".join(result["errors"][:3])}'

            logger.warning(f"Partial failure sending approval reminders for Change {change.id}: {error_msg}")

            return JsonResponse({
                'success': False,
                'error': error_msg,
                'sent_count': result['sent_count'],
                'failed_count': result['failed_count']
            }, status=500)

    except ServiceError as e:
        logger.error(f"Service error sending approval reminders for Change {change.id}: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)

    except Exception as e:
        logger.exception(f"Unexpected error sending approval reminders for Change {change.id}: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': f'Unexpected error: {str(e)}'
        }, status=500)


@login_required
@require_http_methods(["POST"])
def change_update_approver(request, id, approval_id):
    """Update approver details including new fields and attachment."""
    from core.services.storage.service import AttachmentStorageService
    
    change = get_object_or_404(Change, id=id)
    approval = get_object_or_404(ChangeApproval, id=approval_id, change=change)
    
    # Maximum file size: 10MB
    MAX_FILE_SIZE = 10 * 1024 * 1024
    
    try:
        # Update approved_at
        approved_at = request.POST.get('approved_at')
        if approved_at:
            approval.approved_at = approved_at
        else:
            approval.approved_at = None
        
        # Update notes (comment is only updated via approve/reject/abstain actions)
        approval.notes = request.POST.get('notes', '')
        
        approval.save()
        
        # Handle file upload if present
        if 'attachment' in request.FILES:
            file = request.FILES['attachment']
            
            # Validate file size
            if file.size > MAX_FILE_SIZE:
                return JsonResponse({
                    'success': False, 
                    'error': f'File size exceeds maximum allowed size of {MAX_FILE_SIZE // (1024*1024)}MB'
                }, status=400)
            
            # Validate file type
            allowed_extensions = ['.pdf', '.png', '.jpg', '.jpeg', '.txt', '.eml', '.msg']
            file_ext = file.name.lower()[file.name.rfind('.'):] if '.' in file.name else ''
            if file_ext not in allowed_extensions:
                return JsonResponse({
                    'success': False,
                    'error': f'File type not allowed. Allowed types: {", ".join(allowed_extensions)}'
                }, status=400)
            
            # Use storage service to save file
            storage_service = AttachmentStorageService()
            attachment = storage_service.save_file(
                file=file,
                user=request.user,
                original_name=file.name
            )
            
            # Create attachment link
            approval_ct = ContentType.objects.get_for_model(ChangeApproval)
            
            AttachmentLink.objects.create(
                attachment=attachment,
                target_content_type=approval_ct,
                target_object_id=approval.id,
                role=AttachmentRole.APPROVER_ATTACHMENT
            )
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='change.approver_updated',
            target=change,
            actor=request.user if request.user.is_authenticated else None,
            summary=f'Approver {approval.approver.name} details updated'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Approver updated successfully',
            'reload': True
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required
@require_http_methods(["POST"])
def change_remove_approver_attachment(request, id, approval_id, attachment_id):
    """Remove an attachment from an approver."""
    change = get_object_or_404(Change, id=id)
    approval = get_object_or_404(ChangeApproval, id=approval_id, change=change)
    attachment = get_object_or_404(Attachment, id=attachment_id)
    
    try:
        # Find and remove the attachment link
        approval_ct = ContentType.objects.get_for_model(ChangeApproval)
        
        link = AttachmentLink.objects.filter(
            attachment=attachment,
            target_content_type=approval_ct,
            target_object_id=approval.id,
            role=AttachmentRole.APPROVER_ATTACHMENT
        ).first()
        
        if link:
            link.delete()
            # Mark attachment as deleted (soft delete)
            attachment.is_deleted = True
            attachment.save()
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='change.approver_attachment_removed',
            target=change,
            actor=request.user if request.user.is_authenticated else None,
            summary=f'Attachment removed from approver {approval.approver.name}'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Attachment removed successfully',
            'reload': True
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@login_required
@require_http_methods(["POST"])
def change_polish_risk_description(request, id):
    """Polish the risk description using AI agent."""
    change = get_object_or_404(Change, id=id)
    
    try:
        # Get current risk description
        risk_description = change.risk_description or ''
        
        if not risk_description.strip():
            return JsonResponse({
                'success': False, 
                'error': 'Risk description is empty. Please add some text first.'
            }, status=400)
        
        # Execute the change-text-polish-agent
        agent_service = AgentService()
        polished_text = agent_service.execute_agent(
            filename='change-text-polish-agent.yml',
            input_text=risk_description,
            user=request.user if request.user.is_authenticated else None,
            client_ip=request.META.get('REMOTE_ADDR')
        )
        
        # Update the risk description
        change.risk_description = polished_text
        change.save()
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='change.ai_risk_polished',
            target=change,
            actor=request.user if request.user.is_authenticated else None,
            summary=f'AI polished risk description for change "{change.title}"'
        )
        
        return JsonResponse({
            'success': True,
            'text': polished_text,
            'message': 'Risk description improved successfully'
        })
    except Exception as e:
        logger.error(f"Error polishing risk description: {e}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
@require_http_methods(["POST"])
def change_optimize_mitigation(request, id):
    """Optimize the mitigation plan using AI agent."""
    change = get_object_or_404(Change, id=id)
    
    try:
        # Get current mitigation plan
        mitigation = change.mitigation or ''
        
        if not mitigation.strip():
            return JsonResponse({
                'success': False, 
                'error': 'Mitigation plan is empty. Please add some text first.'
            }, status=400)
        
        # Execute the text-optimization-agent
        agent_service = AgentService()
        optimized_text = agent_service.execute_agent(
            filename='text-optimization-agent.yml',
            input_text=mitigation,
            user=request.user if request.user.is_authenticated else None,
            client_ip=request.META.get('REMOTE_ADDR')
        )
        
        # Update the mitigation plan
        change.mitigation = optimized_text
        change.save()
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='change.ai_mitigation_optimized',
            target=change,
            actor=request.user if request.user.is_authenticated else None,
            summary=f'AI optimized mitigation plan for change "{change.title}"'
        )
        
        return JsonResponse({
            'success': True,
            'text': optimized_text,
            'message': 'Mitigation plan optimized successfully'
        })
    except Exception as e:
        logger.error(f"Error optimizing mitigation plan: {e}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
@require_http_methods(["POST"])
def change_optimize_rollback(request, id):
    """Optimize the rollback plan using AI agent."""
    change = get_object_or_404(Change, id=id)
    
    try:
        # Get current rollback plan
        rollback_plan = change.rollback_plan or ''
        
        if not rollback_plan.strip():
            return JsonResponse({
                'success': False, 
                'error': 'Rollback plan is empty. Please add some text first.'
            }, status=400)
        
        # Execute the text-optimization-agent
        agent_service = AgentService()
        optimized_text = agent_service.execute_agent(
            filename='text-optimization-agent.yml',
            input_text=rollback_plan,
            user=request.user if request.user.is_authenticated else None,
            client_ip=request.META.get('REMOTE_ADDR')
        )
        
        # Update the rollback plan
        change.rollback_plan = optimized_text
        change.save()
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='change.ai_rollback_optimized',
            target=change,
            actor=request.user if request.user.is_authenticated else None,
            summary=f'AI optimized rollback plan for change "{change.title}"'
        )
        
        return JsonResponse({
            'success': True,
            'text': optimized_text,
            'message': 'Rollback plan optimized successfully'
        })
    except Exception as e:
        logger.error(f"Error optimizing rollback plan: {e}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
@require_http_methods(["POST"])
def change_assess_risk(request, id):
    """Assess risk level automatically using AI agent."""
    change = get_object_or_404(Change, id=id)
    
    try:
        # Combine all relevant fields for risk assessment
        risk_description = change.risk_description or ''
        mitigation = change.mitigation or ''
        rollback_plan = change.rollback_plan or ''
        
        # Build input for the agent
        agent_input = f"""Risk Description:
{risk_description}

Mitigation Plan:
{mitigation}

Rollback Plan:
{rollback_plan}"""
        
        if not risk_description.strip() and not mitigation.strip() and not rollback_plan.strip():
            return JsonResponse({
                'success': False, 
                'error': 'At least one of Risk Description, Mitigation Plan, or Rollback Plan must have content.'
            }, status=400)
        
        # Execute the change-risk-assessment-agent
        agent_service = AgentService()
        assessment_result = agent_service.execute_agent(
            filename='change-risk-assessment-agent.yml',
            input_text=agent_input,
            user=request.user if request.user.is_authenticated else None,
            client_ip=request.META.get('REMOTE_ADDR')
        )
        
        # Parse the result to extract risk level and reason
        # Expected format can be JSON with RiskClass and RiskClassReason fields
        risk_class = RiskLevel.NORMAL  # Default to NORMAL if no risk class can be determined
        risk_reason = None  # Will be set based on parsing
        
        # Try to parse as JSON first
        try:
            # Attempt to parse JSON response
            result_json = json.loads(assessment_result)
            
            # Extract RiskClassReason for the description
            # ONLY use RiskClassReason if present, otherwise use full JSON
            risk_reason = result_json.get('RiskClassReason', assessment_result)
            
            # Extract and normalize RiskClass for the enum
            if 'RiskClass' in result_json and result_json['RiskClass']:
                risk_class_value = result_json['RiskClass'].lower().strip()
                
                # Normalize to RiskLevel enum values
                if risk_class_value in ['very high', 'veryhigh', 'sehr hoch']:
                    risk_class = RiskLevel.VERY_HIGH
                elif risk_class_value in ['low', 'niedrig', 'gering']:
                    risk_class = RiskLevel.LOW
                elif risk_class_value in ['high', 'hoch']:
                    risk_class = RiskLevel.HIGH
                elif risk_class_value in ['normal', 'mittel']:
                    risk_class = RiskLevel.NORMAL
                # If unrecognized value, keep default NORMAL (set above)
        except (json.JSONDecodeError, AttributeError):
            # If not JSON, assessment_result is None, or .lower() fails on None,
            # fall back to text parsing
            risk_reason = assessment_result  # Use full text for non-JSON responses
            assessment_lower = assessment_result.lower() if assessment_result else ''
            if 'very high' in assessment_lower or 'veryhigh' in assessment_lower or 'sehr hoch' in assessment_lower:
                risk_class = RiskLevel.VERY_HIGH
            elif 'low' in assessment_lower or 'niedrig' in assessment_lower or 'gering' in assessment_lower:
                risk_class = RiskLevel.LOW
            elif 'high' in assessment_lower or 'hoch' in assessment_lower:
                # This check is after "very high" to avoid false matches
                risk_class = RiskLevel.HIGH
            # If no match, keep default NORMAL (set above)
        
        # Update risk level
        old_risk = change.risk
        change.risk = risk_class
        
        # Update or append the reasoning to risk description
        # Check if an AI assessment section already exists and replace it
        risk_desc = change.risk_description or ''
        ai_marker = '## AI Risk Assessment'
        
        if ai_marker in risk_desc:
            # Find and replace existing AI assessment
            parts = risk_desc.split(ai_marker)
            # Keep everything before the marker, append new assessment
            change.risk_description = f"{parts[0].rstrip()}\n\n{ai_marker}\n{risk_reason}"
        elif risk_desc:
            # Append new assessment to existing description
            change.risk_description = f"{risk_desc}\n\n{ai_marker}\n{risk_reason}"
        else:
            # No existing description, just add the assessment
            change.risk_description = f"{ai_marker}\n{risk_reason}"
        
        change.save()
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='change.ai_risk_assessed',
            target=change,
            actor=request.user if request.user.is_authenticated else None,
            summary=f'AI assessed risk level for change "{change.title}" from {old_risk} to {risk_class}'
        )
        
        return JsonResponse({
            'success': True,
            'risk_class': risk_class,
            'risk_class_display': change.get_risk_display(),
            'risk_reason': risk_reason,
            'message': f'Risk level set to {change.get_risk_display()}'
        })
    except Exception as e:
        logger.error(f"Error assessing risk: {e}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


def get_human_readable_verb(verb):
    """Convert activity verb to human-readable text."""
    verb_mapping = {
        # Items
        'item.created': 'Item created',
        'item.status_changed': 'Status changed',
        'item.assigned': 'Item assigned',
        'item.updated': 'Item updated',
        
        # Projects
        'project.created': 'Project created',
        'project.status_changed': 'Project status changed',
        'project.updated': 'Project updated',
        
        # GitHub
        'github.issue_created': 'GitHub issue created',
        'github.pr_created': 'GitHub PR created',
        'github.mapping_synced': 'GitHub mapping synced',
        'github.linked': 'GitHub linked',
        
        # Comments
        'comment.added': 'Comment added',
        'comment.updated': 'Comment updated',
        'comment.deleted': 'Comment deleted',
        
        # Attachments
        'attachment.uploaded': 'Attachment uploaded',
        'attachment.deleted': 'Attachment deleted',
        
        # Changes
        'change.created': 'Change created',
        'change.status_changed': 'Change status changed',
        'change.approved': 'Change approved',
        'change.rejected': 'Change rejected',
        
        # AI
        'ai.job_completed': 'AI job completed',
        'ai.job_failed': 'AI job failed',
        
        # Graph API
        'graph.mail_sent': 'Email sent',
        
        # Default fallback
        'created': 'Created',
        'updated': 'Updated',
        'deleted': 'Deleted',
        'approved': 'Approved',
        'rejected': 'Rejected',
    }
    
    return verb_mapping.get(verb, verb.replace('_', ' ').replace('.', ' ').title())


@login_required
def dashboard_in_progress_items(request):
    """HTMX partial for in-progress items list."""
    items = Item.objects.filter(
        status__in=[ItemStatus.WORKING, ItemStatus.TESTING, ItemStatus.READY_FOR_RELEASE]
    ).select_related(
        'project', 'type', 'organisation', 'requester', 'assigned_to'
    ).order_by('-updated_at')[:20]
    
    context = {
        'items': items,
    }
    return render(request, 'partials/dashboard_in_progress.html', context)


@login_required
def dashboard_activity_stream(request):
    """HTMX partial for global activity stream."""
    from django.utils.timesince import timesince
    
    # Get filter parameter
    filter_type = request.GET.get('filter', 'all')
    offset = int(request.GET.get('offset', 0))
    limit = 50
    
    # Base queryset
    activity_service = ActivityService()
    activities = Activity.objects.select_related('actor', 'target_content_type').order_by('-created_at')
    
    # Apply filter
    if filter_type and filter_type != 'all':
        if filter_type == 'items':
            item_ct = ContentType.objects.get_for_model(Item)
            activities = activities.filter(target_content_type=item_ct)
        elif filter_type == 'projects':
            project_ct = ContentType.objects.get_for_model(Project)
            activities = activities.filter(target_content_type=project_ct)
        elif filter_type == 'github':
            activities = activities.filter(verb__startswith='github.')
        elif filter_type == 'changes':
            change_ct = ContentType.objects.get_for_model(Change)
            activities = activities.filter(target_content_type=change_ct)
        elif filter_type == 'ai':
            activities = activities.filter(verb__startswith='ai.')
    
    # Paginate
    activities = activities[offset:offset + limit]
    
    # Build activity list with human-readable verbs and relative times
    activity_list = []
    for activity in activities:
        # Get target URL if possible
        target_url = None
        target_title = None
        if activity.target_content_type.model == 'item' and activity.target_object_id:
            try:
                item = Item.objects.get(id=activity.target_object_id)
                target_url = f'/items/{item.id}/'
                target_title = item.title
            except Item.DoesNotExist:
                pass
        elif activity.target_content_type.model == 'project' and activity.target_object_id:
            try:
                project = Project.objects.get(id=activity.target_object_id)
                target_url = f'/projects/{project.id}/'
                target_title = project.name
            except Project.DoesNotExist:
                pass
        elif activity.target_content_type.model == 'change' and activity.target_object_id:
            try:
                change = Change.objects.get(id=activity.target_object_id)
                target_url = f'/changes/{change.id}/'
                target_title = change.title
            except Change.DoesNotExist:
                pass
        
        activity_list.append({
            'id': activity.id,
            'verb': get_human_readable_verb(activity.verb),
            'actor': activity.actor,
            'summary': activity.summary,
            'created_at': activity.created_at,
            'time_ago': timesince(activity.created_at),
            'target_url': target_url,
            'target_title': target_title,
        })
    
    context = {
        'activities': activity_list,
        'filter_type': filter_type,
        'offset': offset,
        'limit': limit,
        'has_more': len(activities) == limit,
        'next_offset': offset + limit,
    }
    return render(request, 'partials/dashboard_activity_stream.html', context)


@login_required
def search(request):
    """
    Global search view using Weaviate hybrid search.
    
    Query parameters:
    - q: Search query (required, minimum 2 characters)
    - mode: Search mode ('hybrid', 'similar', 'keyword'; default: 'hybrid')
    - type: Optional filter by object type (e.g., 'item', 'project', 'comment')
    - project_id: Optional filter by project ID
    """
    from django.conf import settings
    from core.services.weaviate.service import global_search
    
    query = request.GET.get('q', '').strip()
    search_mode = request.GET.get('mode', 'hybrid').strip()
    object_type = request.GET.get('type', '').strip()
    project_id = request.GET.get('project_id', '').strip()
    
    # Validate search mode
    valid_modes = ['hybrid', 'similar', 'keyword']
    if search_mode not in valid_modes:
        search_mode = 'hybrid'
    
    # Get configuration from settings
    min_query_length = getattr(settings, 'WEAVIATE_SEARCH_MIN_QUERY_LENGTH', 2)
    search_limit = getattr(settings, 'WEAVIATE_SEARCH_LIMIT', 25)
    search_alpha = getattr(settings, 'WEAVIATE_SEARCH_ALPHA', 0.5)
    
    # Initialize context
    context = {
        'query': query,
        'search_mode': search_mode,
        'object_type': object_type,
        'project_id': project_id,
        'results': [],
        'error': None,
        'min_query_length': min_query_length,
    }
    
    # Only search if query is provided and meets minimum length
    if query:
        if len(query) < min_query_length:
            context['error'] = f'Bitte mindestens {min_query_length} Zeichen eingeben.'
        else:
            try:
                # Build filters
                filters = {}
                if object_type:
                    filters['type'] = object_type
                if project_id:
                    filters['project_id'] = project_id
                
                # Execute search with mode
                results = global_search(
                    query=query,
                    limit=search_limit,
                    alpha=search_alpha,
                    filters=filters if filters else None,
                    mode=search_mode,
                )
                
                context['results'] = results
                
            except Exception as e:
                # Log error without including user query for security
                logger.exception("Search error occurred")
                context['error'] = 'Suchfehler: Weaviate ist möglicherweise nicht verfügbar. Bitte versuchen Sie es später erneut.'
    
    return render(request, 'search.html', context)


# ===========================
# Mail Template Views
# ===========================

@login_required
def mail_templates(request):
    """Mail templates list view with search and filter."""
    templates = MailTemplate.objects.all()
    
    # Search by key or subject
    search_query = request.GET.get('q', '').strip()
    if search_query:
        templates = templates.filter(
            Q(key__icontains=search_query) | Q(subject__icontains=search_query)
        )
    
    # Filter by is_active
    is_active_filter = request.GET.get('is_active', '').strip()
    if is_active_filter == 'true':
        templates = templates.filter(is_active=True)
    elif is_active_filter == 'false':
        templates = templates.filter(is_active=False)
    
    context = {
        'templates': templates,
        'search_query': search_query,
        'is_active_filter': is_active_filter,
    }
    return render(request, 'mail_templates.html', context)


@login_required
def mail_template_detail(request, id):
    """Mail template detail view."""
    template = get_object_or_404(MailTemplate, id=id)
    
    context = {
        'template': template,
    }
    return render(request, 'mail_template_detail.html', context)


@login_required
def mail_template_create(request):
    """Show create form for new mail template."""
    context = {
        'template': None,
    }
    return render(request, 'mail_template_form.html', context)


@login_required
def mail_template_edit(request, id):
    """Show edit form for existing mail template."""
    template = get_object_or_404(MailTemplate, id=id)
    
    context = {
        'template': template,
    }
    return render(request, 'mail_template_form.html', context)


@login_required
@require_http_methods(["POST"])
def mail_template_update(request, id):
    """Create or update a mail template."""
    try:
        # Parse form data
        key = request.POST.get('key', '').strip()
        subject = request.POST.get('subject', '').strip()
        message = request.POST.get('message', '').strip()
        from_name = request.POST.get('from_name', '').strip()
        from_address = request.POST.get('from_address', '').strip()
        cc_address = request.POST.get('cc_address', '').strip()
        is_active = request.POST.get('is_active') == 'on'
        action = request.POST.get('action', 'save')
        
        # Validate required fields
        if not key:
            return JsonResponse({'success': False, 'error': 'Key is required'}, status=400)
        if not subject:
            return JsonResponse({'success': False, 'error': 'Subject is required'}, status=400)
        if not message:
            return JsonResponse({'success': False, 'error': 'Message is required'}, status=400)
        
        # Validate key format (lowercase, numbers, hyphens only)
        if not re.match(r'^[a-z0-9\-]+$', key):
            return JsonResponse({
                'success': False, 
                'error': 'Key must contain only lowercase letters, numbers, and hyphens'
            }, status=400)
        
        # Create or update
        if id == 0:
            # Create new template
            # Check if key already exists
            if MailTemplate.objects.filter(key=key).exists():
                return JsonResponse({
                    'success': False, 
                    'error': f'A template with key "{key}" already exists'
                }, status=400)
            
            template = MailTemplate.objects.create(
                key=key,
                subject=subject,
                message=message,
                from_name=from_name,
                from_address=from_address,
                cc_address=cc_address,
                is_active=is_active
            )
            
            # Log activity
            activity_service = ActivityService()
            activity_service.log(
                verb='mail_template.created',
                target=template,
                actor=request.user if request.user.is_authenticated else None,
                summary=f'Created mail template "{template.key}"'
            )
            
            message_text = f'Mail template "{key}" created successfully'
        else:
            # Update existing template
            template = get_object_or_404(MailTemplate, id=id)
            
            # Update fields
            template.subject = subject
            template.message = message
            template.from_name = from_name
            template.from_address = from_address
            template.cc_address = cc_address
            template.is_active = is_active
            template.save()
            
            # Log activity
            activity_service = ActivityService()
            activity_service.log(
                verb='mail_template.updated',
                target=template,
                actor=request.user if request.user.is_authenticated else None,
                summary=f'Updated mail template "{template.key}"'
            )
            
            message_text = f'Mail template "{template.key}" updated successfully'
        
        # Determine redirect based on action
        if action == 'save_close':
            redirect_url = reverse('mail-templates')
        else:
            redirect_url = reverse('mail-template-detail', args=[template.id])
        
        return JsonResponse({
            'success': True,
            'message': message_text,
            'redirect': redirect_url,
            'template_id': template.id
        })
        
    except Exception as e:
        logger.error(f"Error saving mail template: {e}", exc_info=True)
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
@require_http_methods(["POST"])
def mail_template_delete(request, id):
    """Delete a mail template."""
    # get_object_or_404 will raise Http404 if not found, which Django will handle
    template = get_object_or_404(MailTemplate, id=id)
    template_key = template.key
    
    try:
        # Log activity before deletion
        activity_service = ActivityService()
        activity_service.log(
            verb='mail_template.deleted',
            target=template,
            actor=request.user if request.user.is_authenticated else None,
            summary=f'Deleted mail template "{template_key}"'
        )
        
        template.delete()
        
        return JsonResponse({
            'success': True,
            'message': f'Mail template "{template_key}" deleted successfully',
            'redirect': reverse('mail-templates')
        })
        
    except Exception as e:
        logger.error(f"Error deleting mail template: {e}", exc_info=True)
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
@require_http_methods(["POST"])
def mail_template_generate_ai(request, id):
    """Generate subject and message using AI agent."""
    try:
        # Parse JSON body
        data = json.loads(request.body)
        context = data.get('context', '').strip()
        
        if not context:
            return JsonResponse({
                'success': False, 
                'error': 'Context description is required'
            }, status=400)
        
        # Execute the create-mail-template agent
        agent_service = AgentService()
        result = agent_service.execute_agent(
            filename='create-mail-template.yml',
            input_text=context,
            user=request.user if request.user.is_authenticated else None,
            client_ip=request.META.get('REMOTE_ADDR')
        )
        
        # Parse JSON response from agent
        try:
            # Clean up potential markdown code blocks
            result = result.strip()
            if result.startswith('```json'):
                result = result[7:]
            if result.startswith('```'):
                result = result[3:]
            if result.endswith('```'):
                result = result[:-3]
            result = result.strip()
            
            result_json = json.loads(result)
            
            # Extract Subject and Message
            subject = result_json.get('Subject', '')
            message = result_json.get('Message', '')
            
            if not subject or not message:
                return JsonResponse({
                    'success': False, 
                    'error': 'AI agent did not return valid Subject and Message fields'
                }, status=500)
            
            # Log activity if template exists
            if id > 0:
                template = get_object_or_404(MailTemplate, id=id)
                activity_service = ActivityService()
                activity_service.log(
                    verb='mail_template.ai_generated',
                    target=template,
                    actor=request.user if request.user.is_authenticated else None,
                    summary=f'AI generated content for mail template "{template.key}"'
                )
            
            return JsonResponse({
                'success': True,
                'subject': subject,
                'message': message
            })
            
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse AI agent response as JSON: {result}")
            return JsonResponse({
                'success': False, 
                'error': f'AI agent returned invalid JSON: {str(e)}'
            }, status=500)
        
    except Exception as e:
        logger.error(f"Error generating mail template with AI: {e}", exc_info=True)
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


# Mail Action Mapping Views

@login_required
def mail_action_mappings(request):
    """Mail action mappings list view with search and filter."""
    mappings = MailActionMapping.objects.select_related('item_type', 'mail_template').all()
    
    # Search by status or type name
    search_query = request.GET.get('q', '').strip()
    if search_query:
        mappings = mappings.filter(
            Q(item_status__icontains=search_query) | 
            Q(item_type__name__icontains=search_query) |
            Q(item_type__key__icontains=search_query) |
            Q(mail_template__key__icontains=search_query)
        )
    
    # Filter by status
    status_filter = request.GET.get('status', '').strip()
    if status_filter:
        mappings = mappings.filter(item_status=status_filter)
    
    # Filter by type
    type_filter = request.GET.get('type', '').strip()
    if type_filter:
        try:
            type_id = int(type_filter)
            mappings = mappings.filter(item_type_id=type_id)
        except ValueError:
            pass
    
    # Filter by is_active
    is_active_filter = request.GET.get('is_active', '').strip()
    if is_active_filter == 'true':
        mappings = mappings.filter(is_active=True)
    elif is_active_filter == 'false':
        mappings = mappings.filter(is_active=False)
    
    # Apply consistent ordering
    mappings = mappings.order_by('item_status', 'item_type__name')
    
    # Get all item types and statuses for filters
    item_types = ItemType.objects.filter(is_active=True).order_by('name')
    item_statuses = ItemStatus.choices
    
    context = {
        'mappings': mappings,
        'search_query': search_query,
        'status_filter': status_filter,
        'type_filter': type_filter,
        'is_active_filter': is_active_filter,
        'item_types': item_types,
        'item_statuses': item_statuses,
    }
    return render(request, 'mail_action_mappings.html', context)


@login_required
def mail_action_mapping_detail(request, id):
    """Mail action mapping detail view."""
    mapping = get_object_or_404(MailActionMapping.objects.select_related('item_type', 'mail_template'), id=id)
    
    context = {
        'mapping': mapping,
    }
    return render(request, 'mail_action_mapping_detail.html', context)


@login_required
def mail_action_mapping_create(request):
    """Show create form for new mail action mapping."""
    # Get all active item types and mail templates
    item_types = ItemType.objects.filter(is_active=True).order_by('name')
    mail_templates = MailTemplate.objects.filter(is_active=True).order_by('key')
    item_statuses = ItemStatus.choices
    
    context = {
        'mapping': None,
        'item_types': item_types,
        'mail_templates': mail_templates,
        'item_statuses': item_statuses,
    }
    return render(request, 'mail_action_mapping_form.html', context)


@login_required
def mail_action_mapping_edit(request, id):
    """Show edit form for existing mail action mapping."""
    mapping = get_object_or_404(MailActionMapping.objects.select_related('item_type', 'mail_template'), id=id)
    
    # Get all active item types and mail templates
    item_types = ItemType.objects.filter(is_active=True).order_by('name')
    mail_templates = MailTemplate.objects.filter(is_active=True).order_by('key')
    item_statuses = ItemStatus.choices
    
    context = {
        'mapping': mapping,
        'item_types': item_types,
        'mail_templates': mail_templates,
        'item_statuses': item_statuses,
    }
    return render(request, 'mail_action_mapping_form.html', context)


@login_required
@require_http_methods(["POST"])
def mail_action_mapping_update(request, id):
    """Create or update a mail action mapping with uniqueness validation."""
    try:
        # Parse form data
        item_status = request.POST.get('item_status', '').strip()
        item_type_id = request.POST.get('item_type', '').strip()
        mail_template_id = request.POST.get('mail_template', '').strip()
        is_active = request.POST.get('is_active') == 'on'
        action = request.POST.get('action', 'save')
        
        # Validate required fields
        if not item_status:
            return JsonResponse({'success': False, 'error': 'Status ist erforderlich'}, status=400)
        if not item_type_id:
            return JsonResponse({'success': False, 'error': 'Typ ist erforderlich'}, status=400)
        if not mail_template_id:
            return JsonResponse({'success': False, 'error': 'Mail-Template ist erforderlich'}, status=400)
        
        # Validate item_status is valid choice
        valid_statuses = [choice[0] for choice in ItemStatus.choices]
        if item_status not in valid_statuses:
            return JsonResponse({'success': False, 'error': 'Ungültiger Status'}, status=400)
        
        # Get related objects
        try:
            item_type = ItemType.objects.get(id=int(item_type_id))
        except (ItemType.DoesNotExist, ValueError):
            return JsonResponse({'success': False, 'error': 'Ungültiger Typ'}, status=400)
        
        try:
            mail_template = MailTemplate.objects.get(id=int(mail_template_id))
        except (MailTemplate.DoesNotExist, ValueError):
            return JsonResponse({'success': False, 'error': 'Ungültiges Mail-Template'}, status=400)
        
        # Check for uniqueness (status + type combination)
        existing_mapping = MailActionMapping.objects.filter(
            item_status=item_status,
            item_type=item_type
        )
        
        # If editing, exclude the current mapping from uniqueness check
        if id != 0:
            existing_mapping = existing_mapping.exclude(id=id)
        
        if existing_mapping.exists():
            status_display = dict(ItemStatus.choices).get(item_status, item_status)
            return JsonResponse({
                'success': False,
                'error': f'Ein Mapping für Status "{status_display}" und Typ "{item_type.name}" existiert bereits.'
            }, status=400)
        
        # Create or update
        if id == 0:
            # Create new mapping
            mapping = MailActionMapping.objects.create(
                item_status=item_status,
                item_type=item_type,
                mail_template=mail_template,
                is_active=is_active
            )
            
            # Log activity
            activity_service = ActivityService()
            activity_service.log(
                verb='mail_action_mapping.created',
                target=mapping,
                actor=request.user if request.user.is_authenticated else None,
                summary=f'Created mail action mapping: {mapping}'
            )
            
            message_text = f'Mail action mapping created successfully'
        else:
            # Update existing mapping
            mapping = get_object_or_404(MailActionMapping, id=id)
            
            # Update fields
            mapping.item_status = item_status
            mapping.item_type = item_type
            mapping.mail_template = mail_template
            mapping.is_active = is_active
            mapping.save()
            
            # Log activity
            activity_service = ActivityService()
            activity_service.log(
                verb='mail_action_mapping.updated',
                target=mapping,
                actor=request.user if request.user.is_authenticated else None,
                summary=f'Updated mail action mapping: {mapping}'
            )
            
            message_text = f'Mail action mapping updated successfully'
        
        # Determine redirect based on action
        if action == 'save_close':
            redirect_url = reverse('mail-action-mappings')
        else:
            redirect_url = reverse('mail-action-mapping-detail', args=[mapping.id])
        
        return JsonResponse({
            'success': True,
            'message': message_text,
            'redirect': redirect_url,
            'mapping_id': mapping.id
        })
        
    except Exception as e:
        logger.error(f"Error saving mail action mapping: {e}", exc_info=True)
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
@require_http_methods(["POST"])
def mail_action_mapping_delete(request, id):
    """Delete a mail action mapping."""
    mapping = get_object_or_404(MailActionMapping, id=id)
    mapping_str = str(mapping)
    
    try:
        # Log activity before deletion
        activity_service = ActivityService()
        activity_service.log(
            verb='mail_action_mapping.deleted',
            target=mapping,
            actor=request.user if request.user.is_authenticated else None,
            summary=f'Deleted mail action mapping: {mapping_str}'
        )
        
        mapping.delete()
        
        return JsonResponse({
            'success': True,
            'message': f'Mail action mapping deleted successfully',
            'redirect': reverse('mail-action-mappings')
        })
        
    except Exception as e:
        logger.error(f"Error deleting mail action mapping: {e}", exc_info=True)
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


# Email Reply/Forward Views
@login_required
def email_prepare_reply(request, comment_id):
    """Prepare reply data for an email comment."""
    from core.services.mail.email_reply_service import prepare_reply
    
    try:
        comment = get_object_or_404(ItemComment, id=comment_id)
        
        # Only allow reply for email comments
        if comment.kind not in ['EmailIn', 'EmailOut']:
            return JsonResponse({'success': False, 'error': 'Can only reply to email comments'}, status=400)
        
        # Prepare reply data
        reply_data = prepare_reply(comment, current_user=request.user)
        
        return JsonResponse({
            'success': True,
            'data': reply_data,
        })
    except Exception as e:
        logger.error(f"Error preparing reply: {e}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
def email_prepare_reply_all(request, comment_id):
    """Prepare reply-all data for an email comment."""
    from core.services.mail.email_reply_service import prepare_reply_all
    
    try:
        comment = get_object_or_404(ItemComment, id=comment_id)
        
        # Only allow reply for email comments
        if comment.kind not in ['EmailIn', 'EmailOut']:
            return JsonResponse({'success': False, 'error': 'Can only reply to email comments'}, status=400)
        
        # Prepare reply-all data
        reply_data = prepare_reply_all(comment, current_user=request.user)
        
        return JsonResponse({
            'success': True,
            'data': reply_data,
        })
    except Exception as e:
        logger.error(f"Error preparing reply all: {e}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
def email_prepare_forward(request, comment_id):
    """Prepare forward data for an email comment."""
    from core.services.mail.email_reply_service import prepare_forward
    
    try:
        comment = get_object_or_404(ItemComment, id=comment_id)
        
        # Only allow forward for email comments
        if comment.kind not in ['EmailIn', 'EmailOut']:
            return JsonResponse({'success': False, 'error': 'Can only forward email comments'}, status=400)
        
        # Prepare forward data
        forward_data = prepare_forward(comment, current_user=request.user)
        
        return JsonResponse({
            'success': True,
            'data': forward_data,
        })
    except Exception as e:
        logger.error(f"Error preparing forward: {e}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
@require_POST
def email_send_reply(request):
    """Send a reply/forward email from the compose modal."""
    from core.services.graph.mail_service import send_email
    from core.utils.html_sanitization import sanitize_html
    import json
    
    try:
        data = json.loads(request.body)
        
        item_id = data.get('item_id')
        to_addresses = data.get('to', [])
        cc_addresses = data.get('cc', [])
        subject = data.get('subject', '')
        body = data.get('body', '')
        in_reply_to = data.get('in_reply_to', '')
        
        # Validate inputs
        if not item_id:
            return JsonResponse({'success': False, 'error': 'Item ID is required'}, status=400)
        
        if not to_addresses or not isinstance(to_addresses, list):
            return JsonResponse({'success': False, 'error': 'At least one recipient is required'}, status=400)
        
        if not subject:
            return JsonResponse({'success': False, 'error': 'Subject is required'}, status=400)
        
        # Get item
        item = get_object_or_404(Item, id=item_id)
        
        # Sanitize HTML body to prevent XSS attacks while preserving formatting
        sanitized_body = sanitize_html(body)
        
        # Send email
        result = send_email(
            subject=subject,
            body=sanitized_body,
            to=to_addresses,
            cc=cc_addresses if cc_addresses else None,
            body_is_html=True,
            item=item,
            author=request.user,
            visibility='Internal',
        )
        
        if result.success:
            # Log activity
            activity_service = ActivityService()
            activity_service.log(
                verb='email.sent',
                target=item,
                actor=request.user,
                summary=f"Sent email: {subject}",
            )
            
            return JsonResponse({
                'success': True,
                'message': 'Email sent successfully'
            })
        else:
            return JsonResponse({
                'success': False,
                'error': result.error or 'Failed to send email'
            }, status=500)
            
    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'error': 'Invalid JSON data'}, status=400)
    except Exception as e:
        logger.error(f"Error sending email: {e}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


# ==================== Global Settings Views ====================

@login_required
def global_settings_detail(request):
    """Global Settings detail view (singleton)."""
    settings = GlobalSettings.get_instance()
    
    context = {
        'settings': settings,
    }
    return render(request, 'global_settings_detail.html', context)


@login_required
@require_http_methods(["POST"])
def global_settings_update(request):
    """Update Global Settings."""
    settings = GlobalSettings.get_instance()
    
    try:
        # Update fields only if they have a value (not empty string)
        # This allows partial updates and prevents clearing fields when only uploading a logo
        company_name = request.POST.get('company_name', '').strip()
        if company_name:
            settings.company_name = company_name
        
        email = request.POST.get('email', '').strip()
        if email:
            settings.email = email
        
        address = request.POST.get('address', '').strip()
        if address:
            settings.address = address
        
        base_url = request.POST.get('base_url', '').strip()
        if base_url:
            settings.base_url = base_url
        
        # Handle logo upload
        if 'logo' in request.FILES:
            # Delete old logo if exists
            if settings.logo:
                settings.logo.delete(save=False)
            settings.logo = request.FILES['logo']
        
        settings.full_clean()  # Validate before saving
        settings.save()
        
        # Return success toast trigger
        response = HttpResponse()
        response['HX-Trigger'] = 'showToast'
        return response
        
    except ValidationError as e:
        error_msg = ', '.join([f"{k}: {', '.join(v)}" for k, v in e.message_dict.items()])
        return HttpResponse(f"Validation error: {error_msg}", status=400)
    except Exception as e:
        logger.exception("Unexpected error while updating global settings.")
        return HttpResponse("An error occurred while updating settings.", status=400)


def public_logo(request):
    """
    Serve the company logo publicly without authentication.
    This endpoint is accessible at /public/logo.png
    """
    settings = GlobalSettings.get_instance()
    
    if not settings.logo:
        # Return a 404 if no logo is set
        raise Http404("No logo configured")
    
    # Serve the logo file
    try:
        from django.http import FileResponse
        import mimetypes
        
        logo_file = settings.logo.open('rb')
        
        # Determine content type from file extension
        content_type, _ = mimetypes.guess_type(settings.logo.name)
        if not content_type:
            # Fallback to generic image type if detection fails
            content_type = 'image/png'
        
        response = FileResponse(logo_file, content_type=content_type)
        response['Cache-Control'] = 'public, max-age=3600'  # Cache for 1 hour
        return response
    except Exception as e:
        logger.error(f"Error serving logo: {e}")
        raise Http404("Logo file not found")


# System Setting Views

@login_required
def system_setting_detail(request):
    """System Setting detail view (singleton)."""
    setting = SystemSetting.get_instance()
    
    context = {
        'setting': setting,
    }
    return render(request, 'system_setting_detail.html', context)


@login_required
@require_http_methods(["POST"])
def system_setting_update(request):
    """Update System Setting."""
    setting = SystemSetting.get_instance()
    
    try:
        # Update fields only if they have a value (not empty string)
        system_name = request.POST.get('system_name', '').strip()
        if system_name:
            setting.system_name = system_name
        
        company = request.POST.get('company', '').strip()
        if company:
            setting.company = company
        
        email = request.POST.get('email', '').strip()
        if email:
            setting.email = email
        
        # Handle logo upload
        if 'company_logo' in request.FILES:
            # Delete old logo if exists
            if setting.company_logo:
                setting.company_logo.delete(save=False)
            setting.company_logo = request.FILES['company_logo']
        
        setting.full_clean()  # Validate before saving
        setting.save()
        
        # Return success toast trigger
        response = HttpResponse()
        response['HX-Trigger'] = 'showToast'
        return response
        
    except ValidationError as e:
        error_msg = ', '.join([f"{k}: {', '.join(v)}" for k, v in e.message_dict.items()])
        return HttpResponse(f"Validation error: {error_msg}", status=400)
    except Exception as e:
        return HttpResponse(f"Error updating settings: {str(e)}", status=400)


# Change Policy Views

@login_required
def change_policies(request):
    """Change policies list view."""
    policies = ChangePolicy.objects.prefetch_related('policy_roles').all()
    
    # Filter by risk level
    risk_level_filter = request.GET.get('risk_level', '').strip()
    if risk_level_filter:
        policies = policies.filter(risk_level=risk_level_filter)
    
    # Filter by security relevant
    security_filter = request.GET.get('security_relevant', '').strip()
    if security_filter == 'true':
        policies = policies.filter(security_relevant=True)
    elif security_filter == 'false':
        policies = policies.filter(security_relevant=False)
    
    # Filter by release type
    release_type_filter = request.GET.get('release_type', '').strip()
    if release_type_filter:
        if release_type_filter == 'null':
            policies = policies.filter(release_type__isnull=True)
        else:
            policies = policies.filter(release_type=release_type_filter)
    
    context = {
        'policies': policies,
        'risk_levels': RiskLevel.choices,
        'release_types': ReleaseType.choices,
        'risk_level_filter': risk_level_filter,
        'security_filter': security_filter,
        'release_type_filter': release_type_filter,
    }
    return render(request, 'change_policies.html', context)


@login_required
def change_policy_create(request):
    """Show create form for new change policy."""
    context = {
        'policy': None,
        'risk_levels': RiskLevel.choices,
        'release_types': ReleaseType.choices,
        'user_roles': UserRole.choices,
    }
    return render(request, 'change_policy_form.html', context)


@login_required
def change_policy_edit(request, id):
    """Show edit form for existing change policy."""
    policy = get_object_or_404(ChangePolicy, id=id)
    
    # Get current roles
    selected_roles = list(policy.policy_roles.values_list('role', flat=True))
    
    context = {
        'policy': policy,
        'risk_levels': RiskLevel.choices,
        'release_types': ReleaseType.choices,
        'user_roles': UserRole.choices,
        'selected_roles': selected_roles,
    }
    return render(request, 'change_policy_form.html', context)


@login_required
@require_http_methods(["POST"])
def change_policy_update(request, id):
    """Create or update a change policy."""
    try:
        # Parse form data
        risk_level = request.POST.get('risk_level', '').strip()
        security_relevant = request.POST.get('security_relevant') == 'on'
        release_type = request.POST.get('release_type', '').strip()
        roles = request.POST.getlist('roles')
        
        # Validate required fields
        if not risk_level:
            return JsonResponse({'success': False, 'error': 'Risk level is required'}, status=400)
        
        # Handle empty release type (should be None in DB)
        if not release_type or release_type == '':
            release_type = None
        
        # Create or update
        if id == 0:
            # Create new policy
            # Check if policy with same criteria already exists
            existing_query = ChangePolicy.objects.filter(
                risk_level=risk_level,
                security_relevant=security_relevant
            )
            if release_type:
                existing_query = existing_query.filter(release_type=release_type)
            else:
                existing_query = existing_query.filter(release_type__isnull=True)
            
            if existing_query.exists():
                return JsonResponse({
                    'success': False,
                    'error': 'A policy with these criteria already exists'
                }, status=400)
            
            policy = ChangePolicy.objects.create(
                risk_level=risk_level,
                security_relevant=security_relevant,
                release_type=release_type
            )
            
            # Log activity
            activity_service = ActivityService()
            activity_service.log(
                verb='change_policy.created',
                target=policy,
                actor=request.user,
                summary=f'Created change policy: {policy}'
            )
            
            message = 'Change policy created successfully'
        else:
            # Update existing policy
            policy = get_object_or_404(ChangePolicy, id=id)
            
            # Check if updating would create duplicate
            existing_query = ChangePolicy.objects.filter(
                risk_level=risk_level,
                security_relevant=security_relevant
            ).exclude(id=id)
            
            if release_type:
                existing_query = existing_query.filter(release_type=release_type)
            else:
                existing_query = existing_query.filter(release_type__isnull=True)
            
            if existing_query.exists():
                return JsonResponse({
                    'success': False,
                    'error': 'A policy with these criteria already exists'
                }, status=400)
            
            policy.risk_level = risk_level
            policy.security_relevant = security_relevant
            policy.release_type = release_type
            policy.save()
            
            # Log activity
            activity_service = ActivityService()
            activity_service.log(
                verb='change_policy.updated',
                target=policy,
                actor=request.user,
                summary=f'Updated change policy: {policy}'
            )
            
            message = 'Change policy updated successfully'
        
        # Update roles
        # Delete existing roles
        policy.policy_roles.all().delete()
        
        # Create new roles
        for role in roles:
            if role:  # Skip empty values
                ChangePolicyRole.objects.create(
                    policy=policy,
                    role=role
                )
        
        return JsonResponse({
            'success': True,
            'message': message,
            'redirect': reverse('change-policies')
        })
        
    except Exception as e:
        logger.error(f"Error updating change policy: {e}", exc_info=True)
        return JsonResponse(
            {
                'success': False,
                'error': 'An unexpected error occurred while updating the change policy.'
            },
            status=500
        )


@login_required
@require_http_methods(["POST"])
def change_policy_delete(request, id):
    """Delete a change policy."""
    policy = get_object_or_404(ChangePolicy, id=id)
    policy_str = str(policy)
    
    try:
        # Log activity before deletion
        activity_service = ActivityService()
        activity_service.log(
            verb='change_policy.deleted',
            target=policy,
            actor=request.user,
            summary=f'Deleted change policy: {policy_str}'
        )
        
        policy.delete()
        
        return JsonResponse({
            'success': True,
            'message': f'Change policy "{policy_str}" deleted successfully',
            'redirect': reverse('change-policies')
        })
        
    except Exception as e:
        logger.error(f"Error deleting change policy: {e}", exc_info=True)
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


# ============================================
# IssueBlueprint Views
# ============================================

@login_required
def blueprints(request):
    """Blueprint list view with django-tables2 and django-filter."""
    from django_tables2 import RequestConfig
    from .tables import IssueBlueprintTable
    from .filters import IssueBlueprintFilter
    from .models import IssueBlueprint
    
    # Get queryset with related objects
    blueprints_qs = IssueBlueprint.objects.select_related('category', 'created_by').all()
    
    # Apply filters
    blueprint_filter = IssueBlueprintFilter(request.GET, queryset=blueprints_qs)
    
    # Create table
    table = IssueBlueprintTable(blueprint_filter.qs)
    RequestConfig(request, paginate={'per_page': 25}).configure(table)
    
    context = {
        'table': table,
        'filter': blueprint_filter,
    }
    return render(request, 'blueprints.html', context)


@login_required
def blueprint_detail(request, id):
    """Blueprint detail view."""
    from .models import IssueBlueprint, Project
    blueprint = get_object_or_404(IssueBlueprint.objects.select_related('category', 'created_by'), id=id)
    
    # Get all projects user has access to for creating issues
    # Get user's organizations
    user_orgs = request.user.user_organisations.values_list('organisation', flat=True)
    projects = Project.objects.filter(
        clients__id__in=user_orgs
    ).distinct().order_by('name')
    
    context = {
        'blueprint': blueprint,
        'projects': projects,
    }
    return render(request, 'blueprint_detail.html', context)



@login_required
def blueprint_create(request):
    """Show create form for new blueprint."""
    from .models import IssueBlueprintCategory, RiskLevel
    
    # Get active categories
    categories = IssueBlueprintCategory.objects.filter(is_active=True).order_by('name')
    
    context = {
        'blueprint': None,
        'categories': categories,
        'risk_levels': RiskLevel.choices,
    }
    return render(request, 'blueprint_form.html', context)


@login_required
@require_http_methods(["POST"])
def blueprint_create_submit(request):
    """Handle POST request for creating a new blueprint."""
    # Delegate to blueprint_update with id='0' to indicate creation
    return blueprint_update(request, id='0')


@login_required
def blueprint_edit(request, id):
    """Show edit form for existing blueprint."""
    from .models import IssueBlueprint, IssueBlueprintCategory, RiskLevel
    
    blueprint = get_object_or_404(IssueBlueprint, id=id)
    
    # Get active categories
    categories = IssueBlueprintCategory.objects.filter(is_active=True).order_by('name')
    
    context = {
        'blueprint': blueprint,
        'categories': categories,
        'risk_levels': RiskLevel.choices,
    }
    return render(request, 'blueprint_form.html', context)


@login_required
@require_http_methods(["POST"])
def blueprint_update(request, id):
    """Create or update a blueprint."""
    from .models import IssueBlueprint, IssueBlueprintCategory
    
    try:
        # Parse form data
        title = request.POST.get('title', '').strip()
        category_id = request.POST.get('category', '').strip()
        description_md = request.POST.get('description_md', '').strip()
        is_active = request.POST.get('is_active') == 'on'
        version = request.POST.get('version', '1').strip()
        tags_str = request.POST.get('tags', '').strip()
        default_labels_str = request.POST.get('default_labels', '').strip()
        default_risk_level = request.POST.get('default_risk_level', '').strip()
        default_security_relevant = request.POST.get('default_security_relevant', '').strip()
        notes = request.POST.get('notes', '').strip()
        action = request.POST.get('action', 'save')
        
        # Validate required fields
        if not title:
            return JsonResponse({'success': False, 'error': 'Title is required'}, status=400)
        if not category_id:
            return JsonResponse({'success': False, 'error': 'Category is required'}, status=400)
        if not description_md:
            return JsonResponse({'success': False, 'error': 'Description is required'}, status=400)
        
        # Validate category exists and is active
        try:
            category = IssueBlueprintCategory.objects.get(id=category_id, is_active=True)
        except IssueBlueprintCategory.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Invalid or inactive category'}, status=400)
        
        # Parse version
        try:
            version = int(version)
            if version < 1:
                version = 1
        except (ValueError, TypeError):
            version = 1
        
        # Parse tags (comma-separated to list)
        tags = []
        if tags_str:
            tags = [t.strip() for t in tags_str.split(',') if t.strip()]
        
        # Parse default_labels (comma-separated to list)
        default_labels = []
        if default_labels_str:
            default_labels = [l.strip() for l in default_labels_str.split(',') if l.strip()]
        
        # Parse default_security_relevant
        default_security_relevant_val = None
        if default_security_relevant == 'true':
            default_security_relevant_val = True
        elif default_security_relevant == 'false':
            default_security_relevant_val = False
        
        # Create or update
        if str(id) == '0':
            # Create new blueprint
            blueprint = IssueBlueprint.objects.create(
                title=title,
                category=category,
                description_md=description_md,
                is_active=is_active,
                version=version,
                tags=tags if tags else None,
                default_labels=default_labels if default_labels else None,
                default_risk_level=default_risk_level if default_risk_level else None,
                default_security_relevant=default_security_relevant_val,
                notes=notes,
                created_by=request.user if request.user.is_authenticated else None
            )
            
            # Log activity
            activity_service = ActivityService()
            activity_service.log(
                verb='blueprint.created',
                target=blueprint,
                actor=request.user if request.user.is_authenticated else None,
                summary=f'Created blueprint "{blueprint.title}"'
            )
            
            message_text = f'Blueprint "{title}" created successfully'
        else:
            # Update existing blueprint
            blueprint = get_object_or_404(IssueBlueprint, id=id)
            
            # Update fields
            blueprint.title = title
            blueprint.category = category
            blueprint.description_md = description_md
            blueprint.is_active = is_active
            blueprint.version = version
            blueprint.tags = tags if tags else None
            blueprint.default_labels = default_labels if default_labels else None
            blueprint.default_risk_level = default_risk_level if default_risk_level else None
            blueprint.default_security_relevant = default_security_relevant_val
            blueprint.notes = notes
            blueprint.save()
            
            # Log activity
            activity_service = ActivityService()
            activity_service.log(
                verb='blueprint.updated',
                target=blueprint,
                actor=request.user if request.user.is_authenticated else None,
                summary=f'Updated blueprint "{blueprint.title}"'
            )
            
            message_text = f'Blueprint "{blueprint.title}" updated successfully'
        
        # Determine redirect based on action
        if action == 'save_close':
            redirect_url = reverse('blueprints')
        else:
            redirect_url = reverse('blueprint-detail', args=[str(blueprint.id)])
        
        return JsonResponse({
            'success': True,
            'message': message_text,
            'redirect': redirect_url,
            'blueprint_id': str(blueprint.id)
        })
        
    except Exception as e:
        logger.error(f"Error saving blueprint: {e}", exc_info=True)
        return JsonResponse(
            {
                'success': False,
                'error': 'An unexpected error occurred while saving the blueprint.'
            },
            status=500
        )


@login_required
@require_http_methods(["POST"])
def blueprint_delete(request, id):
    """Delete a blueprint."""
    from .models import IssueBlueprint
    
    blueprint = get_object_or_404(IssueBlueprint, id=id)
    blueprint_title = blueprint.title
    
    try:
        # Log activity before deletion
        activity_service = ActivityService()
        activity_service.log(
            verb='blueprint.deleted',
            target=blueprint,
            actor=request.user if request.user.is_authenticated else None,
            summary=f'Deleted blueprint "{blueprint_title}"'
        )
        
        blueprint.delete()
        
        return JsonResponse({
            'success': True,
            'message': f'Blueprint "{blueprint_title}" deleted successfully',
            'redirect': reverse('blueprints')
        })
        
    except Exception as e:
        logger.error(f"Error deleting blueprint: {e}", exc_info=True)
        return JsonResponse(
            {
                'success': False,
                'error': 'An error occurred while deleting the blueprint.'
            },
            status=500
        )


@login_required
def blueprint_export(request, id):
    """Export a blueprint as JSON."""
    from .models import IssueBlueprint
    from .utils.blueprint_serializer import export_blueprint_json, BlueprintSerializationError
    
    blueprint = get_object_or_404(IssueBlueprint, id=id)
    
    try:
        json_data = export_blueprint_json(blueprint, indent=2)
        
        # Create filename from blueprint title
        safe_title = re.sub(r'[^\w\s-]', '', blueprint.title).strip().replace(' ', '_')
        filename = f"blueprint_{safe_title}_v{blueprint.version}.json"
        
        response = HttpResponse(json_data, content_type='application/json')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='blueprint.exported',
            target=blueprint,
            actor=request.user,
            summary=f'Exported blueprint "{blueprint.title}"'
        )
        
        return response
        
    except BlueprintSerializationError as e:
        logger.error(f"Error exporting blueprint: {e}", exc_info=True)
        return JsonResponse(
            {
                'success': False,
                'error': 'Failed to export blueprint due to a validation or formatting error.'
            },
            status=400
        )
    except Exception as e:
        logger.error(f"Unexpected error exporting blueprint: {e}", exc_info=True)
        return JsonResponse({'success': False, 'error': 'An unexpected error occurred'}, status=500)


@login_required
def blueprint_import_form(request):
    """Show import form for blueprints."""
    return render(request, 'blueprint_import.html')


@login_required
@require_http_methods(["POST"])
def blueprint_import(request):
    """Import a blueprint from JSON."""
    from .models import IssueBlueprint
    from .utils.blueprint_serializer import import_blueprint_json, BlueprintDeserializationError
    
    try:
        # Get JSON data from file upload or text input
        json_data = None
        
        if 'json_file' in request.FILES:
            # Handle file upload
            json_file = request.FILES['json_file']
            try:
                json_data = json_file.read().decode('utf-8')
            except UnicodeDecodeError:
                return JsonResponse({
                    'success': False, 
                    'error': 'Invalid file encoding. Please upload a UTF-8 encoded JSON file.'
                }, status=400)
        elif 'json_text' in request.POST:
            # Handle text input
            json_data = request.POST.get('json_text', '').strip()
        
        if not json_data:
            return JsonResponse({
                'success': False,
                'error': 'No JSON data provided. Please upload a file or paste JSON text.'
            }, status=400)
        
        # Parse update_if_exists option
        update_if_exists = request.POST.get('update_if_exists') == 'on'
        
        # Import blueprint
        blueprint, created = import_blueprint_json(
            json_data,
            created_by=request.user,
            update_if_exists=update_if_exists
        )
        
        # Log activity
        activity_service = ActivityService()
        verb = 'blueprint.imported' if created else 'blueprint.updated_from_import'
        action = 'Imported' if created else 'Updated from import'
        activity_service.log(
            verb=verb,
            target=blueprint,
            actor=request.user,
            summary=f'{action} blueprint "{blueprint.title}"'
        )
        
        return JsonResponse({
            'success': True,
            'message': f'Blueprint "{blueprint.title}" {"created" if created else "updated"} successfully',
            'blueprint_id': str(blueprint.id),
            'redirect': reverse('blueprint-detail', kwargs={'id': blueprint.id})
        })
        
    except BlueprintDeserializationError as e:
        logger.warning(f"Blueprint import validation error: {e}")
        return JsonResponse({'success': False, 'error': str(e)}, status=400)
    except Exception as e:
        logger.error(f"Unexpected error importing blueprint: {e}", exc_info=True)
        return JsonResponse(
            {
                'success': False,
                'error': 'An unexpected error occurred while importing the blueprint. Please try again later.'
            },
            status=500
        )


@login_required
@require_http_methods(["POST"])
def blueprint_create_issue(request, id):
    """Create a new issue from a blueprint."""
    from .models import IssueBlueprint, Project, Item
    from .utils.blueprint_variables import extract_variables, replace_variables, validate_variables_from_multiple
    import json
    
    blueprint = get_object_or_404(IssueBlueprint, id=id, is_active=True)
    
    try:
        # Parse form data
        project_id = request.POST.get('project_id', '').strip()
        variables_json = request.POST.get('variables', '{}')
        
        # Validate project
        if not project_id:
            return JsonResponse({'success': False, 'error': 'Project is required'}, status=400)
        
        try:
            project = Project.objects.get(id=project_id)
            # Check if user has access to this project (via client relationship)
            user_orgs = request.user.user_organisations.values_list('organisation', flat=True)
            if not project.clients.filter(id__in=user_orgs).exists():
                return JsonResponse({'success': False, 'error': 'No access to this project'}, status=403)
        except Project.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Invalid project'}, status=400)
        
        # Parse variables
        try:
            variables = json.loads(variables_json)
        except json.JSONDecodeError:
            return JsonResponse({'success': False, 'error': 'Invalid variables data'}, status=400)
        
        # Validate all required variables are provided from both title and description
        is_valid, missing_vars = validate_variables_from_multiple(
            [blueprint.title, blueprint.description_md],
            variables
        )
        if not is_valid:
            return JsonResponse({
                'success': False,
                'error': f'Missing required variables: {", ".join(missing_vars)}'
            }, status=400)
        
        # Replace variables in blueprint content
        title = replace_variables(blueprint.title, variables)
        description = replace_variables(blueprint.description_md, variables)
        
        # Get default item type if available
        default_type = None
        if hasattr(project, 'default_item_type') and project.default_item_type:
            default_type = project.default_item_type
        else:
            # Try to get the first active item type
            from .models import ItemType
            default_type = ItemType.objects.filter(is_active=True).first()
        
        # Create the new item
        item = Item.objects.create(
            project=project,
            title=title,
            description=description,
            type=default_type,
        )
        
        # Apply optional blueprint fields if they exist
        if blueprint.default_risk_level:
            item.risk_level = blueprint.default_risk_level
        if blueprint.default_security_relevant is not None:
            item.security_relevant = blueprint.default_security_relevant
        
        item.save()
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='item.created_from_blueprint',
            target=item,
            actor=request.user,
            summary=f'Created issue #{item.id} from blueprint "{blueprint.title}"'
        )
        
        return JsonResponse({
            'success': True,
            'message': f'Issue #{item.id} created successfully from blueprint',
            'redirect': reverse('item-detail', args=[item.id])
        })
        
    except Exception as e:
        logger.error(f"Error creating issue from blueprint: {e}", exc_info=True)
        return JsonResponse(
            {
                'success': False,
                'error': 'An internal error occurred while creating the issue from the blueprint.'
            },
            status=500
        )


@login_required
def item_create_blueprint(request, item_id):
    """Show form to create a blueprint from an existing issue."""
    from .models import IssueBlueprintCategory
    
    item = get_object_or_404(Item, id=item_id)
    
    # Get active categories
    categories = IssueBlueprintCategory.objects.filter(is_active=True).order_by('name')
    
    context = {
        'item': item,
        'categories': categories,
        'risk_levels': RiskLevel.choices,
    }
    return render(request, 'item_create_blueprint_form.html', context)


@login_required
@require_http_methods(["POST"])
def item_create_blueprint_submit(request, item_id):
    """Create a blueprint from an existing issue."""
    from .models import IssueBlueprint, IssueBlueprintCategory
    
    item = get_object_or_404(Item, id=item_id)
    
    try:
        # Parse form data
        title = request.POST.get('title', '').strip()
        category_id = request.POST.get('category', '').strip()
        description_md = request.POST.get('description_md', '').strip()
        is_active = request.POST.get('is_active') == 'on'
        version = request.POST.get('version', '1').strip()
        
        # Validate required fields
        if not title:
            return JsonResponse({'success': False, 'error': 'Title is required'}, status=400)
        if not category_id:
            return JsonResponse({'success': False, 'error': 'Category is required'}, status=400)
        if not description_md:
            return JsonResponse({'success': False, 'error': 'Description is required'}, status=400)
        
        # Validate category exists and is active
        try:
            category = IssueBlueprintCategory.objects.get(id=category_id, is_active=True)
        except IssueBlueprintCategory.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Invalid or inactive category'}, status=400)
        
        # Parse version
        try:
            version = int(version)
            if version < 1:
                version = 1
        except (ValueError, TypeError):
            version = 1
        
        # Create blueprint
        blueprint = IssueBlueprint.objects.create(
            title=title,
            category=category,
            description_md=description_md,
            is_active=is_active,
            version=version,
            created_by=request.user if request.user.is_authenticated else None
        )
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='blueprint.created_from_issue',
            target=blueprint,
            actor=request.user if request.user.is_authenticated else None,
            summary=f'Created blueprint "{blueprint.title}" from issue #{item.id}'
        )
        
        # Also log on the item
        activity_service.log(
            verb='item.blueprint_created',
            target=item,
            actor=request.user if request.user.is_authenticated else None,
            summary=f'Created blueprint "{blueprint.title}"'
        )
        
        return JsonResponse({
            'success': True,
            'message': f'Blueprint "{title}" created successfully from issue #{item.id}',
            'redirect': reverse('blueprint-detail', args=[str(blueprint.id)]),
            'blueprint_id': str(blueprint.id)
        })
        
    except Exception as e:
        logger.error(f"Error creating blueprint from issue: {e}", exc_info=True)
        return JsonResponse(
            {
                'success': False,
                'error': 'An unexpected error occurred while creating the blueprint.'
            },
            status=500
        )


@login_required
@require_http_methods(["POST"])
def blueprint_category_create_inline(request):
    """Create a new blueprint category inline (for use in modals)."""
    from .models import IssueBlueprintCategory
    from django.utils.text import slugify
    
    try:
        # Parse form data
        name = request.POST.get('name', '').strip()
        
        # Validate required fields
        if not name:
            return JsonResponse({'success': False, 'error': 'Category name is required'}, status=400)
        
        # Auto-generate slug from name
        slug = slugify(name)
        
        # Check if slug already exists and make it unique if needed
        original_slug = slug
        counter = 1
        while IssueBlueprintCategory.objects.filter(slug=slug).exists():
            slug = f"{original_slug}-{counter}"
            counter += 1
        
        # Create category
        category = IssueBlueprintCategory.objects.create(
            name=name,
            slug=slug,
            is_active=True
        )
        
        # Log activity (non-critical, don't fail if logging fails)
        try:
            activity_service = ActivityService()
            activity_service.log(
                verb='blueprint_category.created',
                target=category,
                actor=request.user if request.user.is_authenticated else None,
                summary=f'Created blueprint category "{category.name}"'
            )
        except Exception as log_error:
            logger.warning(f"Failed to log activity for category creation: {log_error}")
        
        return JsonResponse({
            'success': True,
            'message': f'Category "{name}" created successfully',
            'category': {
                'id': str(category.id),
                'name': category.name,
                'slug': category.slug
            }
        })
        
    except IntegrityError as e:
        logger.error(f"Integrity error creating category: {e}", exc_info=True)
        # Could be duplicate name or slug, or race condition
        error_msg = 'A category with this name or similar name already exists'
        if 'name' in str(e).lower():
            error_msg = 'A category with this name already exists'
        elif 'slug' in str(e).lower():
            error_msg = 'A category with a similar name already exists'
        return JsonResponse({'success': False, 'error': error_msg}, status=400)
    except Exception as e:
        logger.error(f"Error creating blueprint category: {e}", exc_info=True)
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
def item_apply_blueprint(request, item_id):
    """Show modal to select and apply a blueprint to an issue."""
    from .models import IssueBlueprint
    
    item = get_object_or_404(Item, id=item_id)
    
    # Get all active blueprints
    blueprints = IssueBlueprint.objects.filter(is_active=True).select_related('category').order_by('-updated_at')
    
    context = {
        'item': item,
        'blueprints': blueprints,
    }
    return render(request, 'item_apply_blueprint_modal.html', context)


@login_required
@require_http_methods(["POST"])
def item_apply_blueprint_submit(request, item_id):
    """Apply a blueprint to an issue."""
    from .models import IssueBlueprint
    from .utils.blueprint_variables import replace_variables, validate_variables_from_multiple
    from datetime import datetime
    import json
    
    item = get_object_or_404(Item, id=item_id)
    
    try:
        # Parse form data
        blueprint_id = request.POST.get('blueprint_id', '').strip()
        replace_description = request.POST.get('replace_description') == 'on'
        use_blueprint_title = request.POST.get('use_blueprint_title') == 'on'
        variables_json = request.POST.get('variables', '{}')
        
        # Validate blueprint_id
        if not blueprint_id:
            return JsonResponse({'success': False, 'error': 'Blueprint is required'}, status=400)
        
        # Get blueprint
        try:
            blueprint = IssueBlueprint.objects.get(id=blueprint_id, is_active=True)
        except IssueBlueprint.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Invalid or inactive blueprint'}, status=400)
        
        # Parse and validate variables
        try:
            variables = json.loads(variables_json)
        except json.JSONDecodeError:
            return JsonResponse({'success': False, 'error': 'Invalid variables data'}, status=400)
        
        # Validate all required variables are provided from both title and description
        is_valid, missing_vars = validate_variables_from_multiple(
            [blueprint.title, blueprint.description_md],
            variables
        )
        if not is_valid:
            return JsonResponse({
                'success': False,
                'error': f'Missing required variables: {", ".join(missing_vars)}'
            }, status=400)
        
        # Replace variables in blueprint content
        blueprint_title = replace_variables(blueprint.title, variables)
        blueprint_description = replace_variables(blueprint.description_md, variables)
        
        # Apply blueprint to item
        if use_blueprint_title:
            item.title = blueprint_title
        
        if replace_description:
            # Replace entire description
            item.description = blueprint_description
        else:
            # Append blueprint description
            current_date = datetime.now().strftime('%Y-%m-%d')
            blueprint_header = f"\n\n## Blueprint angewendet: {blueprint.title} ({current_date})\n\n"
            item.description = (item.description or '') + blueprint_header + blueprint_description
        
        item.save()
        
        # Log activity
        activity_service = ActivityService()
        activity_service.log(
            verb='item.blueprint_applied',
            target=item,
            actor=request.user if request.user.is_authenticated else None,
            summary=f'Applied blueprint "{blueprint.title}"'
        )
        
        return JsonResponse({
            'success': True,
            'message': f'Blueprint "{blueprint.title}" applied successfully to issue #{item.id}',
            'redirect': reverse('item-detail', args=[item.id])
        })
        
    except Exception as e:
        logger.error(f"Error applying blueprint to issue: {e}", exc_info=True)
        return JsonResponse(
            {
                'success': False,
                'error': 'An internal error occurred while applying the blueprint.'
            },
            status=500
        )
